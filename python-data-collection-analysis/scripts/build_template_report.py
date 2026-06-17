"""按原报告模板版式生成最终 Word 报告。

正文内容与 `Python数据收集与分析期末考查报告.md` 保持一致：
套用学校模板的封面、表格框和签字栏，在"项目过程"单元格内按《评分标准》
逐项填入"评分要求 → 实现方式 → 运行结果"，并插入月度、星期、Top10 数据表、
三张分析图和评分点对照总览表。
"""

from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


PROJECT_ROOT = Path(__file__).resolve().parents[1]
TEMPLATE_DOCX = (
    PROJECT_ROOT
    / "student-report-template"
    / "期末考查报告模板（学生一人一份，不打印）.docx"
)
OUTPUT_DOCX = PROJECT_ROOT / "Python数据收集与分析期末考查报告.docx"
TABLES_DIR = PROJECT_ROOT / "generated" / "tables"
FIGURES_DIR = PROJECT_ROOT / "generated" / "figures"


def set_run_font(run, size: int = 11, bold: bool = False) -> None:
    """为一个 Word 文本片段设置可读的中文字体。"""
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.bold = bold


def add_paragraph(cell, text: str = "", bold: bool = False) -> None:
    """向表格单元格追加一段统一格式的文字。"""
    paragraph = cell.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    set_run_font(run, bold=bold)


def add_labeled(cell, label: str, text: str) -> None:
    """追加一段"加粗标签 + 普通正文"的说明（评分要求/实现方式/运行结果）。"""
    paragraph = cell.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(4)
    label_run = paragraph.add_run(label)
    set_run_font(label_run, bold=True)
    text_run = paragraph.add_run(text)
    set_run_font(text_run, bold=False)


def add_score_item(
    cell,
    heading: str,
    requirement: str,
    implementation: str,
    result: str,
) -> None:
    """输出一个评分点：小标题 + 评分要求 / 实现方式 / 运行结果。"""
    add_paragraph(cell, heading, bold=True)
    add_labeled(cell, "评分要求：", requirement)
    add_labeled(cell, "实现方式：", implementation)
    add_labeled(cell, "运行结果：", result)


def clear_cell(cell) -> None:
    """清除模板占位文字，同时保留单元格边框。"""
    cell.text = ""


def add_bullets(cell, items: list[str]) -> None:
    """在 Word 表格单元格内添加简单编号列表。"""
    for index, item in enumerate(items, start=1):
        add_paragraph(cell, f"{index}. {item}")


def add_picture(cell, image_path: Path, caption: str) -> None:
    """向报告表格插入一张图表和说明文字。"""
    if not image_path.exists():
        add_paragraph(cell, f"[缺少图表：{image_path.name}]")
        return

    caption_paragraph = cell.add_paragraph()
    caption_run = caption_paragraph.add_run(caption)
    set_run_font(caption_run, bold=True)
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    image_paragraph = cell.add_paragraph()
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_paragraph.add_run().add_picture(str(image_path), width=Inches(5.6))


def _set_table_borders(table) -> None:
    """为单元格内嵌表格添加黑色网格线，确保在 Word/PDF 中可见。"""
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")
        borders.append(element)
    tbl_pr.append(borders)


def _fill_table_cell(cell, text: str, bold: bool = False, align=None) -> None:
    """设置内嵌表格单元格的文字、字体和对齐。"""
    cell.text = ""
    paragraph = cell.paragraphs[0]
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    set_run_font(run, size=10, bold=bold)


def add_data_table(
    cell,
    headers: list[str],
    rows: list[list[str]],
    aligns: list[str] | None = None,
) -> None:
    """在单元格内插入一张带网格线的数据表（表头加粗）。"""
    table = cell.add_table(rows=1, cols=len(headers))
    _set_table_borders(table)

    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        _fill_table_cell(
            header_cells[index], header, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER
        )

    align_map = {"r": WD_ALIGN_PARAGRAPH.RIGHT, "c": WD_ALIGN_PARAGRAPH.CENTER}
    for row in rows:
        row_cells = table.add_row().cells
        for index, value in enumerate(row):
            align = None
            if aligns is not None:
                align = align_map.get(aligns[index])
            _fill_table_cell(row_cells[index], value, align=align)


def _money(value) -> str:
    """金额保留两位小数。"""
    return f"{float(value):.2f}"


def _int(value) -> str:
    """整数列不带小数。"""
    return f"{int(round(float(value)))}"


def load_results() -> dict[str, object]:
    """读取用于填充模板报告的已生成分析结果。"""
    quality = pd.read_json(TABLES_DIR / "data_quality_summary.json", typ="series")
    monthly = pd.read_csv(TABLES_DIR / "monthly_sales_summary.csv")
    weekday = pd.read_csv(TABLES_DIR / "weekday_sales_summary.csv")
    top_products = pd.read_csv(TABLES_DIR / "top10_products_by_quantity.csv")
    return {
        "quality": quality.to_dict(),
        "monthly": monthly,
        "weekday": weekday,
        "top_products": top_products,
    }


def fill_cover(document: Document) -> None:
    """保留原封面格式，并填入中性占位内容。"""
    replacements = {
        "姓    名：": "姓    名： _____________",
        "学    号：": "学    号： ______________",
        "学    院：": "学    院： 信息工程学院",
        "专    业：": "专    业： 软件技术",
        "年    级：": "年    级： ______________",
        "班    级：": "班    级： ______________",
        "指导教师：": "指导教师： ______________",
    }
    for paragraph in document.paragraphs:
        stripped = paragraph.text.strip()
        for prefix, value in replacements.items():
            if stripped.startswith(prefix):
                paragraph.text = value
                for run in paragraph.runs:
                    set_run_font(run, size=12)


def fill_main_table(document: Document, results: dict[str, object]) -> None:
    """按《评分标准》逐项填充项目名称、项目目的和项目过程单元格。"""
    quality = results["quality"]
    monthly = results["monthly"]
    weekday = results["weekday"]
    top_products = results["top_products"]
    overview = quality["overview"]

    row_count = int(overview["row_count"])
    column_count = int(overview["column_count"])
    duplicate_removed = int(quality["duplicate_rows_removed"])
    definite_removed = int(quality["definite_error_rows_removed"])
    outlier_marks = int(quality["statistical_outlier_marks"])
    cleaned_rows = int(quality["cleaned_rows"])
    date_min = quality["date_min"]
    date_max = quality["date_max"]
    monthly_avg = float(quality["monthly_average_actual_amount"])
    best_month = monthly.sort_values("实收金额合计", ascending=False).iloc[0]
    best_weekday = weekday.sort_values("实收金额合计", ascending=False).iloc[0]
    best_product = top_products.sort_values("销售数量合计", ascending=False).iloc[0]

    first_table = document.tables[0]
    project_name_cell = first_table.rows[1].cells[0]
    purpose_cell = first_table.rows[2].cells[0]
    process_cell = first_table.rows[3].cells[0]

    # ---- 项目名称 ----
    clear_cell(project_name_cell)
    add_paragraph(project_name_cell, "项目名称：药品销售数据处理与分析", bold=True)

    # ---- 项目目的 ----
    clear_cell(purpose_cell)
    add_paragraph(purpose_cell, "项目目的：", bold=True)
    add_bullets(
        purpose_cell,
        [
            "掌握使用 Pandas 读取本地 Excel 数据的方法。",
            "能够查看数据结构、字段类型、缺失值和重复值情况。",
            "能够对销售时间、销售数量、应收金额和实收金额等字段进行规范化处理。",
            "能够识别并处理缺失值、无效日期、负数销售等异常数据。",
            "能够使用分组聚合完成月度、星期、商品维度的销售统计。",
            "能够使用 Matplotlib/Seaborn 生成可读的分析图表。",
            "能够将分析过程、结果和代码思路整理成报告和答辩材料。",
        ],
    )

    # ---- 项目内容及要求 + 项目过程（按评分标准逐项说明）----
    clear_cell(process_cell)
    add_paragraph(process_cell, "项目内容及要求：", bold=True)
    add_paragraph(
        process_cell,
        "利用 Python（Pandas、NumPy、Matplotlib、Seaborn）完成药品销售数据的加载、"
        "处理与分析，覆盖评分标准中的数据加载、数据概览、数据处理和数据分析各项要求，"
        "并形成可用于答辩的报告与图表。",
    )

    add_paragraph(process_cell, "项目过程（按评分标准逐项说明）：", bold=True)
    add_paragraph(
        process_cell,
        "下面对照《评分标准》逐项说明，每一项给出评分要求、实现方式和运行结果。",
    )

    # （一）数据加载（5 分）
    add_score_item(
        process_cell,
        "（一）数据加载（5 分）",
        "采用合适的方法把本地文件中的数据加载至 Python 环境。",
        "data_loading.py 的 load_sales_data() 使用 pandas.read_excel() 读取 "
        "exam-materials/data/药品销售数据.xlsx，并在加载后校验必要字段是否齐全，"
        "文件或字段缺失时直接抛出清晰错误。",
        f"成功读取原始数据 {row_count} 行、{column_count} 列。",
    )

    # （二）数据概览（5 分）
    add_paragraph(process_cell, "（二）数据概览（5 分）", bold=True)
    add_score_item(
        process_cell,
        "1. 正确查看数据基本信息（3 分）",
        "正确查看数据基本信息。",
        "data_overview.py 的 inspect_raw_data() 输出行列数、字段列表、字段类型、"
        "缺失值数量和重复值数量。",
        f"确认数据规模为 {row_count} 行 × {column_count} 列，原始重复行数量为 "
        f"{duplicate_removed}；购药时间、社保卡号、星期各缺失 2 条，其余关键字段各缺失 1 条。",
    )
    add_score_item(
        process_cell,
        "2. 修正列名（2 分）",
        "修正列名。",
        "data_overview.py 的 rename_sales_time() 将原始列名“购药时间”修正为“销售时间”，"
        "并保留“源数据行号”便于在异常记录表中追溯原始 Excel 行。",
        "列名由“购药时间”统一为“销售时间”，后续时间相关处理均基于该列。",
    )

    # （三）数据处理（30 分）
    add_paragraph(process_cell, "（三）数据处理（30 分）", bold=True)
    add_score_item(
        process_cell,
        "1. 重复值处理（5 分）",
        "重复值处理。",
        "data_cleaning.py 基于销售时间、社保卡号、商品编码、商品名称、销售数量、"
        "应收金额、实收金额、星期这一组业务字段检测并删除重复记录。",
        f"本次数据中重复记录为 {duplicate_removed} 行，无需删除。",
    )
    add_score_item(
        process_cell,
        "2. 缺失值处理（5 分）",
        "缺失值处理。",
        "对分析必需字段（销售时间、商品名称、销售数量、应收金额、实收金额）缺失的记录"
        "予以剔除；对不影响销售统计的社保卡号缺失记录，用“未知社保卡”标记保留。",
        "必需字段缺失的记录被剔除（计入异常值处理的剔除总量），社保卡号缺失的记录保留并标记。",
    )

    add_paragraph(process_cell, "3. 异常值处理（20 分）", bold=True)
    add_paragraph(
        process_cell,
        "本项 20 分由四个子项构成：销售数据类型统一为整型（5 分）、销售时间类型统一为 "
        "DateTime 类型（5 分）、检测异常值（5 分）、处理异常值（5 分）。其中两项类型统一"
        "是异常值检测的前置步骤——只有先把字段转为正确类型，才能准确判断数量/金额是否为非正值、"
        "日期是否有效。",
    )
    add_score_item(
        process_cell,
        "(1) 销售数据类型统一为整型（5 分）",
        "销售数据类型统一为整型。",
        "data_cleaning.py 将销售数量转换为整型，应收金额、实收金额转换为数值类型，"
        "保证金额可参与求和与均值计算。",
        "销售数量统一为整型，金额字段统一为数值类型。",
    )
    add_score_item(
        process_cell,
        "(2) 销售时间类型统一为 DateTime 类型（5 分）",
        "销售时间类型统一为 DateTime 类型。",
        "data_cleaning.py 使用 pandas.to_datetime() 将销售时间转换为 DateTime 类型；"
        "星期不直接采用原始列，而是根据标准化后的销售时间重新计算，避免 2022-02-29 "
        "这类错误星期值影响统计。",
        "销售时间统一为 DateTime 类型，无效日期可被识别；星期由销售时间重算。",
    )
    add_score_item(
        process_cell,
        "(3) 检测异常值（5 分）",
        "检测异常值。",
        "分两类检测：确定错误（无效日期、关键字段缺失、销售数量或金额小于等于 0）；"
        "统计异常（使用 IQR 四分位距方法标记大额或大数量记录）。",
        f"IQR 统计异常共标记 {outlier_marks} 条，已输出到 "
        "generated/tables/anomaly_rows.csv 供复核。",
    )
    add_score_item(
        process_cell,
        "(4) 处理异常值（5 分）",
        "处理异常值。",
        "确定错误记录予以剔除；IQR 统计异常的大额/大数量记录保留复核，因为医药销售中"
        "可能存在真实批量购买，不能只因数值大就删除。",
        f"最终剔除确定错误记录 {definite_removed} 行，清洗后保留 {cleaned_rows} 行有效数据；"
        f"统计异常的 {outlier_marks} 条不直接删除，而是输出异常表供复核。",
    )

    # （四）数据分析（30 分）
    add_paragraph(process_cell, "（四）数据分析（30 分）", bold=True)
    add_score_item(
        process_cell,
        "1. 将数据按销售时间排序并重置行索引（5 分）",
        "将数据按照销售时间进行排序，排序之后重置行索引。",
        "data_cleaning.py 将清洗后的数据按销售时间从小到大排序，并重置行索引（reset_index）。",
        f"数据按销售时间升序排列、行索引连续，分析时间范围为 {date_min} 至 {date_max}。",
    )

    add_paragraph(process_cell, "2. 计算月均销售金额（5 分）", bold=True)
    add_labeled(process_cell, "评分要求：", "计算月均销售金额。")
    add_labeled(
        process_cell,
        "实现方式：",
        "monthly_sales.py 按月份对实收金额进行分组聚合，并计算月均实收金额。",
    )
    add_labeled(
        process_cell,
        "运行结果：",
        f"各月份销售统计如下表，月均实收金额为 {_money(monthly_avg)} 元，"
        f"其中 {best_month['销售月份']} 的实收金额最高，为 {_money(best_month['实收金额合计'])} 元。",
    )
    add_data_table(
        process_cell,
        ["月份", "订单数", "销售数量", "实收金额"],
        [
            [
                str(row["销售月份"]),
                _int(row["订单数"]),
                _int(row["销售数量合计"]),
                _money(row["实收金额合计"]),
            ]
            for _, row in monthly.iterrows()
        ],
        aligns=["l", "r", "r", "r"],
    )

    add_paragraph(process_cell, "3. 绘制销售时间与实收金额的关系（5 分）", bold=True)
    add_labeled(process_cell, "评分要求：", "绘制销售时间与实收金额的关系。")
    add_labeled(
        process_cell,
        "实现方式：",
        "time_actual_relationship.py 汇总每日实收金额并生成销售时间与实收金额关系图。",
    )
    add_labeled(
        process_cell,
        "运行结果：",
        "每日实收金额波动明显，部分日期存在高峰；高峰通常与大额或批量购药记录有关，"
        "因此在清洗中作为统计异常保留复核，未直接删除。",
    )
    add_picture(
        process_cell,
        FIGURES_DIR / "sales_time_actual_amount.png",
        "图 1 销售时间与每日实收金额关系",
    )

    add_paragraph(process_cell, "4. 根据星期分组统计销售数量、应收和实收金额（5 分）", bold=True)
    add_labeled(
        process_cell, "评分要求：", "根据星期分组，统计每星期销售数量、应收和实收金额的关系。"
    )
    add_labeled(
        process_cell,
        "实现方式：",
        "weekday_sales.py 按星期（由销售时间重算）分组，统计订单数、销售数量、应收金额和"
        "实收金额，并生成星期统计图。",
    )
    add_labeled(
        process_cell,
        "运行结果：",
        f"各星期销售统计如下表，{best_weekday['星期']}的订单数、销售数量、应收金额和"
        "实收金额均较高，是一周中销售表现最突出的日期。",
    )
    add_data_table(
        process_cell,
        ["星期", "订单数", "销售数量", "应收金额", "实收金额"],
        [
            [
                str(row["星期"]),
                _int(row["订单数"]),
                _int(row["销售数量合计"]),
                _money(row["应收金额合计"]),
                _money(row["实收金额合计"]),
            ]
            for _, row in weekday.iterrows()
        ],
        aligns=["l", "r", "r", "r", "r"],
    )
    add_picture(
        process_cell,
        FIGURES_DIR / "weekday_sales_summary.png",
        "图 2 星期分组销售统计",
    )

    add_paragraph(process_cell, "5. 销售数量前十位药品（10 分）", bold=True)
    add_labeled(
        process_cell,
        "评分要求：",
        "根据商品名称分组，统计销售数量前十位的药品，绘制销售数量前十位的药品名称的销售数量关系。",
    )
    add_labeled(
        process_cell,
        "实现方式：",
        "top_products.py 按商品名称分组统计销售数量，取前十位药品并生成 Top 10 图表。",
    )
    add_labeled(
        process_cell,
        "运行结果：",
        f"销售数量前十位药品如下表，销售数量最高的是 {best_product['商品名称']}"
        f"（销售数量 {_int(best_product['销售数量合计'])}）；"
        "不同药品的价格差异会影响实收金额排名。",
    )
    add_data_table(
        process_cell,
        ["排名", "商品名称", "订单数", "销售数量", "实收金额"],
        [
            [
                str(rank),
                str(row["商品名称"]),
                _int(row["订单数"]),
                _int(row["销售数量合计"]),
                _money(row["实收金额合计"]),
            ]
            for rank, (_, row) in enumerate(top_products.iterrows(), start=1)
        ],
        aligns=["c", "l", "r", "r", "r"],
    )
    add_picture(
        process_cell,
        FIGURES_DIR / "top10_products_by_quantity.png",
        "图 3 销售数量前十位药品",
    )

    # （五）项目汇报（30 分）
    add_paragraph(process_cell, "（五）项目汇报（30 分）", bold=True)
    add_labeled(
        process_cell,
        "评分要求：",
        "根据程序实施的思路完成对项目的汇报，每人汇报 3 分钟，内容主要包括代码完成思路"
        "及相关技术的应用和实现。",
    )
    add_labeled(
        process_cell,
        "实现方式：",
        "汇报材料整理为 答辩.md（3 分钟口头汇报稿）、docs/defense-outline.md（答辩提纲）"
        "和 docs/defense-qa.md（预设问答）。",
    )
    add_labeled(
        process_cell,
        "汇报主线：",
        "按“数据加载 → 数据概览与列名修正 → 数据处理（重复值、缺失值、类型统一、异常值"
        "检测与处理）→ 数据分析（排序、月均、时间关系、星期分组、Top 10）→ 可视化”讲解；"
        "技术上重点说明 Pandas 的 read_excel/to_datetime/分组聚合、NumPy 的 IQR 计算、"
        "Matplotlib/Seaborn 的中文图表配置，其中异常值处理策略是讲解重点。",
    )

    # 评分点对照总览表
    add_paragraph(process_cell, "评分点对照总览表：", bold=True)
    add_data_table(
        process_cell,
        ["评分项", "分值", "对应实现（模块/文件）", "报告位置"],
        [
            ["数据加载", "5", "data_loading.py load_sales_data()", "（一）"],
            ["数据概览 — 正确查看数据基本信息", "3", "data_overview.py inspect_raw_data()", "（二）1"],
            ["数据概览 — 修正列名", "2", "data_overview.py rename_sales_time()", "（二）2"],
            ["数据处理 — 重复值处理", "5", "data_cleaning.py", "（三）1"],
            ["数据处理 — 缺失值处理", "5", "data_cleaning.py", "（三）2"],
            ["异常值处理 — 销售数据类型统一为整型", "5", "data_cleaning.py", "（三）3(1)"],
            ["异常值处理 — 销售时间统一为 DateTime", "5", "data_cleaning.py", "（三）3(2)"],
            ["异常值处理 — 检测异常值", "5", "data_cleaning.py（IQR）", "（三）3(3)"],
            ["异常值处理 — 处理异常值", "5", "data_cleaning.py", "（三）3(4)"],
            ["数据分析 — 按销售时间排序并重置索引", "5", "data_cleaning.py", "（四）1"],
            ["数据分析 — 计算月均销售金额", "5", "monthly_sales.py", "（四）2"],
            ["数据分析 — 销售时间与实收金额关系", "5", "time_actual_relationship.py", "（四）3"],
            ["数据分析 — 星期分组统计", "5", "weekday_sales.py", "（四）4"],
            ["数据分析 — 销售数量前十位药品", "10", "top_products.py", "（四）5"],
            ["项目汇报", "30", "答辩.md、docs/defense-outline.md、docs/defense-qa.md", "（五）"],
            ["总分", "100", "—", "—"],
        ],
        aligns=["l", "c", "l", "c"],
    )

    # 单元格内最后一个元素若为表格，补一个空段落，避免 Word 文档结构异常。
    process_cell.add_paragraph()


def fill_summary_table(document: Document) -> None:
    """保留签字行，并填充个人总结部分。"""
    second_table = document.tables[1]
    summary_cell = second_table.rows[0].cells[0]
    clear_cell(summary_cell)
    add_paragraph(summary_cell, "个人总结：", bold=True)
    add_paragraph(
        summary_cell,
        "通过本次项目，我完整实践了从数据读取、数据概览、数据清洗、异常处理、分组统计"
        "到图表生成的分析流程。项目中最需要注意的是异常值处理不能简单“一删了之”："
        "负数销售、无效日期和关键字段缺失属于确定错误，应当剔除；但大额销售或大数量销售"
        "可能是真实业务现象，因此更适合作为统计异常保留复核。",
    )
    add_paragraph(
        summary_cell,
        "本项目最终形成了可运行代码、清洗数据、统计表、图表、报告和答辩文档，"
        "能够逐项对应《Python数据收集与分析》期末考查的评分标准。",
    )


def build_report() -> None:
    """生成套用原模板格式的最终报告。"""
    if not TEMPLATE_DOCX.exists():
        raise FileNotFoundError(
            "缺少转换后的模板 docx，请先用 LibreOffice 将原 .doc 模板转换为 .docx："
            f"{TEMPLATE_DOCX}"
        )

    document = Document(TEMPLATE_DOCX)
    results = load_results()
    fill_cover(document)
    fill_main_table(document, results)
    fill_summary_table(document)
    document.save(OUTPUT_DOCX)
    print(f"模板版报告已生成：{OUTPUT_DOCX}")


if __name__ == "__main__":
    build_report()
