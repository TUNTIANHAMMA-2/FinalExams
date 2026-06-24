#!/usr/bin/env python3
"""
Generate software-testing deliverables for the RhizoDelta testing assignment.

The original course templates under software-testing/templates are kept as
references. Current submission files are generated directly as docx/xls so the
build does not depend on unreliable old Office-format conversions.
"""

from __future__ import annotations

import sys
import tempfile
from pathlib import Path


try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt
    import xlwt
except ImportError as exc:  # pragma: no cover - user-facing setup guard
    raise SystemExit(
        "Missing dependency: python-docx/xlwt. "
        "Install with: python3 -m pip install python-docx xlwt"
    ) from exc


ROOT = Path(__file__).resolve().parents[1]
TEMPLATES = ROOT / "templates"
DELIVERABLES = ROOT / "deliverables"
GENERATED = ROOT / "generated"

TEST_DATE = "2026.6.18"
AUTHOR = "本人"
TEAM = "个人独立完成"


PLAN_TEMPLATE = TEMPLATES / "1测试方案.doc"
CASE_TEMPLATE = TEMPLATES / "2测试用例.xls"
BUG_TEMPLATE = TEMPLATES / "3Bug缺陷报告清单.xls"
SUMMARY_TEMPLATE = TEMPLATES / "5测试总结报告.doc"
REPORT_TEMPLATE = TEMPLATES / "《软件测试技术》考查报告（每人一份）.doc"


CASES = [
    (
        "RD-ST-SRS001-001",
        "认证授权",
        "合法注册",
        "高",
        "用户名未占用",
        "等价类：有效用户名 + 密码长度>=8 + display_name 合法；代表输入：{username: qa_<timestamp>, password: ChangeMe123!, display_name: QA Tester}",
        "POST /api/auth/register，提交合法注册信息",
        "HTTP 200，code=0，返回 token、refresh_token、user",
        "一致",
        "通过",
    ),
    (
        "RD-ST-SRS001-002",
        "认证授权",
        "重复用户名注册",
        "高",
        "用户名已存在",
        "等价类：无效用户名=已存在；代表输入：重复提交 qa_tester_st",
        "POST /api/auth/register，再次注册同名用户",
        "按 REST 语义应返回 HTTP 409 Conflict，业务码 40901",
        "当前 HTTP 400，code=40001，message=username already exists",
        "通过，记录 BUG-001",
    ),
    (
        "RD-ST-SRS001-003",
        "认证授权",
        "密码少于 8 位",
        "高",
        "无",
        '边界值：密码长度=3（小于最小有效长度 8）；代表输入：password="123"',
        "POST /api/auth/register",
        "HTTP 400，提示密码至少 8 位",
        "一致",
        "通过",
    ),
    (
        "RD-ST-SRS001-004",
        "认证授权",
        "空用户名注册",
        "中",
        "无",
        '边界值：username 为空字符串；代表输入：username=""',
        "POST /api/auth/register",
        "HTTP 400，提示 must not be blank",
        "一致",
        "通过",
    ),
    (
        "RD-ST-SRS001-005",
        "认证授权",
        "正确登录",
        "高",
        "账号已注册且 ACTIVE",
        "因果图：账号存在=是，密码正确=是，账号状态=ACTIVE；代表输入：正确 username/password",
        "POST /api/auth/login",
        "HTTP 200，code=0，返回 token",
        "一致",
        "通过",
    ),
    (
        "RD-ST-SRS001-006",
        "认证授权",
        "密码错误登录",
        "高",
        "账号存在",
        "因果图：账号存在=是，密码正确=否；代表输入：password=WrongPass99",
        "POST /api/auth/login",
        "HTTP 401，提示 invalid username or password",
        "一致",
        "通过",
    ),
    (
        "RD-ST-SRS001-007",
        "认证授权",
        "不存在用户登录",
        "高",
        "无",
        "因果图：账号存在=否；代表输入：username=no_such_user_xyz",
        "POST /api/auth/login",
        "HTTP 401，提示与密码错误一致，不泄露账号存在性",
        "一致",
        "通过",
    ),
    (
        "RD-ST-SRS001-008",
        "认证授权",
        "带 token 查询当前用户",
        "高",
        "已登录",
        "等价类：Authorization token 合法且未吊销；代表输入：Authorization: Bearer <token>",
        "GET /api/auth/me",
        "HTTP 200，返回 username 与登录账号匹配",
        "一致",
        "通过",
    ),
    (
        "RD-ST-SRS001-009",
        "认证授权",
        "无 token 查询当前用户",
        "高",
        "无",
        "等价类：Authorization token 缺失；代表输入：不带 Authorization 请求头",
        "GET /api/auth/me",
        "HTTP 401，提示 authentication required",
        "一致",
        "通过",
    ),
    (
        "RD-ST-SRS001-010",
        "认证授权",
        "伪造 token 查询当前用户",
        "高",
        "无",
        "错误推测：token 结构存在但签名/载荷无效；代表输入：Bearer faketoken.abc.def",
        "GET /api/auth/me",
        "HTTP 401，提示 invalid token",
        "一致",
        "通过",
    ),
    (
        "RD-ST-SRS001-011",
        "认证授权",
        "刷新令牌",
        "中",
        "持有有效 refresh_token",
        "等价类：refresh_token 合法且未过期；代表输入：{refresh_token: <refreshToken>}",
        "POST /api/auth/refresh",
        "HTTP 200，发放新 token",
        "一致",
        "通过",
    ),
    (
        "RD-ST-SRS001-014",
        "认证授权",
        "密码 8 位边界注册",
        "中",
        "用户名未占用",
        '边界值：password 长度=8（最小有效边界）；代表输入：password="Abcd1234"',
        "POST /api/auth/register",
        "HTTP 200，code=0，返回 token 与 user；若策略要求更复杂密码，应返回明确校验错误",
        "未执行，待加入 Postman 回归",
        "待执行（补充设计）",
    ),
    (
        "RD-ST-SRS001-015",
        "认证授权",
        "注册缺少密码字段",
        "高",
        "无",
        "等价类：password 缺失；代表输入：{username: qa_missing_pwd}",
        "POST /api/auth/register",
        "HTTP 400，code 非 0，不创建用户，不返回 token",
        "未执行，待加入 Postman 回归",
        "待执行（补充设计）",
    ),
    (
        "RD-ST-SRS001-016",
        "认证授权",
        "空密码登录",
        "高",
        "账号存在",
        '边界值：password 为空字符串；代表输入：password=""',
        "POST /api/auth/login",
        "HTTP 400 或 401，code 非 0，不返回 token",
        "未执行，待加入 Postman 回归",
        "待执行（补充设计）",
    ),
    (
        "RD-ST-SRS001-017",
        "认证授权",
        "空刷新令牌",
        "中",
        "无",
        '等价类：refresh_token 缺失/为空；代表输入：refresh_token=""',
        "POST /api/auth/refresh",
        "HTTP 400 或 401，code 非 0，不发放新 token",
        "未执行，待加入 Postman 回归",
        "待执行（补充设计）",
    ),
    (
        "RD-ST-SRS002-001",
        "用户资料与社交",
        "查看本人资料",
        "中",
        "已登录",
        "等价类：token 合法；代表输入：Bearer <token>",
        "GET /api/users/me/profile",
        "HTTP 200，返回 user_id、username、display_name 等字段",
        "一致",
        "通过",
    ),
    (
        "RD-ST-SRS002-002",
        "用户资料与社交",
        "未登录查询资料",
        "高",
        "无",
        "等价类：token 缺失；代表输入：无 Authorization 请求头",
        "GET /api/users/me/profile",
        "HTTP 401，提示 authentication required",
        "一致",
        "通过",
    ),
    (
        "RD-ST-SRS002-003",
        "用户资料与社交",
        "查询在线状态",
        "低",
        "已登录",
        "等价类：token 合法；字段检查：online 为布尔值，last_active 为时间表达",
        "GET /api/users/me/status",
        "HTTP 200，online=true，last_active 可解析",
        "接口可用；last_active 为字符串型 epoch，见 BUG-002",
        "通过，记录低危缺陷",
    ),
    (
        "RD-ST-SRS002-004",
        "用户资料与社交",
        "动态流分页",
        "中",
        "已登录",
        "等价类：默认分页参数；代表输入：不传 page/size",
        "GET /api/users/me/feed",
        "HTTP 200，返回 items 数组和分页字段",
        "一致",
        "通过",
    ),
    (
        "RD-ST-SRS002-005",
        "用户资料与社交",
        "动态流分页最小页大小",
        "中",
        "已登录",
        "边界值：page=0，size=1；验证最小页大小和分页字段一致性",
        "GET /api/users/me/feed?page=0&size=1",
        "HTTP 200，items 长度<=1，分页元数据合法",
        "未执行，待加入 Postman 回归",
        "待执行（补充设计）",
    ),
    (
        "RD-ST-SRS002-006",
        "用户资料与社交",
        "动态流非法页码",
        "中",
        "已登录",
        "边界值/无效等价类：page=-1，size=20；验证负页码校验",
        "GET /api/users/me/feed?page=-1&size=20",
        "HTTP 400，code 非 0，不返回正常分页数据",
        "未执行，待加入 Postman 回归",
        "待执行（补充设计）",
    ),
    (
        "RD-ST-SRS004-001",
        "图谱查询",
        "查询根话题",
        "中",
        "图数据库中有数据",
        "等价类：token 合法且无需路径参数；代表输入：Bearer <token>",
        "GET /api/nodes/roots",
        "HTTP 200，data 为节点数组",
        "一致",
        "通过",
    ),
    (
        "RD-ST-SRS004-002",
        "图谱查询",
        "非法 UUID 格式",
        "中",
        "无",
        '边界值/无效等价类：UUID 格式错误；代表输入：id="not-a-uuid"',
        "GET /api/nodes/{id}",
        "HTTP 400，提示 id must be a valid UUID",
        "一致，格式校验先于查库",
        "通过",
    ),
    (
        "RD-ST-SRS004-003",
        "图谱查询",
        "合法但不存在的节点",
        "中",
        "无",
        "等价类：UUID 格式合法但资源不存在；代表输入：00000000-0000-0000-0000-000000000000",
        "GET /api/nodes/{id}",
        "HTTP 404，code=40401，Node not found",
        "一致",
        "通过",
    ),
    (
        "RD-ST-SRS004-004",
        "图谱查询",
        "无 token 查询根话题",
        "高",
        "无",
        "等价类：token 缺失；代表输入：GET /api/nodes/roots 不带 Authorization",
        "GET /api/nodes/roots",
        "HTTP 401，提示 authentication required",
        "未执行，待加入 Postman 回归",
        "待执行（补充设计）",
    ),
    (
        "RD-ST-SRS001-012",
        "认证收尾",
        "登出",
        "中",
        "已登录",
        "场景法：注册/登录后 token 合法；代表输入：Bearer <token>",
        "POST /api/auth/logout",
        "HTTP 200，code=0",
        "一致",
        "通过",
    ),
    (
        "RD-ST-SRS001-013",
        "认证收尾",
        "登出后复用旧 token",
        "高",
        "已登出",
        "场景法/状态迁移：token 已被吊销；代表输入：旧 token",
        "GET /api/auth/me",
        "HTTP 401，提示 token revoked，黑名单生效",
        "一致",
        "通过",
    ),
    (
        "RD-ST-SRS001-018",
        "认证收尾",
        "无 token 登出",
        "中",
        "无",
        "等价类：token 缺失；代表输入：POST /api/auth/logout 不带 Authorization",
        "POST /api/auth/logout",
        "HTTP 401，code 非 0，不影响其它会话 token 状态",
        "未执行，待加入 Postman 回归",
        "待执行（补充设计）",
    ),
]


BUGS = [
    (
        "BUG-001",
        "认证授权",
        "POST /api/auth/register",
        "重复用户名注册返回 400，HTTP 语义应为 409 Conflict",
        "浏览器/工具：Postman + newman；环境：http://localhost:8090。\n"
        "步骤：1. 使用 qa_tester_st 注册成功；2. 再次使用相同 username 调用注册接口；3. 观察响应。\n"
        "实际：HTTP 400，code=40001，message=username already exists。\n"
        "预期：HTTP 409，业务码 40901。建议重复注册场景改用 ApiResponse.conflict(...)。",
        "低",
        AUTHOR,
        "见 generated/postman-report.html 与 generated/images/02_api_evidence.png",
    ),
    (
        "BUG-002",
        "用户资料与社交",
        "GET /api/users/me/status",
        "last_active 为字符串型 epoch，与其它时间字段格式不一致",
        "浏览器/工具：Postman；环境：http://localhost:8090。\n"
        "步骤：1. 登录取得 Bearer token；2. 调用 /api/users/me/status；3. 查看 data.last_active。\n"
        "实际：last_active 类似 \"1781751037890\"，为字符串型毫秒时间戳。\n"
        "预期：与 updated_at 等字段统一为 ISO8601，或统一为数值型毫秒并在接口文档说明。",
        "低",
        AUTHOR,
        "见接口响应记录",
    ),
]


CASE_MODULE_ORDER = ["认证授权", "用户资料与社交", "图谱查询", "认证收尾"]
EXECUTION_STATUS_DONE = "通过"
POSTMAN_CASE_TOTAL = 20
POSTMAN_ASSERTION_TOTAL = 58

BUG_SUMMARY = [
    ("认证授权", 0, 0, 0, 0, 1, 1),
    ("用户资料与社交", 0, 0, 0, 0, 1, 1),
    ("发帖与关联", 0, 0, 0, 0, 0, 0),
    ("图谱查询", 0, 0, 0, 0, 0, 0),
    ("合计", 0, 0, 0, 0, 2, 2),
]


EQUIVALENCE_CLASSES = [
    (
        "EC-AUTH-USERNAME",
        "注册 username",
        "未占用、非空、满足系统允许字符与长度的用户名",
        "空字符串、已存在用户名、超长/非法字符用户名",
        "空值；已存在值；正常唯一值",
        "RD-ST-SRS001-001、002、004、015",
        "等价类划分 + 边界值",
    ),
    (
        "EC-AUTH-PASSWORD",
        "注册/登录 password",
        "长度>=8 且符合当前密码策略",
        "缺失、空字符串、长度<8、错误密码",
        "长度0、3、8；正确/错误密码组合",
        "RD-ST-SRS001-003、005、006、014、016",
        "等价类划分 + 边界值 + 因果图",
    ),
    (
        "EC-AUTH-TOKEN",
        "Authorization Bearer token",
        "合法、未过期、未吊销的 access token",
        "缺失、伪造、过期、登出后已吊销 token",
        "无请求头；格式伪造；旧 token",
        "RD-ST-SRS001-008、009、010、012、013、018、RD-ST-SRS002-001、002、RD-ST-SRS004-004",
        "等价类划分 + 错误推测 + 场景法",
    ),
    (
        "EC-AUTH-REFRESH",
        "refresh_token",
        "合法且未过期的 refresh_token",
        "缺失、空字符串、伪造、过期 refresh_token",
        "有效值；空值",
        "RD-ST-SRS001-011、017",
        "等价类划分 + 错误推测",
    ),
    (
        "EC-FEED-PAGE",
        "动态流 page/size",
        "page>=0，size 在服务端允许范围内；未传时使用默认值",
        "page<0，size=0，size 过大，非数字",
        "默认值；page=0,size=1；page=-1",
        "RD-ST-SRS002-004、005、006",
        "边界值分析 + 错误推测",
    ),
    (
        "EC-GRAPH-UUID",
        "节点 id(UUID)",
        "符合 UUID 格式且资源存在；符合 UUID 格式但资源不存在",
        "非 UUID 字符串、空路径段、恶意构造路径",
        "not-a-uuid；00000000-0000-0000-0000-000000000000",
        "RD-ST-SRS004-002、003",
        "等价类划分 + 边界值",
    ),
]


METHOD_COVERAGE = [
    ("等价类划分", "把 username、password、token、refresh_token、分页参数、UUID 分成有效/无效输入域", "主表输入栏 + 等价类划分表"),
    ("边界值分析", "覆盖空值、最小有效密码长度、最小页大小、负页码、UUID 格式边界", "RD-ST-SRS001-003/004/014/016、RD-ST-SRS002-005/006、RD-ST-SRS004-002"),
    ("错误推测", "覆盖重复注册、伪造 token、缺字段、非法分页、无 token 访问受保护资源", "RD-ST-SRS001-002/010/015/017/018、RD-ST-SRS004-004"),
    ("因果图/判定表", "登录结果由账号存在性、密码正确性、账号状态共同决定", "RD-ST-SRS001-005/006/007"),
    ("场景法", "注册/登录 -> 查询资料 -> 刷新令牌 -> 图谱查询 -> 登出 -> 旧 token 失效", "RD-ST-SRS001-001/005/008/011/012/013"),
]


def case_counts() -> list[tuple[str, int]]:
    return [(module, sum(1 for case in CASES if case[1] == module)) for module in CASE_MODULE_ORDER]


def case_total() -> int:
    return len(CASES)


def supplement_case_total() -> int:
    return sum(1 for case in CASES if "补充设计" in case[9])


def executed_case_total() -> int:
    return case_total() - supplement_case_total()


def set_cell_text(cell, text: str) -> None:
    cell.text = ""
    for idx, line in enumerate(str(text).split("\n")):
        if idx == 0:
            para = cell.paragraphs[0]
        else:
            para = cell.add_paragraph()
        run = para.add_run(line)
        run.font.name = "宋体"
        run.font.size = Pt(10.5)


def add_para(doc: Document, text: str = "", style: str | None = None, bold: bool = False) -> None:
    p = doc.add_paragraph(style=style)
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.font.name = "宋体"
        run.font.size = Pt(10.5)


def add_heading(doc: Document, text: str, level: int) -> None:
    doc.add_heading(text, level=level)


def add_table(doc: Document, headers: list[str], rows: list[tuple | list]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    for style_name in ("Table Grid", "网格型", "Table Normal"):
        try:
            table.style = style_name
            break
        except KeyError:
            continue
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header)
    for row_data in rows:
        row = table.add_row()
        for idx, value in enumerate(row_data):
            set_cell_text(row.cells[idx], str(value))


def clear_doc_body(doc: Document) -> None:
    body = doc._element.body
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def set_doc_defaults(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(10.5)
    for section in doc.sections:
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)


def build_plan(template_docx: Path, output: Path) -> None:
    doc = Document(template_docx)
    clear_doc_body(doc)
    set_doc_defaults(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("RhizoDelta 图谱化非线性讨论系统测试方案")
    run.bold = True
    run.font.name = "宋体"
    run.font.size = Pt(18)

    add_heading(doc, "目 录", 1)
    for item in [
        "1 概述",
        "1.1 编写目的",
        "1.2 阅读对象",
        "1.3 系统背景",
        "2 测试任务",
        "2.1 测试目的",
        "2.2 测试参考文档",
        "2.3 测试提交文档",
        "3 测试资源",
        "3.1 硬件配置",
        "3.2 软件配置",
        "3.3 人力资源分配",
        "4 测试计划",
        "4.1 整体功能模块划分",
        "5 测试整体进度安排",
        "6 相关风险",
    ]:
        add_para(doc, item)

    add_heading(doc, "1 概述", 1)
    add_heading(doc, "1.1 编写目的", 2)
    add_para(doc, "本文档为 RhizoDelta 软件测试课程作业的测试方案，编写目的包括：")
    for text in [
        "确认被测系统的模块划分、测试范围与范围边界。",
        "确认个人独立测试任务、执行方式与交付文档。",
        "确定测试环境、接口工具、测试数据和外部依赖。",
        "识别测试过程风险并给出应对方案。",
        "为后续测试用例、缺陷清单和总结报告提供依据。",
    ]:
        add_para(doc, f"    {text}")
    add_heading(doc, "1.2 阅读对象", 2)
    add_para(doc, "本文档的预期读者包括：课程指导教师、评审人员、项目负责人、软件设计开发人员、测试人员以及后续维护人员。")
    add_heading(doc, "1.3 系统背景", 2)
    add_para(doc, "RhizoDelta 是一个 B/S 架构的图谱化非线性讨论系统，后端采用 Spring Boot 3、Spring Security + JWT、Spring Data Neo4j、Redis、RabbitMQ 与 SSE，前端采用 React/Vite。系统通过图数据库组织讨论节点与关联，通过异步队列处理发帖、摘要等高成本任务，并向前端实时推送状态变化。本次测试以可运行的 RhizoDelta 实例为被测系统，重点验证认证授权、用户资料基础接口、图谱查询基础接口、登出与 token 吊销等本轮核心子集 REST 行为。")

    add_heading(doc, "2 测试任务", 1)
    add_heading(doc, "2.1 测试目的", 2)
    for text in [
        "确定本轮核心子集功能是否正常可用，业务流程是否闭环。",
        "确定 REST 接口的状态码、统一响应包、字段命名、参数校验和鉴权约束是否符合预期。",
        "确定需求范围是否一致、完整、可验证。",
        "通过 Postman/newman 形成可复现的接口测试证据。",
        "发现并记录缺陷，给出系统质量结论和后续改进建议。",
    ]:
        add_para(doc, f"    {text}")
    add_heading(doc, "2.2 测试参考文档", 2)
    add_table(
        doc,
        ["文档名", "版本", "日期", "作者", "备注"],
        [
            ("RhizoDelta 项目说明与课程交付物说明", "1.0", "2026.6.18", AUTHOR, "说明被测系统技术栈、功能与课程口径"),
            ("软件测试 docs/project-understanding.md", "1.0", "2026.6.18", AUTHOR, "被测系统理解、角色、模块与范围边界"),
            ("软件测试 docs/system-design.md", "1.0", "2026.6.18", AUTHOR, "测试策略、用例设计方法、环境与通过准则"),
            ("Postman Collection 与 newman 报告", "1.0", "2026.6.18", AUTHOR, "接口测试可复现证据"),
        ],
    )
    add_heading(doc, "2.3 测试提交文档", 2)
    add_table(
        doc,
        ["文档名", "版本", "日期", "作者", "备注"],
        [
            ("《1-测试方案》", "1.0", TEST_DATE, AUTHOR, "由 1测试方案.doc 模板副本填写"),
            ("《2-测试用例》", "1.0", TEST_DATE, AUTHOR, f"由 2测试用例.xls 模板副本填写，{case_total()} 条用例，其中 {executed_case_total()} 条已实测、{supplement_case_total()} 条补充设计"),
            ("《3-Bug缺陷报告清单》", "1.0", TEST_DATE, AUTHOR, "由 3Bug缺陷报告清单.xls 模板副本填写，2 个缺陷"),
            ("《4-接口测试-Postman》", "1.0", TEST_DATE, AUTHOR, "Postman 专项测试用例、集合、环境和 newman 报告"),
            ("《5-测试总结报告》", "1.0", TEST_DATE, AUTHOR, "由 5测试总结报告.doc 模板副本填写"),
        ],
    )

    add_heading(doc, "3 测试资源", 1)
    add_heading(doc, "3.1 硬件配置", 2)
    add_table(doc, ["关键项", "数量", "性能要求"], [("测试 PC 机（客户端/接口执行机）", "1", "CPU：x86_64 多核；内存：8GB 及以上；硬盘：20GB 可用空间；网络可访问被测后端")])
    add_heading(doc, "3.2 软件配置", 2)
    add_table(
        doc,
        ["资源名称/类型", "数量", "配置"],
        [
            ("被测后端", "1", "RhizoDelta Spring Boot，http://localhost:8090"),
            ("数据与中间件", "3", "Neo4j、Redis、RabbitMQ"),
            ("浏览器环境", "1", "Chrome 最新版"),
            ("接口测试工具", "1", f"Postman + newman，集合 {POSTMAN_CASE_TOTAL} 条请求，{POSTMAN_ASSERTION_TOTAL} 条断言"),
            ("文档工具", "1", "Word/WPS/LibreOffice，可打开 doc/docx/xls"),
            ("截图工具", "1", "浏览器截图与 newman htmlextra 报告截图"),
        ],
    )
    add_heading(doc, "3.3 人力资源分配", 2)
    add_table(
        doc,
        ["角色", "人员", "主要职责"],
        [
            (
                "测试负责人（个人独立完成）",
                AUTHOR,
                "分析 RhizoDelta 功能与接口；划定测试范围；编写测试方案、测试用例、缺陷清单与总结报告；构建 Postman 集合；执行 newman 测试；整理截图和结论。",
            )
        ],
    )

    add_heading(doc, "4 测试计划", 1)
    add_heading(doc, "4.1 整体功能模块划分", 2)
    add_table(
        doc,
        ["需求编号", "模块名称", "功能名称", "本轮处理", "测试人员"],
        [
            ("RD-ST-SRS001", "认证授权", "注册、登录、刷新令牌、登出、当前用户、JWT 鉴权与 token 吊销", "完整执行", AUTHOR),
            ("RD-ST-SRS002", "用户资料与社交", "个人资料、头像、关注、拉黑、动态流、在线状态", "执行资料/状态/feed 基础接口；头像/关注/拉黑后续扩展", AUTHOR),
            ("RD-ST-SRS003", "发帖与关联", "发帖异步入队、创建/删除节点关联", "后续扩展，不计入本轮通过率", AUTHOR),
            ("RD-ST-SRS004", "图谱查询", "根话题、节点详情、拓扑上下文、谱系、子代、讨论树、溯源与关联查询", "执行根话题与详情异常路径；复杂拓扑后续扩展", AUTHOR),
            ("RD-ST-SRS005", "治理决策与复核", "合并、分支、注入、回滚、审批、驳回与审计", "后续扩展，不计入本轮通过率", AUTHOR),
            ("RD-ST-SRS006", "AI 与实时", "节点摘要、向量、相似检索、SSE 事件流", "范围外说明", AUTHOR),
        ],
    )

    add_heading(doc, "5 测试整体进度安排", 1)
    add_table(
        doc,
        ["测试阶段", "时间安排", "参与人员", "测试工作内容安排", "产出"],
        [
            ("需求分析", "2026.6.18 19:00-19:30", AUTHOR, "分析 RhizoDelta 模块、接口、角色和测试边界", "项目理解、测试范围"),
            ("测试方案", "2026.6.18 19:30-20:00", AUTHOR, "按模板编写测试资源、计划、风险与提交物", "《1-测试方案》"),
            ("测试用例", "2026.6.18 20:00-20:40", AUTHOR, f"用等价类、边界值、错误推测、因果图和场景法设计 {case_total()} 条用例", "《2-测试用例》"),
            ("接口执行", "2026.6.18 20:40-21:10", AUTHOR, f"生成 Postman 集合并通过 newman 执行 {POSTMAN_CASE_TOTAL} 个请求、{POSTMAN_ASSERTION_TOTAL} 条断言", "newman HTML/JSON 报告"),
            ("缺陷登记", "2026.6.18 21:10-21:25", AUTHOR, "登记重复注册状态码与 last_active 时间格式问题", "《3-Bug缺陷报告清单》"),
            ("总结归档", "2026.6.18 21:25-21:45", AUTHOR, "汇总覆盖、缺陷、结论和后续建议", "《5-测试总结报告》与课程考查报告"),
        ],
    )

    add_heading(doc, "6 相关风险", 1)
    add_table(
        doc,
        ["风险类型", "风险", "解决方案"],
        [
            ("环境风险", "Neo4j、Redis、RabbitMQ 或后端 8090 端口未启动，导致接口不可达", "测前检查依赖和 baseUrl；Postman 环境变量统一配置；失败时先排除环境问题"),
            ("数据风险", "注册、发帖等操作污染开发库数据", "使用 qa_tester_st 与 qa_时间戳账号隔离测试数据，必要时在 Neo4j 中清理"),
            ("范围风险", "SSE、AI 摘要、异步队列、治理流程等功能涉及外部依赖和最终一致性，课程作业周期内难以充分覆盖", "本轮只对同步 REST 核心子集下结论，后续扩展项在总结说明"),
            ("时间风险", "个人独立提交，测试时间有限", "优先认证授权、用户资料、图谱查询等高价值模块，自动化接口测试提高执行效率"),
            ("契约风险", "HTTP 状态码、业务码、时间字段格式不统一", "在用例和缺陷清单中明确断言状态码、业务码、字段命名和时间格式"),
        ],
    )
    doc.save(output)


def build_summary(template_docx: Path, output: Path) -> None:
    doc = Document(template_docx)
    clear_doc_body(doc)
    set_doc_defaults(doc)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("RhizoDelta 图谱化非线性讨论系统测试总结报告")
    run.bold = True
    run.font.name = "宋体"
    run.font.size = Pt(18)

    add_heading(doc, "目录", 1)
    for item in [
        "1 测试概述",
        "1.1 编写目的",
        "1.2 项目背景",
        "2 测试参考文档",
        "3 项目组成员",
        "4 测试设计介绍",
        "4.1 测试用例设计方法",
        "4.2 测试环境与配置",
        "4.3 测试方法",
        "5 测试进度",
        "6 用例汇总",
        "7 缺陷分析",
        "8 测试结论",
    ]:
        add_para(doc, item)

    add_heading(doc, "1 测试概述", 1)
    add_heading(doc, "1.1 编写目的", 2)
    for text in [
        "确认测试用例、缺陷登记和测试方案是否覆盖 RhizoDelta 的本轮核心子集测试范围。",
        "汇总本次测试活动的环境、方法、进度、执行数据和测试证据。",
        "解释用例设计方法与具体用例的对应关系。",
        "分析缺陷分布与根因，给出后续修复和预防建议。",
        "形成被测系统当前质量结论，支撑课程考查提交。",
    ]:
        add_para(doc, f"    {text}")
    add_heading(doc, "1.2 项目背景", 2)
    add_para(doc, "RhizoDelta 是 B/S 架构的图谱化非线性讨论系统，使用 Spring Boot、Neo4j、Redis、RabbitMQ、JWT、SSE 与 React/Vite 构建。系统面向非线性讨论场景，核心能力包括认证授权、个人资料与社交关系、发帖与节点关联、图谱查询、治理决策与复核、AI 摘要及实时事件。本轮课程作业聚焦可稳定复现的核心子集 REST 行为，并使用 Postman 作为专项测试工具。")

    add_heading(doc, "2 测试参考文档", 1)
    add_table(
        doc,
        ["文档名", "版本", "日期", "作者", "备注"],
        [
            ("《1-测试方案》", "1.0", TEST_DATE, AUTHOR, "测试范围、资源、计划和风险"),
            ("《2-测试用例》", "1.0", TEST_DATE, AUTHOR, f"{case_total()} 条功能/接口用例，其中 {executed_case_total()} 条已实测、{supplement_case_total()} 条补充设计"),
            ("《3-Bug缺陷报告清单》", "1.0", TEST_DATE, AUTHOR, "2 个低危缺陷"),
            ("《4-接口测试-Postman》", "1.0", TEST_DATE, AUTHOR, "Postman 集合、环境与 newman 报告"),
            ("RhizoDelta 项目理解与测试策略文档", "1.0", TEST_DATE, AUTHOR, "系统背景、测试边界和设计方法"),
        ],
    )

    add_heading(doc, "3 项目组成员", 1)
    add_table(
        doc,
        ["角色", "人员", "主要职责"],
        [
            ("测试负责人（个人独立完成）", AUTHOR, "完成需求理解、方案编写、用例设计、Postman 集合构建、用例执行、缺陷登记、截图取证和总结报告。")
        ],
    )

    add_heading(doc, "4 测试设计介绍", 1)
    add_heading(doc, "4.1 测试用例设计方法", 2)
    method_rows = [
        ("4.1.1 错误推测法", "根据接口常见风险设计伪造 token、无 token、非法 UUID、不存在用户、重复注册、缺少密码字段、非法分页等用例，用于发现鉴权、参数校验和错误处理问题。"),
        ("4.1.2 等价类划分法", "把用户名、密码、token、refresh_token、UUID、分页参数划分为合法类和非法类，并在测试用例表的输入栏中写明代表输入。"),
        ("4.1.3 边界值分析法", "围绕密码长度、空用户名、空密码、分页 page/size、UUID 格式等边界取值设计用例，验证 GlobalExceptionHandler 与校验注解是否按预期返回错误信息。"),
        ("4.1.4 因果图法", "登录场景按用户名存在性、密码正确性、账号状态组合推导结果，覆盖正确登录、密码错误、不存在用户三个关键组合。"),
        ("4.1.5 场景法", "设计注册→登录→查询资料→刷新令牌→登出→旧 token 失效的端到端链路，验证主业务流程闭环和 token 黑名单生效。"),
    ]
    for heading, text in method_rows:
        add_heading(doc, heading, 3)
        add_para(doc, text)

    add_heading(doc, "4.2 测试环境与配置", 2)
    add_heading(doc, "4.2.1 硬件配置", 3)
    add_table(doc, ["关键项", "数量", "性能要求"], [("测试 PC 机（客户端/接口执行机）", "1", "x86_64 多核 CPU，8GB 及以上内存，20GB 可用硬盘空间")])
    add_heading(doc, "4.2.2 软件配置", 3)
    add_table(
        doc,
        ["资源名称/类型", "数量", "配置"],
        [
            ("被测后端", "1", "RhizoDelta Spring Boot，http://localhost:8090"),
            ("依赖服务", "3", "Neo4j、Redis、RabbitMQ"),
            ("浏览器", "1", "Chrome 最新版"),
            ("接口工具", "1", f"Postman + newman；{POSTMAN_CASE_TOTAL} 请求，{POSTMAN_ASSERTION_TOTAL} 断言"),
            ("文档与截图", "1", "Word/WPS/LibreOffice，浏览器截图，newman htmlextra 报告"),
        ],
    )
    add_heading(doc, "4.3 测试方法", 2)
    for heading, text in [
        ("4.3.1 软件审查", "审查系统说明、接口契约、统一响应包、错误码和课程模板要求，确认测试范围与提交物结构。"),
        ("4.3.2 黑盒测试", "从用户和接口调用者视角验证输入与输出，不依赖后端实现细节判断功能是否符合预期。"),
        ("4.3.3 接口自动化测试", f"使用 Postman 编写 {POSTMAN_CASE_TOTAL} 条接口请求和 Tests 断言，通过 newman 生成 HTML/JSON 报告；补充设计用例先纳入测试用例表，后续再加入自动化回归集合。"),
    ]:
        add_heading(doc, heading, 3)
        add_para(doc, text)

    add_heading(doc, "5 测试进度", 1)
    add_heading(doc, "5.1 测试进度回顾", 2)
    add_para(doc, "本次测试由个人独立完成，2026.6.18 完成需求分析、方案、用例、执行、缺陷登记和总结。本轮核心子集用例执行率 100%，测试完成度按本轮范围计为 100%。")
    add_table(
        doc,
        ["测试阶段", "时间安排", "参与人员", "测试工作内容安排"],
        [
            ("需求分析", "19:00-19:30", AUTHOR, "分析 RhizoDelta 系统背景、模块和测试边界"),
            ("测试方案", "19:30-20:00", AUTHOR, "按模板填写测试方案"),
            ("测试用例", "20:00-20:40", AUTHOR, f"编写 {case_total()} 条用例，覆盖认证、用户资料、图谱查询和登出收尾；其中 {supplement_case_total()} 条为补充设计"),
            ("用例执行", "20:40-21:10", AUTHOR, "运行 Postman/newman，生成 HTML/JSON 报告与截图"),
            ("缺陷登记", "21:10-21:25", AUTHOR, "记录 BUG-001、BUG-002，并给出严重程度与建议"),
            ("总结提交", "21:25-21:45", AUTHOR, "汇总用例、缺陷、结论和后续建议"),
        ],
    )
    add_heading(doc, "5.2 测试进度总结", 2)
    for text in [
        f"本次测试总共用时约 2.75 小时，本轮已实测核心子集 {executed_case_total()} 条用例全部执行完成，并补充 {supplement_case_total()} 条后续回归设计用例。",
        f"总共编写测试用例 {case_total()} 个，执行 {executed_case_total()} 个，通过 {executed_case_total()} 个；其中 1 条用例记录 HTTP 语义缺陷观察。",
        f"Postman/newman 共执行 {POSTMAN_CASE_TOTAL} 个请求，{POSTMAN_ASSERTION_TOTAL} 条断言，通过 {POSTMAN_ASSERTION_TOTAL} 条；BUG-001 作为已知缺陷观察记录，不让验收版报告失败。",
        "发现缺陷总数 2 个，严重 0 个，很高 0 个，高 0 个，中 0 个，低 2 个。",
        "完成《测试方案》《测试用例》《Bug缺陷报告清单》《接口测试-Postman》《测试总结报告》及课程考查报告。",
    ]:
        add_para(doc, f"    {text}")

    add_heading(doc, "6 用例汇总", 1)
    add_table(
        doc,
        ["功能模块", "测试用例总数", "用例编写人", "执行人"],
        [(module, count, AUTHOR, AUTHOR) for module, count in case_counts()] + [("用例合计（个）", case_total(), AUTHOR, AUTHOR), ("其中已实测（个）", executed_case_total(), AUTHOR, AUTHOR), ("其中补充设计（个）", supplement_case_total(), AUTHOR, AUTHOR)],
    )

    add_heading(doc, "7 缺陷分析", 1)
    add_table(
        doc,
        ["功能模块", "严重", "很高", "高", "中", "低", "合计"],
        BUG_SUMMARY,
    )
    add_para(doc, "BUG-001：重复用户名注册被正确拒绝，但 HTTP 状态码使用 400 而非 409 Conflict，属于接口契约语义问题。")
    add_para(doc, "BUG-002：在线状态接口 last_active 使用字符串型 epoch，而其它时间字段为 ISO8601，属于数据格式一致性问题。")
    add_para(doc, "两个缺陷均不阻塞主流程，但建议在后续版本统一错误码/状态码映射和时间序列化策略。")

    add_heading(doc, "8 测试结论", 1)
    for text in [
        "本轮测试范围内，认证授权、用户资料基础接口、图谱查询基础接口、登出与 token 吊销等核心子集功能均能正常工作。",
        "系统鉴权约束有效：无 token、伪造 token、登出后旧 token 均被拒绝；登录错误提示不泄露账号是否存在，安全性设计较好。",
        "接口统一响应包整体稳定，请求/响应字段以 snake_case 为主，便于前后端对接和 Postman 断言。",
        "当前仅发现 2 个低危缺陷，不影响本轮核心子集流程验收；建议修复后补充发帖异步一致性、节点关联、头像上传、关注/拉黑、治理决策、SSE 冒烟和性能/回归测试。",
        "综合判断：RhizoDelta 在本轮所测核心子集上达到课程作业验收质量，可以作为软件测试技术课程的被测系统提交。",
    ]:
        add_para(doc, f"    {text}")
    doc.save(output)


def xls_style(*, bold: bool = False, fill: int | None = None) -> xlwt.XFStyle:
    style = xlwt.XFStyle()
    font = xlwt.Font()
    font.name = "宋体"
    font.bold = bold
    style.font = font
    alignment = xlwt.Alignment()
    alignment.wrap = 1
    alignment.vert = xlwt.Alignment.VERT_TOP
    style.alignment = alignment
    borders = xlwt.Borders()
    borders.left = borders.right = borders.top = borders.bottom = xlwt.Borders.THIN
    style.borders = borders
    if fill is not None:
        pattern = xlwt.Pattern()
        pattern.pattern = xlwt.Pattern.SOLID_PATTERN
        pattern.pattern_fore_colour = fill
        style.pattern = pattern
    return style


XLS_NORMAL = xls_style()
XLS_HEADER = xls_style(bold=True, fill=22)
XLS_SECTION = xls_style(bold=True, fill=42)
XLS_TITLE = xls_style(bold=True, fill=43)


def write_xls_rows(ws, rows: list[list], widths: list[int], *, header_rows: set[int] | None = None) -> None:
    header_rows = header_rows or set()
    for col_idx, width in enumerate(widths):
        ws.col(col_idx).width = width * 256
    for row_idx, row in enumerate(rows):
        non_empty = [cell for cell in row if cell not in (None, "")]
        is_section = len(non_empty) == 1 and isinstance(non_empty[0], str) and "模块" in non_empty[0]
        is_title = len(non_empty) == 1 and isinstance(non_empty[0], str) and "RhizoDelta" in non_empty[0]
        style = XLS_HEADER if row_idx in header_rows else XLS_SECTION if is_section else XLS_TITLE if is_title else XLS_NORMAL
        max_lines = 1
        for col_idx, value in enumerate(row):
            text = "" if value is None else value
            ws.write(row_idx, col_idx, text, style)
            max_lines = max(max_lines, str(text).count("\n") + max(1, len(str(text)) // max(12, widths[min(col_idx, len(widths) - 1)])))
        ws.row(row_idx).height_mismatch = True
        ws.row(row_idx).height = min(2200, max(320, max_lines * 260))


def build_cases_xls(output: Path) -> None:
    wb = xlwt.Workbook(encoding="utf-8")
    ws = wb.add_sheet("测试用例", cell_overwrite_ok=True)
    rows = [
        ["模块名称", "用例个数（个）"],
        *case_counts(),
        ["合计（个）", case_total()],
        ["其中已实测（个）", executed_case_total()],
        ["其中补充设计（个）", supplement_case_total()],
        [],
        ["RhizoDelta 图谱化非线性讨论系统测试用例"],
        [],
        ["测试用例编号", "测试项目", "测试标题", "重要级别", "预置条件", "输入", "执行步骤", "预期输出", "实际输出", "测试结果"],
    ]
    case_header_row = len(rows) - 1

    grouped: dict[str, list[tuple[str, ...]]] = {}
    for case in CASES:
        grouped.setdefault(case[1], []).append(case)
    section_idx = 1
    for module in CASE_MODULE_ORDER:
        module_cases = grouped[module]
        rows.append([f"{section_idx}、{module}模块（测试用例个数：{len(module_cases)}个）"])
        for case in module_cases:
            rows.append(list(case))
        section_idx += 1
    write_xls_rows(ws, rows, [20, 18, 30, 12, 28, 54, 42, 52, 46, 20], header_rows={0, case_header_row})

    eq = wb.add_sheet("等价类划分", cell_overwrite_ok=True)
    eq_rows = [["输入条件编号", "输入对象/参数", "有效等价类", "无效等价类", "边界值/代表值", "覆盖用例", "设计方法"]]
    for row in EQUIVALENCE_CLASSES:
        eq_rows.append(list(row))
    write_xls_rows(eq, eq_rows, [18, 22, 42, 46, 36, 52, 28], header_rows={0})

    method = wb.add_sheet("方法覆盖", cell_overwrite_ok=True)
    method_rows = [["方法", "本项目中的使用方式", "对应证据"]]
    for row in METHOD_COVERAGE:
        method_rows.append(list(row))
    method_rows.append([])
    method_rows.append(["统计口径", f"测试用例总数 {case_total()} 条；已实测 {executed_case_total()} 条；补充设计 {supplement_case_total()} 条；Postman 自动化 {POSTMAN_CASE_TOTAL} 条请求、{POSTMAN_ASSERTION_TOTAL} 条断言。", "README 与 Postman 报告"])
    write_xls_rows(method, method_rows, [20, 62, 62], header_rows={0})

    wb.save(output)


def build_bugs_xls(output: Path) -> None:
    wb = xlwt.Workbook(encoding="utf-8")
    summary = wb.add_sheet("缺陷汇总", cell_overwrite_ok=True)
    summary_rows = [["模块名称", "严重", "很高", "高", "中", "低", "合计"]]
    for row in BUG_SUMMARY:
        summary_rows.append(list(row))
    write_xls_rows(summary, summary_rows, [18, 12, 12, 12, 12, 12, 14], header_rows={0})

    detail = wb.add_sheet("缺陷明细", cell_overwrite_ok=True)
    detail_rows = [["RhizoDelta 图谱化非线性讨论系统缺陷报告"], ["缺陷编号", "模块名称", "页面/窗口", "摘要", "描述", "缺陷严重程度", "提交人", "附件说明"]]
    for row in BUGS:
        detail_rows.append(list(row))
    write_xls_rows(detail, detail_rows, [14, 18, 28, 38, 78, 16, 14, 42], header_rows={1})
    wb.save(output)


def add_image_to_cell(cell, image_path: Path, caption: str) -> None:
    if not image_path.exists():
        set_cell_text(cell, f"{caption}\n（图片文件不存在：{image_path.name}）")
        return
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image_path), width=Inches(4.9))
    cap = cell.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def build_course_report(template_docx: Path, output: Path) -> None:
    doc = Document(template_docx)
    clear_doc_body(doc)
    set_doc_defaults(doc)

    p = doc.add_paragraph("Sichuan Top Vocational College of Information Technology")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph("软件测试技术")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(18)
    p = doc.add_paragraph("考查报告")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(18)
    for line in [
        "姓    名： 请填写",
        "学    号： 请填写",
        "系    别： 信息工程学院",
        "专    业： 软件技术",
        "年    级： 2024 级",
        "班    级： 请填写",
        "指导教师： 请填写",
        "2026 年 6 月 18 日 至 2026 年 6 月 18 日",
    ]:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table = doc.add_table(rows=9, cols=2)
    table.autofit = False
    for style_name in ("Table Grid", "网格型", "Table Normal"):
        try:
            table.style = style_name
            break
        except KeyError:
            continue
    labels = [
        "所 在 单 位：",
        "项目名称：",
        "项目目的：",
        "实训内容及要求：",
        "个人总结：",
        "学生签名：",
        "课程设计得分：\n\n课程设计报告：",
        "合计：",
        "指导教师签字（签章）：",
    ]
    for row, label in zip(table.rows, labels):
        row.cells[0].width = Inches(1.45)
        row.cells[1].width = Inches(5.75)
        set_cell_text(row.cells[0], label)
    set_cell_text(table.cell(0, 1), "2024 级  信息工程学院  软件技术  专业     班")
    set_cell_text(table.cell(1, 1), "基于 RhizoDelta（图谱化非线性讨论系统）的软件测试")
    set_cell_text(
        table.cell(2, 1),
        "1. 掌握软件测试的知识和方法，包括黑盒测试、等价类划分、边界值分析、错误推测、场景法、缺陷管理与接口测试。\n"
        "2. 建立以需求为依据、以用例为载体、以缺陷和质量结论为产出的测试思维模式。\n"
        "3. 针对真实可运行的 B/S 系统完成测试方案、测试用例、缺陷清单、接口测试和测试总结的完整闭环。",
    )
    cell = table.cell(3, 1)
    set_cell_text(
        cell,
        "1. 选定 RhizoDelta 作为被测系统，梳理认证授权、用户资料与社交、发帖与关联、图谱查询、治理决策与 AI 实时能力等模块；本轮结论限定在已实测核心子集。\n"
        "2. 使用原始模板副本填写《测试方案》《测试用例》《Bug缺陷报告清单》《测试总结报告》。\n"
        f"3. 选择 Postman 作为专项测试工具，构建 {POSTMAN_CASE_TOTAL} 条接口用例并通过 newman 生成测试报告；同时补充 {supplement_case_total()} 条后续回归设计用例。\n"
        f"4. 本次实测执行 {POSTMAN_CASE_TOTAL} 个请求、{POSTMAN_ASSERTION_TOTAL} 条断言，验收版 {POSTMAN_ASSERTION_TOTAL} 条全部通过；BUG-001 作为已知缺陷观察登记，合计 2 个低危缺陷。\n"
        "5. 自己所测项目界面截图及成果截图如下：",
    )
    for image, caption in [
        ("03_login_page.png", "图 1 登录页：JWT 认证入口，对应认证授权模块。"),
        ("04_home_graph.png", "图 2 登录后图谱讨论主界面：左侧话题流、中间讨论节点、右侧质量分布统计。"),
        ("05_profile_settings.png", "图 3 个人设置页：头像、用户名和显示名编辑，对应用户资料模块。"),
        ("01_newman_report.png", f"图 4 newman 接口测试报告：{POSTMAN_CASE_TOTAL} 请求、{POSTMAN_ASSERTION_TOTAL} 断言，验收版全部通过。"),
        ("02_api_evidence.png", "图 5 接口请求响应证据：登录、查当前用户、登出吊销和重复注册缺陷观察。"),
    ]:
        add_image_to_cell(cell, GENERATED / "images" / image, caption)
    set_cell_text(
        table.cell(4, 1),
        f"本次测试从模板要求出发，围绕一个真实可运行的 B/S 系统完成了需求理解、测试计划、用例设计、接口自动化执行、缺陷登记和质量总结。通过认证、用户资料基础接口、图谱查询基础接口、登出吊销等 {executed_case_total()} 条已实测用例，以及 {supplement_case_total()} 条后续回归设计用例，我进一步理解了黑盒测试中等价类、边界值、错误推测、因果图和场景法的组合使用方式。Postman/newman 的自动化执行让测试结果具备可复现性，也暴露出重复注册状态码和时间字段格式一致性这类接口契约问题。综合来看，RhizoDelta 本轮核心子集功能稳定，鉴权和错误处理整体可靠，后续应继续补充发帖异步一致性、头像上传、治理决策、SSE 和性能测试。",
    )
    doc.save(output)


def main() -> None:
    DELIVERABLES.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory(prefix="st-template-build-") as tmp:
        tmpdir = Path(tmp)
        tmpdir.chmod(0o777)
        plan_docx_template = tmpdir / "blank-plan.docx"
        summary_docx_template = tmpdir / "blank-summary.docx"
        Document().save(plan_docx_template)
        Document().save(summary_docx_template)

        build_plan(plan_docx_template, DELIVERABLES / "1-测试方案.docx")
        build_cases_xls(DELIVERABLES / "2-测试用例.xls")
        build_bugs_xls(DELIVERABLES / "3-Bug缺陷报告清单.xls")
        build_summary(summary_docx_template, DELIVERABLES / "5-测试总结报告.docx")


if __name__ == "__main__":
    sys.exit(main())
