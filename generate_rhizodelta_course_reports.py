#!/usr/bin/env python3
"""Generate RhizoDelta course-report deliverables.

The generated reports are based on the reviewed RhizoDelta documentation and
source layout. The Vue course report follows the local course-design template
structure and formatting requirements while mapping the actual React frontend
implementation to the course's Vue.js assessment points.
"""

from __future__ import annotations

import base64
import html
import json
import os
import shutil
import subprocess
import sys
from copy import deepcopy
from pathlib import Path
from typing import Iterable

sys.path.insert(0, "/tmp/finalexams-pydeps")

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parent
JAVAEE_DIR = ROOT / "javaee-web-app"
VUE_DIR = ROOT / "vuejs-app-dev"
VUE_TEMPLATE_DOC = VUE_DIR / "final-course-design-report-template.doc"
RHIZODELTA_REPO_URL = "https://github.com/TUNTIANHAMMA-2/RhizoDelta"
RHIZODELTA_LOCAL_PATH = "/home/tthm/workspace/RhizoDelta"

FONT_PATH = Path("/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc")
if not FONT_PATH.exists():
    FONT_PATH = Path("/usr/share/fonts/opentype/unifont/unifont.otf")
CODE_FONT_PATH = Path("/usr/share/fonts/truetype/dejavu/DejaVuSansMono.ttf")
if not CODE_FONT_PATH.exists():
    CODE_FONT_PATH = FONT_PATH


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    return ImageFont.truetype(str(FONT_PATH), size)


def code_font(size: int) -> ImageFont.FreeTypeFont:
    return ImageFont.truetype(str(CODE_FONT_PATH), size)


F_TITLE = font(36)
F_H = font(26)
F = font(21)
F_S = font(18)
F_XS = font(15)
F_CODE = code_font(15)
F_CODE_CJK = font(15)

BODY_FONT_SIZE = 12
BODY_FIRST_LINE_INDENT = Pt(24)
FIXED_LINE_SPACING = Pt(22)
HEADING1_SPACE = Pt(18)
HEADING2_SPACE = Pt(9)
CODE_IMAGE_WIDTH = Cm(10.8)
REPORT_IMAGE_WIDTH = Cm(14.2)


def text_size(draw: ImageDraw.ImageDraw, text: str, fnt: ImageFont.FreeTypeFont) -> tuple[int, int]:
    if not text:
        return (0, 0)
    box = draw.textbbox((0, 0), text, font=fnt)
    return box[2] - box[0], box[3] - box[1]


def wrap(draw: ImageDraw.ImageDraw, text: str, max_width: int, fnt: ImageFont.FreeTypeFont) -> list[str]:
    lines: list[str] = []
    for raw in text.split("\n"):
        if not raw:
            lines.append("")
            continue
        current = ""
        for ch in raw:
            trial = current + ch
            if text_size(draw, trial, fnt)[0] <= max_width or not current:
                current = trial
            else:
                lines.append(current)
                current = ch
        if current:
            lines.append(current)
    return lines


def center_text(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    text: str,
    fnt: ImageFont.FreeTypeFont = F,
    fill: str = "#172033",
    gap: int = 5,
) -> None:
    x1, y1, x2, y2 = box
    lines = wrap(draw, text, max(1, x2 - x1 - 18), fnt)
    heights = [text_size(draw, line, fnt)[1] for line in lines]
    total_h = sum(heights) + max(0, len(lines) - 1) * gap
    y = y1 + (y2 - y1 - total_h) / 2
    for line, h in zip(lines, heights):
        w, _ = text_size(draw, line, fnt)
        draw.text((x1 + (x2 - x1 - w) / 2, y), line, font=fnt, fill=fill)
        y += h + gap


def box(
    draw: ImageDraw.ImageDraw,
    rect: tuple[int, int, int, int],
    text: str,
    fill: str = "#F8FAFC",
    outline: str = "#2F5EAA",
    fnt: ImageFont.FreeTypeFont = F,
) -> None:
    draw.rounded_rectangle(rect, radius=12, fill=fill, outline=outline, width=3)
    center_text(draw, rect, text, fnt)


def line_arrow(
    draw: ImageDraw.ImageDraw,
    start: tuple[int, int],
    end: tuple[int, int],
    label: str = "",
    color: str = "#475467",
) -> None:
    draw.line((start, end), fill=color, width=3)
    dx = end[0] - start[0]
    dy = end[1] - start[1]
    length = max((dx * dx + dy * dy) ** 0.5, 1)
    ux, uy = dx / length, dy / length
    left = (-uy, ux)
    p1 = (end[0] - ux * 16 + left[0] * 8, end[1] - uy * 16 + left[1] * 8)
    p2 = (end[0] - ux * 16 - left[0] * 8, end[1] - uy * 16 - left[1] * 8)
    draw.polygon((end, p1, p2), fill=color)
    if label:
        mx = (start[0] + end[0]) // 2
        my = (start[1] + end[1]) // 2 - 26
        w, h = text_size(draw, label, F_XS)
        draw.rounded_rectangle((mx - w // 2 - 8, my - 5, mx + w // 2 + 8, my + h + 5), radius=6, fill="white")
        draw.text((mx - w // 2, my), label, font=F_XS, fill=color)


def diagram_canvas(title: str, width: int = 1500, height: int = 900) -> tuple[Image.Image, ImageDraw.ImageDraw]:
    img = Image.new("RGB", (width, height), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    draw.rectangle((0, 0, width - 1, height - 1), outline="#D7DBE8", width=2)
    tw, _ = text_size(draw, title, F_TITLE)
    draw.text(((width - tw) // 2, 28), title, font=F_TITLE, fill="#101827")
    return img, draw


def save_diagram(img: Image.Image, out_dir: Path, name: str) -> Path:
    path = out_dir / "images" / name
    path.parent.mkdir(parents=True, exist_ok=True)
    img.save(path)
    return path


def architecture_diagram(out_dir: Path) -> Path:
    img, draw = diagram_canvas("RhizoDelta 全栈架构与数据流")
    box(draw, (70, 180, 330, 310), "浏览器\nReact/Vite 前端", "#F0FDF4", "#2D8F6F")
    box(draw, (460, 150, 760, 340), "Spring Boot 后端\nController + Security\n统一 API 响应", "#EFF6FF", "#2F5EAA")
    box(draw, (870, 120, 1140, 250), "RabbitMQ\n异步发帖队列", "#FFF7ED", "#B45309")
    box(draw, (870, 310, 1140, 440), "PostConsumer\n落库 + Embedding\n质量评估 + AI 路由", "#F8FAFC", "#475467")
    box(draw, (1220, 110, 1450, 250), "Neo4j 5\nDAG + 语义关联\n向量索引", "#F5F3FF", "#7C3AED")
    box(draw, (1220, 330, 1450, 460), "Redis\n复核 TTL\nToken/状态辅助", "#FFF1F2", "#E11D48")
    box(draw, (460, 570, 760, 740), "SSE 事件总线\nNODE_CREATED\nEDGE_CREATED\nDECISION_*", "#ECFEFF", "#0891B2")
    box(draw, (870, 570, 1140, 740), "LangGraph4j 工作流\n召回 → 裁决 → 反思\nMERGE / BRANCH\nREVIEW", "#F0F9FF", "#0369A1")
    line_arrow(draw, (330, 245), (460, 245), "REST /api")
    line_arrow(draw, (760, 245), (870, 185), "POST /api/posts")
    line_arrow(draw, (1005, 250), (1005, 310), "consume")
    line_arrow(draw, (1140, 350), (1220, 180), "write/query")
    line_arrow(draw, (1140, 400), (1220, 395), "review TTL")
    line_arrow(draw, (1005, 440), (1005, 570), "orchestrate")
    line_arrow(draw, (870, 655), (760, 655), "events")
    line_arrow(draw, (460, 655), (330, 275), "SSE stream")
    return save_diagram(img, out_dir, "01_fullstack_architecture.png")


def graph_model_diagram(out_dir: Path) -> Path:
    img, draw = diagram_canvas("图数据模型：版本演进 DAG + 语义关联层")
    box(draw, (90, 180, 320, 300), "Human_Post\n用户原始观点", "#EFF6FF", "#2563EB")
    box(draw, (590, 160, 840, 320), "AI_Consensus\nAI 共识摘要", "#F5F3FF", "#7C3AED")
    box(draw, (1120, 180, 1370, 300), "Result\n阶段性结果", "#ECFDF5", "#059669")
    box(draw, (90, 540, 320, 660), "UserAccount\nUserProfile", "#FFF7ED", "#B45309")
    box(draw, (590, 520, 840, 680), "Topic\nPreferenceEvent", "#FEFCE8", "#CA8A04")
    box(draw, (1120, 540, 1370, 660), "语义关联边\nRELATES_TO\nCONCEPTUAL_OVERLAP", "#F8FAFC", "#475467")
    line_arrow(draw, (320, 240), (590, 240), "SYNTHESIZED_FROM / MERGED_INTO")
    line_arrow(draw, (840, 240), (1120, 240), "MATERIALIZED_FROM / JOIN")
    line_arrow(draw, (210, 540), (210, 300), "AUTHORED")
    line_arrow(draw, (320, 600), (590, 600), "FOLLOWS / MUTED / PREFERS")
    line_arrow(draw, (1120, 600), (840, 280), "语义层不参与拓扑排序")
    line_arrow(draw, (590, 280), (320, 280), "BRANCHED_FROM / CONTINUES_FROM")
    center_text(
        draw,
        (420, 720, 1080, 820),
        "核心原则：历史节点默认不可变；修订、合并、分叉均通过新增节点和显式关系边表达。",
        F,
        "#334155",
    )
    return save_diagram(img, out_dir, "02_graph_data_model.png")


def frontend_diagram(out_dir: Path) -> Path:
    img, draw = diagram_canvas("前端模块结构与运行链路")
    box(draw, (80, 160, 350, 290), "App.tsx\nReact Router 7\nRequireAuth", "#EFF6FF", "#2563EB")
    box(draw, (520, 120, 800, 250), "页面组件\nLogin / Home\nWorkspace / Settings", "#F0FDF4", "#16A34A")
    box(draw, (520, 340, 800, 500), "GraphWorkspace\nDesktopGraphWorkspace\nMobileDiscussionTreeView", "#F5F3FF", "#7C3AED")
    box(draw, (950, 120, 1220, 280), "Zustand Stores\nauth / graph / home\nsse / ui / tree", "#FFF7ED", "#B45309")
    box(draw, (950, 380, 1220, 520), "API 封装\nclient.ts\nnodes/posts/auth/events", "#ECFEFF", "#0891B2")
    box(draw, (1260, 250, 1460, 390), "Spring Boot API\n/api/**\nSSE stream", "#F8FAFC", "#475467")
    line_arrow(draw, (350, 225), (520, 185), "route")
    line_arrow(draw, (350, 225), (520, 420), "route")
    line_arrow(draw, (800, 185), (950, 200), "state")
    line_arrow(draw, (800, 420), (950, 200), "state")
    line_arrow(draw, (1085, 280), (1085, 380), "actions")
    line_arrow(draw, (1220, 450), (1260, 330), "fetch")
    line_arrow(draw, (1260, 295), (1220, 200), "events")
    center_text(
        draw,
        (200, 650, 1300, 780),
        "桌面端使用 React Flow 展示 Lineage / Explore 图谱；移动端使用 discussion-tree API 展示嵌套阅读视图，避免在小屏挂载复杂画布。",
        F,
        "#334155",
    )
    return save_diagram(img, out_dir, "03_frontend_architecture.png")


def backend_flow_diagram(out_dir: Path) -> Path:
    img, draw = diagram_canvas("后端发帖与 AI 编排时序")
    steps = [
        ("1. 用户提交帖子", "#EFF6FF", "#2563EB"),
        ("2. JWT 鉴权\n参数校验", "#F0FDF4", "#16A34A"),
        ("3. RabbitMQ 入队\n返回 202", "#FFF7ED", "#B45309"),
        ("4. PostConsumer\n创建 Human_Post", "#F8FAFC", "#475467"),
        ("5. Embedding\n质量评估", "#F5F3FF", "#7C3AED"),
        ("6. AI 路由\nMERGE / BRANCH\nREVIEW", "#ECFEFF", "#0891B2"),
        ("7. Neo4j 写入\n关系边和审计", "#FEFCE8", "#CA8A04"),
        ("8. SSE 推送\n前端增量刷新", "#FFF1F2", "#E11D48"),
    ]
    positions = [
        (90, 190, 340, 320),
        (440, 190, 690, 320),
        (790, 190, 1040, 320),
        (1140, 190, 1390, 320),
        (1140, 480, 1390, 610),
        (790, 480, 1040, 610),
        (440, 480, 690, 610),
        (90, 480, 340, 610),
    ]
    for i, ((text, fill, outline), rect) in enumerate(zip(steps, positions)):
        box(draw, rect, text, fill, outline, F_S)
        if i < len(steps) - 1:
            x1, y1, x2, y2 = rect
            nx1, ny1, nx2, ny2 = positions[i + 1]
            if i == 3:
                line_arrow(draw, ((x1 + x2) // 2, y2), ((nx1 + nx2) // 2, ny1))
            elif i < 3:
                line_arrow(draw, (x2, (y1 + y2) // 2), (nx1, (ny1 + ny2) // 2))
            else:
                line_arrow(draw, (x1, (y1 + y2) // 2), (nx2, (ny1 + ny2) // 2))
    center_text(
        draw,
        (180, 690, 1320, 810),
        "这个链路把 HTTP 请求线程、消息队列、图数据库写入、LLM 调用和前端实时反馈解耦，适合课程报告中说明 JavaEE 企业级应用的分层、异步和安全设计。",
        F,
        "#334155",
    )
    return save_diagram(img, out_dir, "04_backend_async_flow.png")


def code_snapshot_image(out_dir: Path, name: str, title: str, code: str) -> Path:
    width = 1350
    line_height = 22
    padding_x = 50
    padding_y = 58
    lines = code.strip("\n").splitlines()
    height = padding_y * 2 + 44 + len(lines) * line_height
    img = Image.new("RGB", (width, height), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    draw.rectangle((0, 0, width - 1, height - 1), outline="#D7DBE8", width=2)
    draw.rounded_rectangle((28, 24, width - 28, height - 24), radius=8, fill="#FFFFFF", outline="#CBD5E1", width=2)
    draw.text((padding_x, 42), title, font=F_S, fill="#101827")
    y = 88
    for idx, line in enumerate(lines, 1):
        draw_code_text(draw, (padding_x, y), f"{idx:>2}  {line}", fill="#172033")
        y += line_height
    return save_diagram(img, out_dir, name)


def draw_code_text(draw: ImageDraw.ImageDraw, xy: tuple[int, int], text: str, fill: str) -> None:
    x, y = xy
    for char in text:
        fnt = F_CODE_CJK if ord(char) > 127 else F_CODE
        draw.text((x, y), char, font=fnt, fill=fill)
        x += text_size(draw, char, fnt)[0]


def frontend_package_versions() -> dict[str, str]:
    package_json = Path(RHIZODELTA_LOCAL_PATH) / "frontend" / "package.json"
    package = json.loads(package_json.read_text(encoding="utf-8"))
    deps = package.get("dependencies", {}) | package.get("devDependencies", {})
    return {name: deps.get(name, "未声明").lstrip("^~") for name in deps}


def source_excerpt_snapshot(
    out_dir: Path,
    name: str,
    title: str,
    source: Path,
    ranges: list[tuple[int, int]],
) -> Path:
    lines = source.read_text(encoding="utf-8").splitlines()
    excerpt: list[str] = []
    for idx, (start, end) in enumerate(ranges):
        if idx:
            excerpt.append("...")
        for line_no in range(start, end + 1):
            if 1 <= line_no <= len(lines):
                excerpt.append(f"{line_no:>3}  {lines[line_no - 1]}")
    return code_snapshot_image(
        out_dir,
        name,
        f"{title}（{source.relative_to(Path(RHIZODELTA_LOCAL_PATH) / 'frontend')}）",
        "\n".join(excerpt),
    )


def route_source_snapshot(out_dir: Path) -> Path:
    return source_excerpt_snapshot(
        out_dir,
        "05_route_auth_code.png",
        "路由鉴权真实代码",
        Path(RHIZODELTA_LOCAL_PATH) / "frontend" / "src" / "App.tsx",
        [(11, 29), (48, 56)],
    )


def sse_source_snapshot(out_dir: Path) -> Path:
    return source_excerpt_snapshot(
        out_dir,
        "06_sse_store_code.png",
        "SSE 事件处理真实代码",
        Path(RHIZODELTA_LOCAL_PATH) / "frontend" / "src" / "hooks" / "useSse.ts",
        [(148, 165), (181, 194), (237, 248)],
    )


def workspace_source_snapshot(out_dir: Path) -> Path:
    return source_excerpt_snapshot(
        out_dir,
        "07_workspace_code.png",
        "图谱工作区真实代码",
        Path(RHIZODELTA_LOCAL_PATH) / "frontend" / "src" / "components" / "DesktopGraphWorkspace.tsx",
        [(50, 64), (78, 99), (165, 179)],
    )


def api_client_source_snapshot(out_dir: Path) -> Path:
    return source_excerpt_snapshot(
        out_dir,
        "08_api_client_code.png",
        "API 请求封装真实代码",
        Path(RHIZODELTA_LOCAL_PATH) / "frontend" / "src" / "api" / "client.ts",
        [(6, 19), (21, 37), (45, 57)],
    )


def home_source_snapshot(out_dir: Path) -> Path:
    return source_excerpt_snapshot(
        out_dir,
        "09_home_page_code.png",
        "首页入口真实代码",
        Path(RHIZODELTA_LOCAL_PATH) / "frontend" / "src" / "components" / "home" / "HomePage.tsx",
        [(14, 30), (50, 64)],
    )


def post_form_source_snapshot(out_dir: Path) -> Path:
    return source_excerpt_snapshot(
        out_dir,
        "10_post_form_code.png",
        "发布编辑真实代码",
        Path(RHIZODELTA_LOCAL_PATH) / "frontend" / "src" / "components" / "forms" / "PostForm.tsx",
        [(8, 28), (34, 48), (79, 88)],
    )


def node_detail_source_snapshot(out_dir: Path) -> Path:
    return source_excerpt_snapshot(
        out_dir,
        "11_node_detail_code.png",
        "节点详情治理真实代码",
        Path(RHIZODELTA_LOCAL_PATH) / "frontend" / "src" / "components" / "panels" / "NodeDetailPanel.tsx",
        [(51, 76), (147, 156)],
    )


def mobile_tree_source_snapshot(out_dir: Path) -> Path:
    return source_excerpt_snapshot(
        out_dir,
        "12_mobile_tree_code.png",
        "移动端讨论树真实代码",
        Path(RHIZODELTA_LOCAL_PATH) / "frontend" / "src" / "components" / "mobile" / "MobileDiscussionTreeView.tsx",
        [(22, 45), (128, 136)],
    )


def command_palette_source_snapshot(out_dir: Path) -> Path:
    return source_excerpt_snapshot(
        out_dir,
        "13_command_palette_code.png",
        "检索与命令面板真实代码",
        Path(RHIZODELTA_LOCAL_PATH) / "frontend" / "src" / "components" / "search" / "CommandPalette.tsx",
        [(37, 55), (101, 116), (153, 162)],
    )


def settings_source_snapshot(out_dir: Path) -> Path:
    return source_excerpt_snapshot(
        out_dir,
        "14_settings_code.png",
        "个人设置与界面偏好真实代码",
        Path(RHIZODELTA_LOCAL_PATH) / "frontend" / "src" / "components" / "settings" / "SettingsPage.tsx",
        [(10, 32), (98, 120)],
    )


def refresh_extra_vue_source_snapshots() -> dict[str, Path]:
    VUE_DIR.mkdir(exist_ok=True)
    (VUE_DIR / "images").mkdir(exist_ok=True)
    return {
        "home_code": home_source_snapshot(VUE_DIR),
        "post_form_code": post_form_source_snapshot(VUE_DIR),
        "node_detail_code": node_detail_source_snapshot(VUE_DIR),
        "mobile_tree_code": mobile_tree_source_snapshot(VUE_DIR),
        "command_palette_code": command_palette_source_snapshot(VUE_DIR),
        "settings_code": settings_source_snapshot(VUE_DIR),
    }


VUE_CODE_SNIPPETS = {
    "route": """function RequireAuth() {
  const token = useAuthStore((s) => s.token);
  const verifyToken = useAuthStore((s) => s.verifyToken);
  useEffect(() => { verifyToken(); }, [verifyToken]);
  if (!token) return <Navigate to=\"/login\" replace />;
  return <Outlet />;
}
<Route element={<RequireAuth />}>
  <Route path=\"/workspace/:rhizomeId\" element={<GraphWorkspace />} />
</Route>""",
    "workspace": """const loadTopologyContext = useGraphStore((s) => s.loadTopologyContext);
const canvasMode = useUiStore((s) => s.canvasMode);
useEffect(() => {
  if (rhizomeId) return loadGraphForRoot(rhizomeId, { loadTopologyContext });
  loadRhizomes().then(async () => {
    const rootId = useGraphStore.getState().rhizomes[0]?.node_id;
    if (rootId) await loadGraphForRoot(rootId, { loadTopologyContext });
  });
}, [rhizomeId, loadTopologyContext]);
{canvasMode === \"lineage\" ? <DagCanvas /> : <ExploreCanvas />}""",
    "api_state": """const res = await fetch(`${BASE_URL}${path}`, {
  ...options,
  headers: {
    ...(isFormData ? {} : { \"Content-Type\": \"application/json\" }),
    ...(token ? { Authorization: `Bearer ${token}` } : {}),
    ...options?.headers,
  },
});
if (res.status === 401) useAuthStore.getState().clearToken();
if (body.code !== 0) throw new Error(body.message);
return body.data;""",
    "sse": """function handleSseEvent(event: SseEvent) {
  const graphStore = useGraphStore.getState();
  switch (event.type) {
    case \"NODE_CREATED\":
      fetchNode(payload.node_id).then((node) => graphStore.addNode(node));
      useDiscussionTreeStore.getState().refreshTree();
      break;
    case \"EDGE_CREATED\":
      graphStore.addEdge(edge);
      graphStore.scheduleFlushLayout();
  }
}""",
    "home": """export function HomePage() {
  const loadRhizomes = useGraphStore((s) => s.loadRhizomes);
  const rightPanelMode = useUiStore((s) => s.rightPanelMode);
  useCommandPalette();
  useEffect(() => { loadRhizomes(); }, [loadRhizomes]);
  useSse();
  return <div className=\"min-h-screen flex bg-bg-canvas font-ui\">
    <HomeSidebar /><HomeMainColumn /><HomeRightRail />
    <CommandPalette /><ToastContainer />
  </div>;
}""",
    "post_form": """const handleSubmit = async (e: React.FormEvent) => {
  e.preventDefault();
  if (!content.replace(/<[^>]+>/g, \"\").trim() || !userId) return;
  await createPost({
    request_id: crypto.randomUUID(),
    author_id: userId,
    content,
    target_node_id: selectedNodeId ?? undefined,
  });
  addToast({ type: \"success\", message: selectedNodeId ? \"回复已排队\" : \"发布已排队\" });
};""",
    "node_detail": """useEffect(() => {
  if (!payloadNodeId) return;
  fetchNode(payloadNodeId)
    .then((freshNode) => useGraphStore.getState().addNode(freshNode))
    .catch(() => addToast({ type: \"error\", message: \"节点详情加载失败\" }));
}, [payloadNodeId, addToast]);
const updateNodePreferenceState = (patch: Partial<typeof node>) => {
  useGraphStore.getState().addNode({ ...node, ...patch });
};
{TABS.map((tab) => <button onClick={() => setActiveTab(tab.id)}>{tab.label}</button>)}""",
    "mobile_tree": """useEffect(() => {
  const resolveRoot = async () => {
    if (rhizomeId) return rhizomeId;
    await loadRhizomes();
    return useGraphStore.getState().rhizomes[0]?.node_id ?? null;
  };
  resolveRoot().then((rootId) => rootId ? loadTree(rootId) : navigate(\"/\"));
}, [rhizomeId, loadRhizomes, loadTree, navigate]);
{loadingState === \"loaded\" && rootId && <CommentTreeItem nodeId={rootId} depth={0} />}
<MobileReplyComposer /><LongPressMenu />""",
    "command_palette": """useEffect(() => {
  const query = debouncedQuery.trim();
  if (!query) return setSearchResults([]);
  searchSimilar({ query, top_k: MAX_RESULTS })
    .then((results) => setSearchResults(results ?? []))
    .catch((err) => setSearchError(err.message));
}, [debouncedQuery]);
const { vector } = await fetchEmbedding(selectedNodeId);
const results = await searchSimilar({ vector, top_k: 20 });""",
    "settings": """useEffect(() => {
  getMyProfile()
    .then((p) => { setProfile(p); setDisplayName(p.display_name ?? \"\"); })
    .catch((err) => setLoadError(err.message))
    .finally(() => setLoading(false));
}, []);
const updatedProfile = await updateProfile({ display_name: displayName || null });
setProfile(updatedProfile);
<AvatarUpload profile={profile} onProfileChange={setProfile} />
<RadiusModeToggle />""",
}


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "宋体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    r.font.size = Pt(BODY_FONT_SIZE)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_doc_defaults(doc: Document) -> None:
    styles = doc.styles
    styles["Normal"].font.name = "宋体"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(BODY_FONT_SIZE)
    for name in ["Heading 1", "Heading 2", "Heading 3"]:
        try:
            style = styles[name]
        except KeyError:
            continue
        style.font.name = "黑体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")


def add_page_number(section) -> None:
    footer = section.footer
    for paragraph in footer.paragraphs:
        paragraph.clear()
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char)
    run._r.append(instr)
    run._r.append(fld_char2)


def add_cover(doc: Document, course: str, title: str, subtitle: str) -> None:
    sec = doc.sections[0]
    sec.top_margin = Cm(2.4)
    sec.bottom_margin = Cm(2.2)
    sec.left_margin = Cm(2.4)
    sec.right_margin = Cm(2.4)

    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(course)
    r.bold = True
    r.font.name = "黑体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    r.font.size = Pt(22)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.name = "黑体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    r.font.size = Pt(20)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle)
    r.font.name = "宋体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    r.font.size = Pt(13)

    for _ in range(6):
        doc.add_paragraph()

    table = doc.add_table(rows=6, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    fields = [
        ("项目名称", "RhizoDelta 图谱化非线性讨论系统"),
        ("学生姓名", "（请填写）"),
        ("学号", "（请填写）"),
        ("班级", "（请填写）"),
        ("指导教师", "（请填写）"),
        ("完成时间", "2026 年 5 月"),
    ]
    for row, (k, v) in zip(table.rows, fields):
        set_cell_text(row.cells[0], k, True)
        set_cell_text(row.cells[1], v)

    doc.add_page_break()


def add_toc(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("目  录")
    r.bold = True
    r.font.name = "黑体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    r.font.size = Pt(16)
    doc.add_paragraph("（提交前可在 Word/WPS 中右键更新目录域。）")
    p = doc.add_paragraph()
    run = p.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "separate")
    fld_char3 = OxmlElement("w:fldChar")
    fld_char3.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char)
    run._r.append(instr)
    run._r.append(fld_char2)
    run._r.append(fld_char3)
    doc.add_page_break()


def para(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.first_line_indent = Pt(21)
    p.paragraph_format.line_spacing = 1.35


def bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, True)
    for row_values in rows:
        row = table.add_row()
        for idx, value in enumerate(row_values):
            set_cell_text(row.cells[idx], value)


def add_image(doc: Document, path: Path, caption: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Cm(15.5))
    c = doc.add_paragraph(caption)
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER


def libreoffice_convert(source: Path, target_ext: str, out_dir: Path) -> Path:
    office = shutil.which("libreoffice") or shutil.which("soffice")
    if not office:
        raise RuntimeError("LibreOffice is required to generate template-based Word deliverables.")

    out_dir.mkdir(parents=True, exist_ok=True)
    runtime_dir = Path("/tmp/libreoffice-runtime")
    cache_dir = Path("/tmp/libreoffice-cache")
    home_dir = Path("/tmp/libreoffice-home")
    profile_dir = Path(f"/tmp/libreoffice-profile-{source.stem}-{target_ext}")
    for path in [runtime_dir, cache_dir, home_dir, profile_dir]:
        path.mkdir(parents=True, exist_ok=True)
    runtime_dir.chmod(0o700)

    output = out_dir / f"{source.stem}.{target_ext}"
    output.unlink(missing_ok=True)
    env = os.environ.copy()
    env.update(
        {
            "HOME": str(home_dir),
            "XDG_RUNTIME_DIR": str(runtime_dir),
            "XDG_CACHE_HOME": str(cache_dir),
            "SAL_USE_VCLPLUGIN": "svp",
        }
    )
    command = [
        office,
        f"-env:UserInstallation=file://{profile_dir}",
        "--headless",
        "--nologo",
        "--nofirststartwizard",
        "--convert-to",
        target_ext,
        "--outdir",
        str(out_dir),
        str(source),
    ]
    result = subprocess.run(command, env=env, capture_output=True, text=True, timeout=60)
    if result.returncode != 0 or not output.exists():
        raise RuntimeError(
            "LibreOffice conversion failed:\n"
            f"command: {' '.join(command)}\nstdout: {result.stdout}\nstderr: {result.stderr}"
        )
    return output


def vue_template_document() -> Document:
    template_docx = libreoffice_convert(VUE_TEMPLATE_DOC, "docx", Path("/tmp/finalexams-vue-template"))
    template_copy = template_docx.with_name("vue-report-template-copy.docx")
    shutil.copyfile(template_docx, template_copy)
    doc = Document(str(template_copy))
    set_doc_defaults(doc)
    return doc


def set_run_font(run, size: float = BODY_FONT_SIZE, east_asia: str = "宋体", latin: str = "Times New Roman", bold: bool = False) -> None:
    run.bold = bold
    run.font.name = latin
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)


def apply_fixed_line_spacing(paragraph: Paragraph) -> None:
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    paragraph.paragraph_format.line_spacing = FIXED_LINE_SPACING


def apply_image_paragraph_spacing(paragraph: Paragraph) -> None:
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph.paragraph_format.line_spacing = 1
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(2)


def delete_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def remove_paragraph_child(paragraph: Paragraph, tag: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    child = p_pr.find(qn(tag))
    if child is not None:
        p_pr.remove(child)


def reset_template_paragraph(paragraph: Paragraph, outline_level: int | None = None) -> None:
    paragraph.style = "Normal"
    for tag in ["w:numPr", "w:outlineLvl"]:
        remove_paragraph_child(paragraph, tag)
    if outline_level is not None:
        outline = OxmlElement("w:outlineLvl")
        outline.set(qn("w:val"), str(outline_level))
        paragraph._p.get_or_add_pPr().append(outline)


def set_paragraph_text(paragraph: Paragraph, text: str, kind: str = "body") -> None:
    paragraph.clear()
    if kind == "heading1":
        reset_template_paragraph(paragraph, outline_level=0)
    elif kind == "heading":
        reset_template_paragraph(paragraph, outline_level=1)
    else:
        reset_template_paragraph(paragraph)

    run = paragraph.add_run(text)
    if kind == "heading1":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.first_line_indent = Pt(0)
        paragraph.paragraph_format.space_before = HEADING1_SPACE
        paragraph.paragraph_format.space_after = HEADING1_SPACE
        paragraph.paragraph_format.page_break_before = True
        apply_fixed_line_spacing(paragraph)
        set_run_font(run, size=16, east_asia="黑体", bold=True)
    elif kind == "heading":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.first_line_indent = Pt(0)
        paragraph.paragraph_format.space_before = HEADING2_SPACE
        paragraph.paragraph_format.space_after = HEADING2_SPACE
        paragraph.paragraph_format.page_break_before = False
        apply_fixed_line_spacing(paragraph)
        set_run_font(run, size=14, east_asia="黑体")
    elif kind == "caption":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.first_line_indent = Pt(0)
        paragraph.paragraph_format.page_break_before = False
        apply_fixed_line_spacing(paragraph)
        set_run_font(run, size=BODY_FONT_SIZE, east_asia="黑体")
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.first_line_indent = BODY_FIRST_LINE_INDENT
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.page_break_before = False
        apply_fixed_line_spacing(paragraph)
        set_run_font(run, size=BODY_FONT_SIZE)


def insert_paragraph_after(reference: Paragraph, text: str = "", kind: str = "body") -> Paragraph:
    new_element = OxmlElement("w:p")
    p_pr = reference._p.pPr
    if p_pr is not None:
        new_element.append(deepcopy(p_pr))
    reference._p.addnext(new_element)
    paragraph = Paragraph(new_element, reference._parent)
    if text:
        set_paragraph_text(paragraph, text, kind)
    return paragraph


def format_image_paragraph(paragraph: Paragraph) -> None:
    reset_template_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.page_break_before = False
    apply_image_paragraph_spacing(paragraph)


def find_paragraph(doc: Document, text: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == text:
            return paragraph
    raise ValueError(f"Paragraph not found: {text}")


def find_paragraph_index(doc: Document, text: str) -> int:
    for idx, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip() == text:
            return idx
    raise ValueError(f"Paragraph not found: {text}")


def paragraph_has_media(paragraph: Paragraph) -> bool:
    return bool(paragraph._p.xpath(".//w:drawing|.//w:pict"))


def remove_template_textboxes(doc: Document) -> None:
    for textbox in list(doc._element.xpath(".//w:txbxContent")):
        ancestor = textbox
        while ancestor is not None and ancestor.tag != qn("w:r"):
            ancestor = ancestor.getparent()
        if ancestor is not None and ancestor.getparent() is not None:
            ancestor.getparent().remove(ancestor)


def remove_template_toc_controls(doc: Document) -> None:
    for gallery in list(doc._element.xpath(".//w:docPartGallery")):
        if gallery.get(qn("w:val")) != "Table of Contents":
            continue
        ancestor = gallery
        while ancestor is not None and ancestor.tag != qn("w:sdt"):
            ancestor = ancestor.getparent()
        if ancestor is not None and ancestor.getparent() is not None:
            ancestor.getparent().remove(ancestor)


def remove_template_body_media(doc: Document) -> None:
    try:
        body_start = find_paragraph_index(doc, "1 项目概述")
    except ValueError:
        body_start = 0
    for paragraph in list(doc.paragraphs[body_start:]):
        if paragraph_has_media(paragraph):
            delete_paragraph(paragraph)


def replace_heading(doc: Document, old: str, new: str, level: int = 2) -> None:
    set_paragraph_text(find_paragraph(doc, old), new, "heading1" if level == 1 else "heading")


def normalize_vue_template_headings(doc: Document) -> None:
    for heading, _page in VUE_TOC_ENTRIES:
        try:
            paragraph = find_paragraph(doc, heading)
        except ValueError:
            continue
        prefix = heading.split()[0]
        level = 2 if "." in prefix else 1
        set_paragraph_text(paragraph, heading, "heading1" if level == 1 else "heading")


def replace_section_body(
    doc: Document,
    heading: str,
    next_heading: str,
    body_texts: list[str],
    images: list[tuple[Path, str]] | None = None,
) -> None:
    start_idx = find_paragraph_index(doc, heading)
    end_idx = find_paragraph_index(doc, next_heading)
    start_paragraph = doc.paragraphs[start_idx]
    for paragraph in list(doc.paragraphs[start_idx + 1 : end_idx]):
        delete_paragraph(paragraph)

    cursor = start_paragraph
    for text in body_texts:
        cursor = insert_paragraph_after(cursor, text, "body")

    for path, caption in images or []:
        image_paragraph = insert_paragraph_after(cursor)
        format_image_paragraph(image_paragraph)
        image_paragraph.add_run().add_picture(str(path), width=REPORT_IMAGE_WIDTH)
        cursor = insert_paragraph_after(image_paragraph, caption, "caption")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_table_borders(table, color: str = "94A3B8", size: str = "8") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is not None:
        tbl_pr.remove(borders)
    borders = OxmlElement("w:tblBorders")
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)
        borders.append(tag)
    tbl_pr.append(borders)


def set_cell_margins(cell, margin: str = "70") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side in ["top", "left", "bottom", "right"]:
        item = tc_mar.find(qn(f"w:{side}"))
        if item is None:
            item = OxmlElement(f"w:{side}")
            tc_mar.append(item)
        item.set(qn("w:w"), margin)
        item.set(qn("w:type"), "dxa")


def set_code_paragraph(paragraph: Paragraph, text: str) -> None:
    paragraph.clear()
    reset_template_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.left_indent = Pt(0)
    paragraph.paragraph_format.right_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    paragraph.paragraph_format.line_spacing = Pt(8.5)
    run = paragraph.add_run(text)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(7.2)


def add_code_table_after(reference: Paragraph, code: str) -> Paragraph:
    table = reference._parent.add_table(rows=1, cols=1, width=Cm(14.2))
    reference._p.addnext(table._tbl)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(14.2)
    set_table_borders(table, size="6")
    cell = table.cell(0, 0)
    cell.width = Cm(14.2)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(cell, "FFFFFF")
    set_cell_margins(cell)
    lines = code.strip("\n").splitlines()
    set_code_paragraph(cell.paragraphs[0], lines[0] if lines else "")
    for line in lines[1:]:
        set_code_paragraph(cell.add_paragraph(), line)
    anchor = OxmlElement("w:p")
    table._tbl.addnext(anchor)
    return Paragraph(anchor, reference._parent)


def add_placeholder_box_after(reference: Paragraph, text: str, caption_text: str) -> Paragraph:
    table = reference._parent.add_table(rows=1, cols=1, width=Cm(14.8))
    reference._p.addnext(table._tbl)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(cell, "F8FAFC")
    paragraph = cell.paragraphs[0]
    set_paragraph_text(paragraph, text, "caption")
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(12)
    caption_element = OxmlElement("w:p")
    table._tbl.addnext(caption_element)
    caption = Paragraph(caption_element, reference._parent)
    set_paragraph_text(caption, caption_text, "caption")
    return caption


def replace_module_section(
    doc: Document,
    heading: str,
    next_heading: str,
    body_text: str,
    effect_placeholder: str | None,
    effect_caption: str | None,
    code_image: Path | None,
    code_caption: str,
    code_text: str | None = None,
) -> Paragraph:
    start_idx = find_paragraph_index(doc, heading)
    end_idx = find_paragraph_index(doc, next_heading)
    start_paragraph = doc.paragraphs[start_idx]
    for paragraph in list(doc.paragraphs[start_idx + 1 : end_idx]):
        delete_paragraph(paragraph)

    cursor = insert_paragraph_after(start_paragraph, body_text, "body")
    if effect_placeholder and effect_caption:
        cursor = add_placeholder_box_after(cursor, effect_placeholder, effect_caption)
    if code_text is not None:
        anchor = add_code_table_after(cursor, code_text)
        return insert_paragraph_after(anchor, code_caption, "caption")
    if code_image is None:
        return cursor
    image_paragraph = insert_paragraph_after(cursor)
    format_image_paragraph(image_paragraph)
    image_paragraph.add_run().add_picture(str(code_image), width=CODE_IMAGE_WIDTH)
    return insert_paragraph_after(image_paragraph, code_caption, "caption")


def insert_module_after(
    reference: Paragraph,
    heading: str,
    body_texts: list[str],
    code_image: Path | None,
    code_caption: str,
    code_text: str | None = None,
) -> Paragraph:
    cursor = insert_paragraph_after(reference, heading, "heading")
    for text in body_texts:
        cursor = insert_paragraph_after(cursor, text, "body")
    if code_text is not None:
        anchor = add_code_table_after(cursor, code_text)
        return insert_paragraph_after(anchor, code_caption, "caption")
    if code_image is None:
        return cursor
    image_paragraph = insert_paragraph_after(cursor)
    format_image_paragraph(image_paragraph)
    image_paragraph.add_run().add_picture(str(code_image), width=CODE_IMAGE_WIDTH)
    return insert_paragraph_after(image_paragraph, code_caption, "caption")


def delete_section_body(doc: Document, heading: str, next_heading: str) -> None:
    start_idx = find_paragraph_index(doc, heading)
    end_idx = find_paragraph_index(doc, next_heading)
    for paragraph in list(doc.paragraphs[start_idx:end_idx]):
        delete_paragraph(paragraph)


def replace_section_body_to_end(doc: Document, heading: str, body_texts: list[str]) -> None:
    start_idx = find_paragraph_index(doc, heading)
    start_paragraph = doc.paragraphs[start_idx]
    for paragraph in list(doc.paragraphs[start_idx + 1 :]):
        delete_paragraph(paragraph)
    cursor = start_paragraph
    for text in body_texts:
        cursor = insert_paragraph_after(cursor, text, "body")


VUE_TOC_ENTRIES = [
    ("1 项目概述", "- 1 -"),
    ("1.1 项目背景", "- 1 -"),
    ("1.2 项目目标", "- 1 -"),
    ("2 技术选型", "- 2 -"),
    ("2.1 前端技术栈", "- 2 -"),
    ("2.2 后端技术栈（可选）", "- 2 -"),
    ("3 需求分析", "- 3 -"),
    ("3.1 功能需求", "- 3 -"),
    ("3.2 非功能性需求", "- 3 -"),
    ("4 详细设计", "- 4 -"),
    ("4.1 开发环境", "- 4 -"),
    ("4.2 路由与鉴权模块", "- 5 -"),
    ("4.3 图谱工作区模块", "- 7 -"),
    ("4.4 API 封装与状态管理模块", "- 9 -"),
    ("4.5 实时更新模块", "- 11 -"),
    ("4.6 首页与话题入口模块", "- 12 -"),
    ("4.7 发布与回复编辑模块", "- 13 -"),
    ("4.8 节点详情与治理模块", "- 14 -"),
    ("4.9 移动端讨论树模块", "- 15 -"),
    ("4.10 检索与命令面板模块", "- 16 -"),
    ("4.11 个人设置与界面偏好模块", "- 17 -"),
    ("4.x 难点与解决方案", "- 18 -"),
    ("5 总结与展望", "- 19 -"),
    ("5.1 项目总结", "- 19 -"),
    ("5.2 未来展望", "- 19 -"),
]


def replace_vue_template_cover(doc: Document) -> None:
    fields = [
        "RhizoDelta 图谱化非线性讨论系统前端实现",
        "（请填写）",
        "（请填写）",
        "信息工程学院",
        "软件技术",
        "（请填写）",
        "唐俊",
        "2026年6月x日",
    ]
    table = doc.tables[0]
    for row, value in zip(table.rows, fields):
        set_cell_text(row.cells[1], value)


def replace_vue_template_toc(doc: Document) -> None:
    remove_template_toc_controls(doc)
    toc = find_paragraph(doc, "目  录")
    body_start = find_paragraph_index(doc, "1 项目概述")
    toc_idx = find_paragraph_index(doc, "目  录")
    for paragraph in list(doc.paragraphs[toc_idx + 1 : body_start]):
        delete_paragraph(paragraph)

    cursor = toc
    for title, page in VUE_TOC_ENTRIES:
        paragraph = insert_paragraph_after(cursor)
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        paragraph.paragraph_format.line_spacing = Pt(22)
        paragraph.paragraph_format.tab_stops.add_tab_stop(Cm(15.0), WD_TAB_ALIGNMENT.RIGHT)
        left = paragraph.add_run(title)
        set_run_font(left)
        paragraph.add_run("\t")
        right = paragraph.add_run(page)
        set_run_font(right)
        cursor = paragraph


def javaee_markdown(diagrams: dict[str, Path]) -> str:
    img = lambda key: f"images/{diagrams[key].name}"
    return f"""# 《JavaEE企业级Web应用开发实战》期末课程设计报告

项目名称：RhizoDelta 图谱化非线性讨论系统

代码仓库：{RHIZODELTA_REPO_URL}

本地开发目录：`{RHIZODELTA_LOCAL_PATH}`

## 一、项目概述

RhizoDelta 是一个基于图谱的非线性讨论系统。它把传统论坛中的线性聊天记录组织为“共识主干 + 异议分支”的知识图谱：用户提交观点后，后端通过 Spring Boot 接收请求，RabbitMQ 解耦异步处理，Neo4j 保存不可变 DAG，AI 编排层对内容进行召回、裁决、反思和落库，前端通过 SSE 接收增量事件并刷新图谱视图。

## 二、需求分析

- 用户可以注册、登录，并通过 JWT 访问受保护接口。
- 用户可以发布新话题或回复已有节点，系统返回 `202 Accepted`，后台异步写入图谱。
- 系统需要支持节点详情、谱系 lineage、后代 children、溯源 provenance、移动端 discussion-tree 等查询。
- 管理或 Agent 角色可以执行合并、分支、注入、物化、回滚、语义关联等治理操作。
- 系统需要保证历史节点不可变，所有演进通过新增节点和关系边表达。
- 后端需要提供 SSE，让前端实时接收节点创建、边创建、决策完成、质量评分等事件。

## 三、技术选型与开发环境

| 层次 | 技术 |
|---|---|
| 后端语言与框架 | Java 17, Spring Boot 3.2.3 |
| 数据访问 | Spring Data Neo4j, Neo4jClient |
| 数据库 | Neo4j 5.22 |
| 消息队列 | RabbitMQ 3-management |
| 缓存/复核 | Redis 7 |
| 鉴权 | Spring Security, JWT, BCrypt |
| AI 编排 | LangChain4j 0.36.2, LangGraph4j 1.8.10 |
| 监控 | Spring Boot Actuator, Micrometer, Prometheus, Grafana |
| 构建 | Maven |

## 四、系统总体设计

![全栈架构]({img('architecture')})

系统采用分层架构：`api` 层负责 HTTP 输入输出，`service` 层承载业务规则，`repository` 和 `Neo4jClient` 承担图数据库访问，`infrastructure` 包统一放置安全、消息、SSE、持久化初始化和可观测性能力。

## 五、数据模型与图数据库设计

![图数据模型]({img('graph')})

核心节点包括 `Human_Post`、`AI_Consensus`、`Result`、`UserAccount`、`UserProfile`、`Topic` 和 `PreferenceEvent`。版本演进关系包括 `BRANCHED_FROM`、`MERGED_INTO`、`SYNTHESIZED_FROM`、`CONTINUES_FROM`、`CONVERGED_FROM`、`MATERIALIZED_FROM`、`CROSS_SYNTHESIZED_FROM`。语义关联关系包括 `CONCEPTUAL_OVERLAP` 和 `RELATES_TO`。用户域关系包括 `AUTHORED`、`HAS_PROFILE`、`FOLLOWS`、`MUTED`、`PREFERS`。

## 六、后端功能实现

### 6.1 认证与权限

`SecurityConfig` 关闭 CSRF，使用无状态会话，并把 `JwtAuthenticationFilter` 放入 Spring Security 过滤链。公开接口包括登录、注册、刷新 token、健康检查和头像读取；写入决策、创建语义关联、复核和回滚接口按 `USER`、`AGENT`、`ADMIN` 角色区分权限。

### 6.2 发帖异步处理

`PostController` 负责校验请求、绑定认证用户、检查目标节点、生成稳定 `event_id`，再把 `PostEventMessage` 投递到 RabbitMQ。HTTP 层等待 publisher confirm，消息确认后返回 `202 Accepted`。`PostConsumer` 消费消息后创建 `Human_Post`，异步生成 embedding、质量评分，发布 `NODE_CREATED` 和 `EDGE_CREATED` 事件，并触发 AI 路由编排。

![后端时序]({img('flow')})

### 6.3 图谱查询

`NodeQueryController` 提供只读查询：根话题、节点详情、`lineage`、`children`、`topology-context`、`discussion-tree`、`provenance` 和 `associations`。这些接口把图数据库中的节点和关系转换为前端可直接渲染的 DTO。

### 6.4 AI 编排与治理

`AiRoutingWorkflowService` 使用 LangGraph4j 定义状态图：加载帖子、确保 embedding、向量召回、上下文裁剪、规则预过滤、LLM 裁决、反思校验、提交前守卫、执行合并/分支或创建复核任务。规则层可以在高相似或低相似场景跳过 LLM，从而降低延迟和成本。

### 6.5 实时推送

`EventController` 暴露 `/api/events/stream`，通过 `SseEventService` 为当前用户注册 `SseEmitter`。后端在发帖排队、节点创建、边创建、决策完成、摘要生成和质量评分时发布事件，前端据此做增量刷新。

## 七、部署与运行

本地运行需要 JDK 17、Maven、Node.js、npm、Docker Compose 和可用的 `DASHSCOPE_API_KEY`。核心命令：

```bash
docker compose up -d neo4j rabbitmq redis
export DASHSCOPE_API_KEY=your_api_key_here
mvn spring-boot:run
cd frontend && npm install && npm run dev
```

后端默认端口为 `8090`，前端开发端口为 `5173`，Vite 会把 `/api` 代理到 `http://localhost:8090`。

## 八、测试与质量保障

后端测试使用 JUnit 5、Spring Boot Test 和 Testcontainers。当前仓库包含 92 个 Java 测试文件，覆盖认证、发帖、决策、图查询、SSE、用户画像、偏好聚合、AI 路由、数据库初始化等模块。常用验证命令：

```bash
mvn test -Dspring.profiles.active=test
```

## 九、总结

本项目从 JavaEE 企业级 Web 应用的角度，体现了分层架构、REST API、Spring Security 鉴权、异步消息、图数据库建模、实时事件推送、自动化测试和可观测性等综合能力。项目难点不在单表 CRUD，而在如何把非线性讨论、不可变历史、AI 编排和前端实时反馈组合成一个可运行的工程系统。
"""


def legacy_vue_markdown(diagrams: dict[str, Path]) -> str:
    img = lambda key: f"images/{diagrams[key].name}"
    return f"""# 《Vue.js应用开发》期末课程设计报告

项目名称：RhizoDelta 图谱化非线性讨论系统前端

说明：按课程提交要求，本报告保留 Vue.js 课程封面与文件名；正文依据实际项目填写。实际前端技术栈为 React 19 + TypeScript + Vite。

代码仓库：{RHIZODELTA_REPO_URL}

本地开发目录：`{RHIZODELTA_LOCAL_PATH}/frontend`

## 一、项目概述

RhizoDelta 前端是一个图谱化讨论系统的交互界面。用户可以登录或注册，浏览根话题，进入图谱工作区查看观点的演进关系，发布新话题，对节点执行延续注入或分叉，并通过右侧详情面板查看溯源、关联和审计信息。桌面端使用 React Flow 展示 DAG 图谱，移动端使用嵌套讨论树降低复杂图形渲染成本。

## 二、需求分析

- 提供登录、注册和受保护路由。
- 首页展示根话题、信息流和右侧辅助信息。
- 工作区支持桌面图谱模式和移动端讨论树模式。
- 节点可查看详情、确权溯源、语义关联和审计时间线。
- 用户可以发布新话题、回复节点、执行延续注入和分叉。
- 前端需要监听 SSE，实时处理节点创建、边创建、决策完成、质量评分等事件。
- UI 需要适配桌面和移动端，保证大规模图谱场景下的交互可用性。

## 三、技术栈与开发环境

| 类别 | 技术 |
|---|---|
| UI 框架 | React 19 |
| 类型系统 | TypeScript 5.x |
| 构建工具 | Vite 8 |
| 路由 | React Router 7 |
| 图谱画布 | @xyflow/react, @dagrejs/dagre, d3-force |
| 状态管理 | Zustand 5 |
| 编辑器 | TipTap 3 + Markdown |
| 样式 | Tailwind CSS 4, CSS Design Tokens |
| 测试 | Vitest, Testing Library |

## 四、前端总体架构

![前端架构]({img('frontend')})

前端入口为 `src/main.tsx` 和 `src/App.tsx`。`App.tsx` 管理路由和鉴权保护：`/login` 是公开路由，`/`、`/workspace`、`/workspace/:rhizomeId`、`/settings` 都需要有效 token。`GraphWorkspace` 根据视口宽度选择桌面图谱工作区或移动端讨论树。

## 五、页面与组件设计

### 5.1 登录/注册页

`LoginPage.tsx` 调用 `/api/auth/login` 和 `/api/auth/register`，成功后把 access token 写入 `localStorage.jwt_token`，并交给 `authStore` 维护用户身份、角色和会话状态。

### 5.2 首页

首页由 `HomePage` 及 `HomeSidebar`、`HomeMainColumn`、`HomeRightRail` 等组件组成，用于展示根话题、动态信息和个人化入口。

### 5.3 工作区

桌面端 `DesktopGraphWorkspace` 组合左侧话题列表、中间 React Flow 画布和右侧详情面板。画布支持版本视图和探索视图，节点点击后可打开 `NodeDetailPanel`，工具条可触发注入和分叉操作。

### 5.4 移动端讨论树

移动端不直接挂载复杂图谱画布，而是请求 `/api/nodes/{{id}}/discussion-tree`，由 `MobileDiscussionTreeView`、`CommentTreeItem`、`MobileReplyComposer` 和 `LongPressMenu` 展示更适合小屏阅读的嵌套讨论结构。

## 六、API 封装与状态管理

`src/api/client.ts` 是统一请求入口，自动附加 `Authorization: Bearer <token>`，并处理统一响应结构 `{{code, message, data}}`。领域 API 被拆分为 `auth.ts`、`nodes.ts`、`posts.ts`、`decisions.ts`、`associations.ts`、`audit.ts`、`profile.ts`、`feed.ts`、`events.ts` 等。

Zustand Store 按职责拆分：`authStore` 管理登录状态，`graphStore` 管理节点、边、选择态和图谱视图，`sseStore` 管理连接状态，`homeStore` 管理首页数据，`discussionTreeStore` 管理移动端讨论树，`uiStore` 管理侧边栏、右面板和 toast。

## 七、实时更新与交互流程

前端通过 `useSse.ts` 请求 `/api/events/stream`。连接建立后解析 SSE 文本块，并根据事件类型更新 store：

- `NODE_CREATED`：拉取新节点并加入图谱。
- `EDGE_CREATED`：补齐端点节点后加入关系边。
- `DECISION_COMPLETE`：替换乐观节点并展示结果。
- `ORCHESTRATION_STATUS`：更新编排状态。
- `SUMMARY_GENERATED`、`QUALITY_SCORED`：刷新摘要和质量徽章。

![全栈数据流]({img('architecture')})

## 八、界面风格与响应式设计

项目采用 Botanical Observatory / Wikipedia-Notion 风格：暖纸色背景、植物学绿色强调色、内容区衬线字体、控件区无衬线字体。核心 design token 位于 `src/styles/tokens.css`。界面在桌面端提供三栏工作台，在移动端切换为讨论树和轻量操作菜单。

## 九、运行、构建与测试

前端运行命令：

```bash
cd frontend
npm install
npm run dev
```

构建命令：

```bash
npm run build
```

Lint 与测试：

```bash
npm run lint
npx vitest
```

当前前端包含 118 个 TypeScript/CSS 源文件和多组测试文件，覆盖表单、节点操作、移动端组件、Markdown 工具、SSE 事件、viewport hook 和 Zustand store。

## 十、总结

本项目前端重点解决了图谱型应用中“数据结构复杂、实时事件多、桌面与移动体验差异大”的问题。通过 API 分层、Zustand 状态拆分、React Flow 图谱渲染、SSE 增量更新和移动端讨论树降级方案，系统形成了较完整的前端工程实践。虽然课程名称为 Vue.js 应用开发，本报告正文按照学校允许的实际项目内容填写，保留课程封面和文件名要求。
"""


def add_javaee_content(doc: Document, diagrams: dict[str, Path]) -> None:
    doc.add_heading("一、项目概述", 1)
    para(doc, "RhizoDelta 是一个基于图谱的非线性讨论系统。它把传统论坛中的线性聊天记录组织为“共识主干 + 异议分支”的知识图谱：用户提交观点后，后端通过 Spring Boot 接收请求，RabbitMQ 解耦异步处理，Neo4j 保存不可变 DAG，AI 编排层对内容进行召回、裁决、反思和落库，前端通过 SSE 接收增量事件并刷新图谱视图。")
    para(doc, f"代码仓库为 {RHIZODELTA_REPO_URL}，本地开发目录为 {RHIZODELTA_LOCAL_PATH}。")

    doc.add_heading("二、需求分析", 1)
    for item in [
        "用户可以注册、登录，并通过 JWT 访问受保护接口。",
        "用户可以发布新话题或回复已有节点，系统返回 202 Accepted，后台异步写入图谱。",
        "系统需要支持节点详情、谱系 lineage、后代 children、溯源 provenance、移动端 discussion-tree 等查询。",
        "管理或 Agent 角色可以执行合并、分支、注入、物化、回滚、语义关联等治理操作。",
        "系统需要保证历史节点不可变，所有演进通过新增节点和关系边表达。",
        "后端需要提供 SSE，让前端实时接收节点创建、边创建、决策完成、质量评分等事件。",
    ]:
        bullet(doc, item)

    doc.add_heading("三、技术选型与开发环境", 1)
    add_table(doc, ["层次", "技术"], [
        ["后端语言与框架", "Java 17, Spring Boot 3.2.3"],
        ["数据访问", "Spring Data Neo4j, Neo4jClient"],
        ["数据库", "Neo4j 5.22"],
        ["消息队列", "RabbitMQ 3-management"],
        ["缓存/复核", "Redis 7"],
        ["鉴权", "Spring Security, JWT, BCrypt"],
        ["AI 编排", "LangChain4j 0.36.2, LangGraph4j 1.8.10"],
        ["监控", "Spring Boot Actuator, Micrometer, Prometheus, Grafana"],
        ["构建", "Maven"],
    ])

    doc.add_heading("四、系统总体设计", 1)
    add_image(doc, diagrams["architecture"], "图 1 RhizoDelta 全栈架构与数据流")
    para(doc, "系统采用分层架构：api 层负责 HTTP 输入输出，service 层承载业务规则，repository 和 Neo4jClient 承担图数据库访问，infrastructure 包统一放置安全、消息、SSE、持久化初始化和可观测性能力。")

    doc.add_heading("五、数据模型与图数据库设计", 1)
    add_image(doc, diagrams["graph"], "图 2 图数据模型")
    para(doc, "核心节点包括 Human_Post、AI_Consensus、Result、UserAccount、UserProfile、Topic 和 PreferenceEvent。版本演进关系包括 BRANCHED_FROM、MERGED_INTO、SYNTHESIZED_FROM、CONTINUES_FROM、CONVERGED_FROM、MATERIALIZED_FROM、CROSS_SYNTHESIZED_FROM。语义关联关系包括 CONCEPTUAL_OVERLAP 和 RELATES_TO。用户域关系包括 AUTHORED、HAS_PROFILE、FOLLOWS、MUTED、PREFERS。")

    doc.add_heading("六、后端功能实现", 1)
    doc.add_heading("6.1 认证与权限", 2)
    para(doc, "SecurityConfig 关闭 CSRF，使用无状态会话，并把 JwtAuthenticationFilter 放入 Spring Security 过滤链。公开接口包括登录、注册、刷新 token、健康检查和头像读取；写入决策、创建语义关联、复核和回滚接口按 USER、AGENT、ADMIN 角色区分权限。")
    doc.add_heading("6.2 发帖异步处理", 2)
    para(doc, "PostController 负责校验请求、绑定认证用户、检查目标节点、生成稳定 event_id，投递 PostEventMessage 到 RabbitMQ。HTTP 层等待 publisher confirm，消息确认后返回 202 Accepted。PostConsumer 消费消息后创建 Human_Post，异步生成 embedding、质量评分，发布 NODE_CREATED 和 EDGE_CREATED 事件，并触发 AI 路由编排。")
    add_image(doc, diagrams["flow"], "图 3 后端发帖与 AI 编排时序")
    doc.add_heading("6.3 图谱查询", 2)
    para(doc, "NodeQueryController 提供只读查询：根话题、节点详情、lineage、children、topology-context、discussion-tree、provenance 和 associations。这些接口把图数据库中的节点和关系转换为前端可直接渲染的 DTO。")
    doc.add_heading("6.4 AI 编排与治理", 2)
    para(doc, "AiRoutingWorkflowService 使用 LangGraph4j 定义状态图：加载帖子、确保 embedding、向量召回、上下文裁剪、规则预过滤、LLM 裁决、反思校验、提交前守卫、执行合并/分支或创建复核任务。规则层可以在高相似或低相似场景跳过 LLM，从而降低延迟和成本。")
    doc.add_heading("6.5 实时推送", 2)
    para(doc, "EventController 暴露 /api/events/stream，通过 SseEventService 为当前用户注册 SseEmitter。后端在发帖排队、节点创建、边创建、决策完成、摘要生成和质量评分时发布事件，前端据此做增量刷新。")

    doc.add_heading("七、部署与运行", 1)
    para(doc, "本地运行需要 JDK 17、Maven、Node.js、npm、Docker Compose 和可用的 DASHSCOPE_API_KEY。后端默认端口为 8090，前端开发端口为 5173，Vite 会把 /api 代理到 http://localhost:8090。")
    add_table(doc, ["步骤", "命令"], [
        ["启动基础设施", "docker compose up -d neo4j rabbitmq redis"],
        ["设置模型密钥", "export DASHSCOPE_API_KEY=your_api_key_here"],
        ["启动后端", "mvn spring-boot:run"],
        ["启动前端", "cd frontend && npm install && npm run dev"],
    ])

    doc.add_heading("八、测试与质量保障", 1)
    para(doc, "后端测试使用 JUnit 5、Spring Boot Test 和 Testcontainers。当前仓库包含 92 个 Java 测试文件，覆盖认证、发帖、决策、图查询、SSE、用户画像、偏好聚合、AI 路由、数据库初始化等模块。常用验证命令为 mvn test -Dspring.profiles.active=test。")

    doc.add_heading("九、总结", 1)
    para(doc, "本项目从 JavaEE 企业级 Web 应用的角度，体现了分层架构、REST API、Spring Security 鉴权、异步消息、图数据库建模、实时事件推送、自动化测试和可观测性等综合能力。项目难点不在单表 CRUD，而在如何把非线性讨论、不可变历史、AI 编排和前端实时反馈组合成一个可运行的工程系统。")


def legacy_add_vue_content(doc: Document, diagrams: dict[str, Path]) -> None:
    doc.add_heading("一、项目概述", 1)
    para(doc, "按课程提交要求，本报告保留 Vue.js 课程封面与文件名；正文依据实际项目填写。实际前端技术栈为 React 19 + TypeScript + Vite。RhizoDelta 前端是一个图谱化讨论系统的交互界面。用户可以登录或注册，浏览根话题，进入图谱工作区查看观点的演进关系，发布新话题，对节点执行延续注入或分叉，并通过右侧详情面板查看溯源、关联和审计信息。")
    para(doc, f"代码仓库为 {RHIZODELTA_REPO_URL}，前端目录为 {RHIZODELTA_LOCAL_PATH}/frontend。")

    doc.add_heading("二、需求分析", 1)
    for item in [
        "提供登录、注册和受保护路由。",
        "首页展示根话题、信息流和右侧辅助信息。",
        "工作区支持桌面图谱模式和移动端讨论树模式。",
        "节点可查看详情、确权溯源、语义关联和审计时间线。",
        "用户可以发布新话题、回复节点、执行延续注入和分叉。",
        "前端需要监听 SSE，实时处理节点创建、边创建、决策完成、质量评分等事件。",
        "UI 需要适配桌面和移动端，保证大规模图谱场景下的交互可用性。",
    ]:
        bullet(doc, item)

    doc.add_heading("三、技术栈与开发环境", 1)
    add_table(doc, ["类别", "技术"], [
        ["UI 框架", "React 19"],
        ["类型系统", "TypeScript 5.x"],
        ["构建工具", "Vite 8"],
        ["路由", "React Router 7"],
        ["图谱画布", "@xyflow/react, @dagrejs/dagre, d3-force"],
        ["状态管理", "Zustand 5"],
        ["编辑器", "TipTap 3 + Markdown"],
        ["样式", "Tailwind CSS 4, CSS Design Tokens"],
        ["测试", "Vitest, Testing Library"],
    ])

    doc.add_heading("四、前端总体架构", 1)
    add_image(doc, diagrams["frontend"], "图 1 前端模块结构与运行链路")
    para(doc, "前端入口为 src/main.tsx 和 src/App.tsx。App.tsx 管理路由和鉴权保护：/login 是公开路由，/、/workspace、/workspace/:rhizomeId、/settings 都需要有效 token。GraphWorkspace 根据视口宽度选择桌面图谱工作区或移动端讨论树。")

    doc.add_heading("五、页面与组件设计", 1)
    doc.add_heading("5.1 登录/注册页", 2)
    para(doc, "LoginPage.tsx 调用 /api/auth/login 和 /api/auth/register，成功后把 access token 写入 localStorage.jwt_token，并交给 authStore 维护用户身份、角色和会话状态。")
    doc.add_heading("5.2 首页", 2)
    para(doc, "首页由 HomePage 及 HomeSidebar、HomeMainColumn、HomeRightRail 等组件组成，用于展示根话题、动态信息和个人化入口。")
    doc.add_heading("5.3 工作区", 2)
    para(doc, "桌面端 DesktopGraphWorkspace 组合左侧话题列表、中间 React Flow 画布和右侧详情面板。画布支持版本视图和探索视图，节点点击后可打开 NodeDetailPanel，工具条可触发注入和分叉操作。")
    doc.add_heading("5.4 移动端讨论树", 2)
    para(doc, "移动端不直接挂载复杂图谱画布，而是请求 /api/nodes/{id}/discussion-tree，由 MobileDiscussionTreeView、CommentTreeItem、MobileReplyComposer 和 LongPressMenu 展示更适合小屏阅读的嵌套讨论结构。")

    doc.add_heading("六、API 封装与状态管理", 1)
    para(doc, "src/api/client.ts 是统一请求入口，自动附加 Authorization: Bearer <token>，并处理统一响应结构 {code, message, data}。领域 API 被拆分为 auth.ts、nodes.ts、posts.ts、decisions.ts、associations.ts、audit.ts、profile.ts、feed.ts、events.ts 等。")
    para(doc, "Zustand Store 按职责拆分：authStore 管理登录状态，graphStore 管理节点、边、选择态和图谱视图，sseStore 管理连接状态，homeStore 管理首页数据，discussionTreeStore 管理移动端讨论树，uiStore 管理侧边栏、右面板和 toast。")

    doc.add_heading("七、实时更新与交互流程", 1)
    para(doc, "前端通过 useSse.ts 请求 /api/events/stream。连接建立后解析 SSE 文本块，并根据事件类型更新 store：NODE_CREATED 拉取新节点并加入图谱；EDGE_CREATED 补齐端点节点后加入关系边；DECISION_COMPLETE 替换乐观节点并展示结果；ORCHESTRATION_STATUS 更新编排状态；SUMMARY_GENERATED、QUALITY_SCORED 刷新摘要和质量徽章。")
    add_image(doc, diagrams["architecture"], "图 2 全栈数据流与 SSE 实时更新")

    doc.add_heading("八、界面风格与响应式设计", 1)
    para(doc, "项目采用 Botanical Observatory / Wikipedia-Notion 风格：暖纸色背景、植物学绿色强调色、内容区衬线字体、控件区无衬线字体。核心 design token 位于 src/styles/tokens.css。界面在桌面端提供三栏工作台，在移动端切换为讨论树和轻量操作菜单。")

    doc.add_heading("九、运行、构建与测试", 1)
    add_table(doc, ["用途", "命令"], [
        ["安装依赖", "cd frontend && npm install"],
        ["开发运行", "npm run dev"],
        ["生产构建", "npm run build"],
        ["代码检查", "npm run lint"],
        ["单元测试", "npx vitest"],
    ])
    para(doc, "当前前端包含 118 个 TypeScript/CSS 源文件和多组测试文件，覆盖表单、节点操作、移动端组件、Markdown 工具、SSE 事件、viewport hook 和 Zustand store。")

    doc.add_heading("十、总结", 1)
    para(doc, "本项目前端重点解决了图谱型应用中“数据结构复杂、实时事件多、桌面与移动体验差异大”的问题。通过 API 分层、Zustand 状态拆分、React Flow 图谱渲染、SSE 增量更新和移动端讨论树降级方案，系统形成了较完整的前端工程实践。虽然课程名称为 Vue.js 应用开发，本报告正文按照学校允许的实际项目内容填写，保留课程封面和文件名要求。")


def vue_react_report_content(versions: dict[str, str]) -> dict[str, object]:
    background = (
        f"《Vue.js应用开发》课程要求提交前端课程设计报告和 Word 文档，文件名与课程目录按要求保留 Vue.js，"
        f"但本项目真实代码实现是 React 前端，因此正文全部按 /home/tthm/workspace/RhizoDelta/frontend 中的 React、"
        f"TypeScript 和 Vite 源码编写，不把现有系统伪写成 Vue 项目。RhizoDelta 图谱化非线性讨论系统前端选取"
        f"“多人围绕同一议题进行发散讨论与共识沉淀”这一场景。传统论坛、群聊和评论区通常按时间线排列内容，"
        f"用户很难在长讨论中快速看清观点之间的继承、分歧、合并和转化关系；随着讨论规模扩大，重复观点、"
        f"分支争论、证据补充和阶段性总结会混在一起，后续参与者需要反复回看历史消息，管理者也难以判断结论"
        f"来自哪些原始发言。项目源码中的 React 页面、React Router 路由、Zustand 状态、fetch API 客户端、SSE "
        f"事件处理和 React Flow 图谱画布共同解决这个问题：节点代表用户观点、AI 共识或阶段结果，边代表延续、"
        f"分叉、合并、语义关联等关系，用户可以从根话题进入工作区查看观点演进，也可以在移动端使用讨论树阅读。"
        f"本报告从真实 React {versions['react']}、TypeScript {versions['typescript']}、Vite {versions['vite']} "
        f"工程出发，总结组件化、路由鉴权、接口封装、状态管理、响应式布局和实时更新等前端开发能力。"
    )
    target = (
        "本项目的功能目标是完成一套可运行的图谱化讨论前端：用户能够注册、登录和保持会话；首页能够展示根话题、"
        "信息流和辅助入口；工作区能够按桌面端图谱模式和移动端讨论树模式呈现节点关系；节点详情面板能够展示正文、"
        "作者、质量评分、溯源、关联和审计信息；用户能够发布新话题、回复节点，并根据权限执行延续注入、分叉和复核"
        "等操作；前端还要通过 SSE 实时接收后端事件，使节点创建、边创建、AI 裁决完成和质量评分可以增量刷新。上述目标"
        "来自 src/App.tsx、src/components/GraphWorkspace.tsx、src/components/DesktopGraphWorkspace.tsx、"
        "src/components/mobile/MobileDiscussionTreeView.tsx、src/api/client.ts 和 src/hooks/useSse.ts 等真实源码。"
        "技术目标是掌握 React 函数组件、Hooks、React Router 嵌套路由、Zustand store、统一 fetch 请求封装、"
        "TypeScript 类型约束、Vite 构建、Tailwind 样式和 Vitest 测试等现代前端工程能力。"
    )
    return {
        "background": background,
        "target": target,
        "tech_stack": [
            f"UI 框架与语言：真实工程在 package.json 中声明 React {versions['react']}、TypeScript {versions['typescript']} 和 Vite {versions['vite']}。页面由函数组件和 Hooks 组织，入口在 src/main.tsx 与 src/App.tsx，业务组件分布在 src/components/ 下。",
            f"路由与页面组织：项目使用 React Router {versions['react-router-dom']} 管理 /login、/、/workspace、/workspace/:rhizomeId 和 /settings 等页面。src/App.tsx 中通过 BrowserRouter、Routes、Route、Navigate、Outlet、RequireAuth 和 PublicOnlyRoute 实现公开路由与受保护路由。",
            f"状态管理：项目使用 Zustand {versions['zustand']}，按职责拆分 authStore、graphStore、sseStore、homeStore、discussionTreeStore 和 uiStore。认证、图谱、SSE、首页、移动端讨论树和界面状态互相独立，组件通过 store action 更新状态。",
            f"UI 与图谱渲染：项目使用 Tailwind CSS {versions['tailwindcss']}、CSS Design Tokens、@xyflow/react {versions['@xyflow/react']}、@dagrejs/dagre {versions['@dagrejs/dagre']} 和 d3-force {versions['d3-force']} 实现布局、主题、图谱节点边渲染和布局计算。",
            f"接口、编辑器与测试：统一 API 客户端在 src/api/client.ts 中使用原生 fetch，负责附加 JWT token、处理 401 和解析 {{code, message, data}} 响应结构；TipTap {versions['@tiptap/react']} 与 Markdown 工具用于内容编辑；Vitest {versions['vitest']}、Testing Library 和 ESLint {versions['eslint']} 用于质量保障。",
        ],
        "backend": "项目后端由 Java 17、Spring Boot 3.2.3、Spring Security、Neo4j 5、RabbitMQ、Redis 和 SSE 事件服务组成。前端通过 /api/** 访问后端 REST 接口，通过 /api/events/stream 订阅实时事件。后端负责用户认证、节点查询、发帖异步处理、AI 路由决策、图数据库写入和审计记录，前端负责把这些数据转换为用户可理解、可操作的 React 页面状态。",
        "functional": [
            "（1）用户模块：支持注册、登录、退出和 token 持久化。登录成功后进入受保护页面，未登录用户访问工作区或设置页时自动跳转到登录页。",
            "（2）首页模块：展示根话题列表、动态信息和右侧辅助入口，帮助用户快速找到正在讨论的主题，并进入指定话题的工作区。",
            "（3）图谱工作区模块：桌面端展示图谱画布、话题列表和右侧详情面板；用户可以缩放、选择节点、查看节点关系、切换谱系视图和探索视图。",
            "（4）移动端讨论树模块：小屏设备不直接挂载复杂图谱画布，而是请求 discussion-tree 接口，用嵌套评论树展示讨论结构，并提供长按菜单和移动端回复输入。",
            "（5）节点详情与治理模块：节点详情面板展示正文、作者、质量评分、溯源、关联和审计时间线；具备权限的用户可以触发延续注入、分叉、复核和其他治理操作。",
            "（6）实时更新模块：前端监听 SSE 事件，并根据 NODE_CREATED、EDGE_CREATED、DECISION_COMPLETE、ORCHESTRATION_STATUS、SUMMARY_GENERATED 和 QUALITY_SCORED 等事件增量更新页面。",
            "（7）检索与命令面板模块：支持最近节点、文本相似检索、向量相似检索和键盘选择，帮助用户在大规模讨论图谱中快速定位内容。",
            "（8）个人设置与偏好模块：支持头像、展示名称和界面圆角偏好设置，使系统具备完整的用户资料维护和界面个性化能力。",
        ],
        "nonfunctional": "性能方面，桌面端图谱需要避免无意义重绘，移动端需要使用讨论树降低复杂画布渲染成本；接口请求应使用统一封装，避免重复处理 token 和错误提示。兼容性方面，项目面向现代浏览器，支持桌面和移动端响应式布局。安全性方面，受保护路由必须检查认证状态，请求头必须携带 Bearer token，前端不能把权限判断作为唯一安全边界。可维护性方面，页面组件、API 模块、store 和工具函数按职责拆分，测试覆盖表单、节点操作、移动端组件、Markdown 工具、SSE 事件和状态管理。",
        "environment": f"本项目开发环境为 Linux 工作站，代码编辑工具为 Visual Studio Code 或同类编辑器，前端目录位于 {RHIZODELTA_LOCAL_PATH}/frontend。项目使用 Node.js、npm、Vite {versions['vite']}、TypeScript {versions['typescript']}、ESLint {versions['eslint']}、Vitest {versions['vitest']} 和 Testing Library 完成开发、构建、代码检查和测试。常用命令来自 package.json，包括 npm run dev、npm run build、npm run lint 和测试命令 npx vitest。本报告内容逐项核对了 package.json、src/App.tsx、src/api/client.ts、src/api/nodes.ts、src/hooks/useSse.ts、src/stores/*.ts、src/components/GraphWorkspace.tsx、src/components/DesktopGraphWorkspace.tsx 和 src/components/mobile/MobileDiscussionTreeView.tsx。后端联调时需要先启动 Spring Boot 服务，并保证 Vite 将 /api 代理到 http://localhost:8090。",
        "route": "路由与鉴权模块负责把公开页面和受保护页面分开。/login 允许未登录用户访问，首页、工作区和设置页都需要有效 token。用户登录成功后，认证信息写入本地存储并同步到 authStore；刷新页面时，RequireAuth 调用 verifyToken 恢复会话状态。该模块的实现重点是避免未认证用户直接进入业务页面，同时避免每个页面重复编写登录判断逻辑，因此把鉴权逻辑抽象到 React Router 外层组件中。",
        "workspace": "图谱工作区是前端最核心的功能模块。桌面端由左侧话题列表、中间 React Flow 图谱画布和右侧详情面板组成：左侧帮助用户切换根话题，中间画布负责展示节点和边，右侧详情面板展示当前选中节点的内容、作者、溯源、关联和审计信息。画布支持谱系视图和探索视图，前者强调从根话题到当前结论的演进关系，后者强调语义关联和分支扩展。移动端由于屏幕空间有限，不直接挂载复杂画布，而是使用 discussion-tree 接口把讨论呈现为嵌套阅读结构。",
        "api_state": "API 封装模块以 src/api/client.ts 为统一入口，自动附加 Authorization: Bearer <token>，并统一处理后端返回的 {code, message, data} 结构。业务接口按领域拆分为 auth、nodes、posts、decisions、associations、audit、profile、feed、events、search、follows 和 mutes 等文件，页面组件不直接拼接底层请求细节。状态管理模块按职责拆分 store：authStore 管理登录状态，graphStore 管理节点、边、选择态和图谱视图，sseStore 管理连接状态，homeStore 管理首页数据，discussionTreeStore 管理移动端讨论树，uiStore 管理侧边栏、右面板、命令面板和 toast。API 与 store 的分层让组件只关注业务交互，错误处理、token 处理和数据建模集中管理，也便于对表单、移动端组件、SSE 和 store 单独测试。",
        "sse": "实时更新模块通过 useSse.ts 请求 /api/events/stream。连接建立后，前端持续解析后端推送的事件，并根据事件类型决定更新策略：节点创建事件会补拉节点并加入图谱，边创建事件会补齐端点关系，决策完成事件会替换乐观节点，摘要生成和质量评分事件会刷新节点徽章。该模块的难点是保证实时性和一致性之间的平衡，既不能因为频繁刷新造成页面抖动，也不能因为缓存过旧导致图谱状态和后端不一致。",
        "home": [
            "首页与话题入口模块负责登录后的第一层信息组织。HomePage 进入页面后调用 graphStore 的 loadRhizomes 拉取根话题，用 HomeSidebar、HomeMainColumn 和 HomeRightRail 组成左侧导航、中间主列、右侧辅助栏的三栏结构；移动端则把侧边导航切换为底部弹层，避免小屏空间被固定侧栏占用。",
            "实现过程中，首页组件本身只承担装配职责，数据读取和界面状态分别交给 graphStore 与 uiStore 管理，同时接入 useSse 和 useCommandPalette，使首屏可以响应实时事件并打开全局检索。这个模块体现了项目的信息架构能力：用户不是直接面对复杂图谱画布，而是先通过根话题、动态信息和快捷入口理解系统当前状态，再进入具体工作区。",
        ],
        "post_form": [
            "发布与回复编辑模块负责把用户输入转化为后端可异步处理的发帖请求。PostForm 使用 MarkdownEditor 管理正文，根据 selectedNodeId 判断当前是发布新话题还是回复已有节点，并在界面上展示“正在回复”的目标摘要，用户也可以取消回复回到新发布状态。",
            "提交时，handleSubmit 会先去除 HTML 标签检查空内容，再校验当前登录用户，随后调用 createPost 传入 request_id、author_id、content 和 target_node_id。接口返回后前端显示“发布已排队”或“回复已排队”的 toast，不阻塞等待 AI 裁决完成，而是由后续 SSE 事件推动图谱更新。该流程展示了项目对异步业务的处理能力：用户交互保持及时反馈，复杂决策交给后端队列和实时事件链路完成。",
        ],
        "node_detail": [
            "节点详情与治理模块是图谱工作区中承载深度信息的关键页面。用户选中节点后，NodeDetailPanel 会根据 rightPanelPayload 打开右侧面板，展示节点类型、作者、创建时间、正文内容、AI 决策说明、质量状态和编排状态，并通过“详情、确权溯源、关联、审计”标签页组织不同维度的信息。",
            "实现过程中，面板会在打开时调用 fetchNode 获取最新节点数据，并用 graphStore.addNode 写回本地状态，避免详情面板展示过期缓存。关注、取消关注、屏蔽、取消屏蔽和摘要生成等操作都通过独立 API 完成，操作后立即更新 store 并给出 toast 反馈。该模块把节点从简单图形元素提升为可审计、可治理、可追踪的业务对象，是本项目区别于普通评论列表的重要亮点。",
        ],
        "mobile_tree": [
            "移动端讨论树模块解决了复杂图谱在小屏设备上难以阅读和操作的问题。GraphWorkspace 根据视口宽度选择桌面图谱或移动端讨论树；MobileDiscussionTreeView 根据路由参数确定根节点，没有指定根节点时先加载根话题并选择第一个话题进入。",
            "实现过程中，移动端视图通过 discussionTreeStore 加载 discussion-tree 数据，使用 Skeleton 展示加载态，加载完成后由 CommentTreeItem 递归渲染嵌套评论结构，并配合 MobileReplyComposer、LongPressMenu 和 ToastContainer 完成回复、长按菜单和反馈提示。页面在浏览器重新可见时自动 refreshTree，保证移动端返回页面后内容仍然及时。该模块体现了项目对响应式体验的深入处理：桌面保留图谱探索能力，移动端则转为更适合阅读和回复的线性树结构。",
        ],
        "command_palette": [
            "检索与命令面板模块用于解决大规模图谱中节点定位困难的问题。CommandPalette 打开后会聚合最近节点、关键词相似检索和当前节点向量相似检索，用户可以通过输入框、方向键和回车快速跳转到目标节点。",
            "实现过程中，模块对输入内容做 150ms 防抖，调用 searchSimilar 发起文本检索；当用户选择“查找相似节点”时，先通过 fetchEmbedding 取得当前节点向量，再提交向量相似检索。如果目标节点当前不在本地图谱中，模块会先调用 loadTopologyContext 补拉上下文，再 selectNode、requestFocusNode 并打开详情面板。该模块体现了项目对复杂信息检索的工程化处理，让图谱系统不仅能展示关系，也能高效定位和导航。",
        ],
        "settings": [
            "个人设置与界面偏好模块用于补齐系统的用户资料维护能力。SettingsPage 进入后调用 getMyProfile 加载当前用户资料，展示用户名、头像、展示名称和界面外观设置，并通过 RadiusModeToggle 支持界面圆角偏好切换。",
            "实现过程中，页面分别处理 loading、saving、loadError 和 message 状态，保存展示名称时调用 updateProfile 并把返回的 UserProfile 写回本地状态；头像上传交给 AvatarUpload 组件完成，认证用户名来自 authStore。该模块虽然不是图谱核心链路，但体现了应用完整性：系统不仅能完成讨论和治理，也提供了真实产品中必需的个人资料和界面偏好入口。",
        ],
        "module_placeholders": {
            "route": "效果图占位：登录页、登录后首页或未登录跳转登录页截图。",
            "workspace": "效果图占位：桌面端图谱工作区三栏布局截图，包含左侧话题列表、中间图谱画布和右侧详情面板。",
            "sse": "效果图占位：发帖后节点增量出现、边关系刷新或质量评分更新截图。",
        },
        "effect_captions": {
            "route": "图4-2 路由与鉴权模块效果图占位（请后期替换为真实页面截图）",
            "workspace": "图4-4 图谱工作区模块效果图占位（请后期替换为真实页面截图）",
            "sse": "图4-7 实时更新模块效果图占位（请后期替换为真实页面截图）",
        },
        "difficulties": "第一个难点是图谱数据复杂。节点、边、语义关联和审计记录来自不同接口，如果组件直接处理所有格式，代码会迅速膨胀。解决方法是在 API 层和 store 层完成数据转换，让画布组件只关心可渲染节点和边。第二个难点是桌面端和移动端交互差异大。复杂图谱适合大屏，而移动端更适合线性阅读，因此项目按视口选择桌面图谱或移动讨论树。第三个难点是实时事件多。项目通过事件类型分发和局部更新减少全量刷新，同时保留必要的补拉逻辑，保证用户看到的数据最终与后端一致。第四个难点是复杂业务入口多，如果所有操作都堆在画布上会影响学习成本，因此项目通过首页、命令面板、详情面板和设置页分担入口，让用户可以按任务路径进入相应功能。",
        "summary": "通过本项目的开发，可以系统理解 React 前端工程中函数组件、Hooks、路由管理、状态管理、接口封装、响应式布局、实时事件处理和测试验证等核心技术。项目优点是模块边界清晰、桌面与移动端体验分层、实时更新链路完整；不足是图谱场景对首次理解成本较高，真实页面截图、端到端测试和自动化部署说明仍可进一步补充。虽然提交文件名按课程要求保留 Vue.js，但报告内容已经按真实 React/Vite 代码实现编写。",
        "future": "后续可以从三个方向继续优化。第一，增加更多端到端测试，覆盖登录、发帖、图谱切换、移动端回复和 SSE 增量刷新等完整用户流程。第二，补充真实运行截图、关键 React 源码白底截图和部署截图，使课程报告中的“先文后图”材料更接近最终答辩展示。第三，继续优化 React Flow 大图谱性能，例如节点虚拟化、布局缓存、SSE 事件批处理和移动端数据预取，从而提升大规模讨论场景下的交互稳定性。",
    }


def vue_markdown(diagrams: dict[str, Path]) -> str:
    """Template-compliant report body for the real React frontend."""
    img = lambda key: f"images/{diagrams[key].name}"
    content = vue_react_report_content(frontend_package_versions())
    tech_stack = "\n\n".join(content["tech_stack"])
    functional = "\n\n".join(content["functional"])
    placeholders = content["module_placeholders"]
    return f"""# 《Vue.js应用开发》期末课程设计报告

项目名称：RhizoDelta 图谱化非线性讨论系统前端实现

代码仓库：{RHIZODELTA_REPO_URL}

本地开发目录：`{RHIZODELTA_LOCAL_PATH}/frontend`

## 1 项目概述

### 1.1 项目背景

{content['background']}

### 1.2 项目目标

{content['target']}

## 2 技术选型

### 2.1 前端技术栈

{tech_stack}

### 2.2 后端技术栈（可选）

{content['backend']}

## 3 需求分析

### 3.1 功能需求

{functional}

### 3.2 非功能性需求

{content['nonfunctional']}

## 4 详细设计

### 4.1 开发环境

{content['environment']}

![前端模块结构与运行链路]({img('frontend')})

### 4.2 路由与鉴权模块

{content['route']}

> {placeholders['route']}

![路由鉴权真实源码截图]({img('route_code')})

### 4.3 图谱工作区模块

{content['workspace']}

> {placeholders['workspace']}

![图谱工作区真实源码截图]({img('workspace_code')})

### 4.4 API 封装与状态管理模块

{content['api_state']}

![API 请求封装真实源码截图]({img('api_client_code')})

### 4.5 实时更新模块

{content['sse']}

> {placeholders['sse']}

![SSE 事件处理真实源码截图]({img('sse_code')})

### 4.6 首页与话题入口模块

{content['home'][0]}

{content['home'][1]}

![首页入口真实源码截图]({img('home_code')})

### 4.7 发布与回复编辑模块

{content['post_form'][0]}

{content['post_form'][1]}

![发布与回复编辑真实源码截图]({img('post_form_code')})

### 4.8 节点详情与治理模块

{content['node_detail'][0]}

{content['node_detail'][1]}

![节点详情治理真实源码截图]({img('node_detail_code')})

### 4.9 移动端讨论树模块

{content['mobile_tree'][0]}

{content['mobile_tree'][1]}

![移动端讨论树真实源码截图]({img('mobile_tree_code')})

### 4.10 检索与命令面板模块

{content['command_palette'][0]}

{content['command_palette'][1]}

![检索与命令面板真实源码截图]({img('command_palette_code')})

### 4.11 个人设置与界面偏好模块

{content['settings'][0]}

{content['settings'][1]}

![个人设置与界面偏好真实源码截图]({img('settings_code')})

### 4.x 难点与解决方案

{content['difficulties']}

## 5 总结与展望

### 5.1 项目总结

{content['summary']}

### 5.2 未来展望

{content['future']}
"""


def add_vue_content(doc: Document, diagrams: dict[str, Path]) -> None:
    content = vue_react_report_content(frontend_package_versions())
    placeholders = content["module_placeholders"]
    effect_captions = content["effect_captions"]

    remove_template_textboxes(doc)
    remove_template_body_media(doc)
    replace_vue_template_cover(doc)
    replace_vue_template_toc(doc)

    replace_heading(doc, "4.2 功能模块名称", "4.2 路由与鉴权模块")
    replace_heading(doc, "4.3 功能模块名称", "4.3 图谱工作区模块")
    replace_heading(doc, "4.4 功能模块名称", "4.4 API 封装与状态管理模块")
    replace_heading(doc, "4.5 功能模块名称", "4.5 实时更新模块")
    normalize_vue_template_headings(doc)

    replace_section_body(doc, "1.1 项目背景", "1.2 项目目标", [content["background"]])
    replace_section_body(doc, "1.2 项目目标", "2 技术选型", [content["target"]])
    replace_section_body(doc, "2.1 前端技术栈", "2.2 后端技术栈（可选）", content["tech_stack"])
    replace_section_body(doc, "2.2 后端技术栈（可选）", "3 需求分析", [content["backend"]])
    replace_section_body(doc, "3.1 功能需求", "3.2 非功能性需求", content["functional"])
    replace_section_body(doc, "3.2 非功能性需求", "4 详细设计", [content["nonfunctional"]])
    replace_section_body(
        doc,
        "4.1 开发环境",
        "4.2 路由与鉴权模块",
        [content["environment"]],
        [(diagrams["frontend"], "图4-1 前端模块结构与运行链路")],
    )
    replace_module_section(
        doc,
        "4.2 路由与鉴权模块",
        "4.3 图谱工作区模块",
        content["route"],
        placeholders["route"],
        effect_captions["route"],
        None,
        "图4-3 路由鉴权源码片段",
        VUE_CODE_SNIPPETS["route"],
    )
    replace_module_section(
        doc,
        "4.3 图谱工作区模块",
        "4.4 API 封装与状态管理模块",
        content["workspace"],
        placeholders["workspace"],
        effect_captions["workspace"],
        None,
        "图4-5 图谱工作区源码片段",
        VUE_CODE_SNIPPETS["workspace"],
    )
    replace_module_section(
        doc,
        "4.4 API 封装与状态管理模块",
        "4.5 实时更新模块",
        content["api_state"],
        None,
        None,
        None,
        "图4-6 API 请求封装源码片段",
        VUE_CODE_SNIPPETS["api_state"],
    )
    cursor = replace_module_section(
        doc,
        "4.5 实时更新模块",
        "4.6 ………",
        content["sse"],
        placeholders["sse"],
        effect_captions["sse"],
        None,
        "图4-8 SSE 事件处理源码片段",
        VUE_CODE_SNIPPETS["sse"],
    )
    delete_section_body(doc, "4.6 ………", "4.x 难点与解决方案")
    cursor = insert_module_after(cursor, "4.6 首页与话题入口模块", content["home"], None, "图4-9 首页入口源码片段", VUE_CODE_SNIPPETS["home"])
    cursor = insert_module_after(cursor, "4.7 发布与回复编辑模块", content["post_form"], None, "图4-10 发布与回复编辑源码片段", VUE_CODE_SNIPPETS["post_form"])
    cursor = insert_module_after(cursor, "4.8 节点详情与治理模块", content["node_detail"], None, "图4-11 节点详情治理源码片段", VUE_CODE_SNIPPETS["node_detail"])
    cursor = insert_module_after(cursor, "4.9 移动端讨论树模块", content["mobile_tree"], None, "图4-12 移动端讨论树源码片段", VUE_CODE_SNIPPETS["mobile_tree"])
    cursor = insert_module_after(cursor, "4.10 检索与命令面板模块", content["command_palette"], None, "图4-13 检索与命令面板源码片段", VUE_CODE_SNIPPETS["command_palette"])
    insert_module_after(cursor, "4.11 个人设置与界面偏好模块", content["settings"], None, "图4-14 个人设置与界面偏好源码片段", VUE_CODE_SNIPPETS["settings"])
    replace_section_body(doc, "4.x 难点与解决方案", "5 总结与展望", [content["difficulties"]])
    replace_section_body(doc, "5.1 项目总结", "5.2 未来展望", [content["summary"]])
    replace_section_body_to_end(doc, "5.2 未来展望", [content["future"]])


def html_doc(title: str, markdown_body: str, image_paths: Iterable[Path]) -> str:
    image_map = {}
    for path in image_paths:
        data_uri = "data:image/png;base64," + base64.b64encode(path.read_bytes()).decode("ascii")
        image_map[path.as_posix()] = data_uri
        image_map[path.name] = data_uri
        image_map[f"images/{path.name}"] = data_uri

    lines = []
    for raw in markdown_body.splitlines():
        line = raw.strip()
        if not line:
            lines.append("")
            continue
        if line.startswith("# "):
            lines.append(f"<h1>{html.escape(line[2:])}</h1>")
        elif line.startswith("## "):
            lines.append(f"<h2>{html.escape(line[3:])}</h2>")
        elif line.startswith("### "):
            lines.append(f"<h3>{html.escape(line[4:])}</h3>")
        elif line.startswith("- "):
            lines.append(f"<p>• {html.escape(line[2:])}</p>")
        elif line.startswith("> "):
            lines.append(f'<p class="placeholder">{html.escape(line[2:])}</p>')
        elif line.startswith("![") and "](" in line and line.endswith(")"):
            alt = line[2:line.find("]")]
            src = line[line.find("(") + 1:-1]
            data = image_map.get(src, src)
            lines.append(f'<p style="text-align:center"><img alt="{html.escape(alt)}" src="{data}" style="max-width:680px;width:100%"></p>')
        elif line.startswith("|"):
            lines.append(f"<p>{html.escape(line)}</p>")
        elif line.startswith("```"):
            lines.append("<hr>")
        else:
            lines.append(f"<p>{html.escape(line)}</p>")

    return """<!doctype html>
<html>
<head>
<meta charset="utf-8">
<title>{title}</title>
<style>
body {{ font-family: SimSun, serif; font-size: 14px; line-height: 1.65; }}
h1 {{ text-align: center; font-family: SimHei, sans-serif; }}
h2, h3 {{ font-family: SimHei, sans-serif; }}
p {{ text-indent: 2em; }}
p.placeholder {{ border: 1px solid #94a3b8; background: #f8fafc; padding: 20px; text-align: center; text-indent: 0; }}
img {{ border: 1px solid #ddd; }}
</style>
</head>
<body>
{body}
</body>
</html>
""".format(title=html.escape(title), body="\n".join(lines))


def build_report(
    out_dir: Path,
    course: str,
    doc_title: str,
    subtitle: str,
    markdown: str,
    image_paths: list[Path],
    content_builder,
    basename: str,
    doc_factory=None,
    convert_doc: bool = False,
) -> None:
    out_dir.mkdir(parents=True, exist_ok=True)
    (out_dir / "generated").mkdir(exist_ok=True)

    md_path = out_dir / f"{basename}.md"
    docx_path = out_dir / f"{basename}.docx"
    doc_path = out_dir / f"{basename}.doc"
    html_path = out_dir / "generated" / f"{basename}.html"

    md_path.write_text(markdown, encoding="utf-8")

    doc = doc_factory() if doc_factory else Document()
    set_doc_defaults(doc)
    if doc_factory is None:
        add_cover(doc, course, doc_title, subtitle)
        add_toc(doc)
    content_builder(doc)
    if doc_factory is None:
        for section in doc.sections:
            add_page_number(section)
    doc.save(docx_path)

    html_content = html_doc(doc_title, markdown, image_paths)
    html_path.write_text(html_content, encoding="utf-8")
    if convert_doc:
        libreoffice_convert(docx_path, "doc", out_dir)
    else:
        doc_path.write_text(html_content, encoding="utf-8")


def write_readme() -> None:
    readme = """# RhizoDelta 课程交付物

本目录下的两门课程报告均基于 RhizoDelta 当前代码与已 review 的项目文档生成。

- `javaee-web-app/`：JavaEE 企业级 Web 应用开发实战报告，侧重 Spring Boot 后端、Neo4j、RabbitMQ、JWT、SSE 和 AI 编排。
- `vuejs-app-dev/`：Vue.js 应用开发课程封面与文件名保留 Vue；正文按学校允许的实际项目内容填写，侧重 React/Vite 前端工程。

注意：报告中的个人姓名、学号、班级、指导教师保留为“请填写”，提交前需要在 Word/WPS 中替换。
"""
    (ROOT / "RhizoDelta课程交付物说明.md").write_text(readme, encoding="utf-8")


def main() -> None:
    JAVAEE_DIR.mkdir(exist_ok=True)
    VUE_DIR.mkdir(exist_ok=True)
    for course_dir in [JAVAEE_DIR, VUE_DIR]:
        (course_dir / "images").mkdir(exist_ok=True)

    shared_dir = Path("/tmp/finalexams-rhizodelta-report-assets")
    if shared_dir.exists():
        shutil.rmtree(shared_dir)
    (shared_dir / "images").mkdir(parents=True, exist_ok=True)

    diagrams = {
        "architecture": architecture_diagram(shared_dir),
        "graph": graph_model_diagram(shared_dir),
        "frontend": frontend_diagram(shared_dir),
        "flow": backend_flow_diagram(shared_dir),
    }

    # Copy diagrams into each course folder so markdown links remain local.
    java_diagrams = {}
    vue_diagrams = {}
    for key, path in diagrams.items():
        jpath = JAVAEE_DIR / "images" / path.name
        vpath = VUE_DIR / "images" / path.name
        shutil.copyfile(path, jpath)
        shutil.copyfile(path, vpath)
        java_diagrams[key] = jpath
        vue_diagrams[key] = vpath
    vue_diagrams["route_code"] = route_source_snapshot(VUE_DIR)
    vue_diagrams["sse_code"] = sse_source_snapshot(VUE_DIR)
    vue_diagrams["workspace_code"] = workspace_source_snapshot(VUE_DIR)
    vue_diagrams["api_client_code"] = api_client_source_snapshot(VUE_DIR)
    vue_diagrams.update(refresh_extra_vue_source_snapshots())

    java_md = javaee_markdown(java_diagrams)
    vue_md = vue_markdown(vue_diagrams)

    java_basename = "JavaEE企业级Web应用开发实战期末课程设计报告"
    vue_basename = "Vue.js应用开发期末课程设计报告"

    build_report(
        JAVAEE_DIR,
        "《JavaEE企业级Web应用开发实战》",
        "期末课程设计报告",
        "RhizoDelta 图谱化非线性讨论系统后端实现",
        java_md,
        list(java_diagrams.values()),
        lambda doc: add_javaee_content(doc, java_diagrams),
        java_basename,
    )
    build_report(
        VUE_DIR,
        "《Vue.js应用开发》",
        "期末课程设计报告",
        "RhizoDelta 图谱化非线性讨论系统前端实现",
        vue_md,
        list(vue_diagrams.values()),
        lambda doc: add_vue_content(doc, vue_diagrams),
        vue_basename,
        doc_factory=vue_template_document,
        convert_doc=True,
    )
    write_readme()
    print("Generated RhizoDelta course reports.")


if __name__ == "__main__":
    main()
