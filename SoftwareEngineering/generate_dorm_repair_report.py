#!/usr/bin/env python3
"""Generate the Software Engineering final report for the dorm repair system.

The script writes an offline-openable .docx report and embedded diagram images.
It expects python-docx and Pillow to be importable. In this Codex environment
they are installed under /tmp/finalexams-pydeps.
"""

from __future__ import annotations

import math
import os
import shutil
import sys
from pathlib import Path
from typing import Iterable, Sequence

sys.path.insert(0, "/tmp/finalexams-pydeps")

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT_DIR = ROOT / "generated"
IMG_DIR = OUT_DIR / "images"
DOCX_PATH = OUT_DIR / "宿舍报修管理系统软件工程导论期末考查报告.docx"
ASCII_DOCX_PATH = OUT_DIR / "dorm_repair_report_compatible.docx"
MD_PATH = OUT_DIR / "宿舍报修管理系统软件工程导论期末考查报告.md"

FONT_PATH = Path("/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc")
if not FONT_PATH.exists():
    FONT_PATH = Path("/usr/share/fonts/opentype/unifont/unifont.otf")


def ensure_dirs() -> None:
    IMG_DIR.mkdir(parents=True, exist_ok=True)


def font(size: int) -> ImageFont.FreeTypeFont:
    return ImageFont.truetype(str(FONT_PATH), size)


F_TITLE = font(34)
F_H = font(25)
F = font(21)
F_S = font(18)
F_XS = font(16)


def text_size(draw: ImageDraw.ImageDraw, text: str, fnt: ImageFont.FreeTypeFont) -> tuple[int, int]:
    if not text:
        return (0, 0)
    box = draw.textbbox((0, 0), text, font=fnt)
    return box[2] - box[0], box[3] - box[1]


def wrap_text(draw: ImageDraw.ImageDraw, text: str, max_width: int, fnt: ImageFont.FreeTypeFont) -> list[str]:
    lines: list[str] = []
    for raw in text.split("\n"):
        if not raw:
            lines.append("")
            continue
        current = ""
        for ch in raw:
            trial = current + ch
            if text_size(draw, trial, fnt)[0] <= max_width or not current:
                current = trial
            else:
                lines.append(current)
                current = ch
        if current:
            lines.append(current)
    return lines


def draw_center_text(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    text: str,
    fnt: ImageFont.FreeTypeFont = F,
    fill: str = "#172033",
    line_gap: int = 5,
) -> None:
    x1, y1, x2, y2 = box
    max_width = max(1, x2 - x1 - 18)
    lines = wrap_text(draw, text, max_width, fnt)
    heights = [text_size(draw, line, fnt)[1] for line in lines]
    total_h = sum(heights) + max(0, len(lines) - 1) * line_gap
    y = y1 + (y2 - y1 - total_h) / 2
    for line, h in zip(lines, heights):
        w, _ = text_size(draw, line, fnt)
        draw.text((x1 + (x2 - x1 - w) / 2, y), line, font=fnt, fill=fill)
        y += h + line_gap


def new_canvas(w: int, h: int, title: str) -> tuple[Image.Image, ImageDraw.ImageDraw]:
    img = Image.new("RGB", (w, h), "white")
    draw = ImageDraw.Draw(img)
    draw.rectangle((0, 0, w - 1, h - 1), outline="#D7DBE8", width=2)
    draw.text((w // 2 - text_size(draw, title, F_TITLE)[0] // 2, 28), title, font=F_TITLE, fill="#101827")
    return img, draw


def save(img: Image.Image, name: str) -> Path:
    path = IMG_DIR / name
    img.save(path)
    return path


def rect(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    text: str,
    fill: str = "#F6F8FB",
    outline: str = "#2E5EAA",
    width: int = 3,
    fnt: ImageFont.FreeTypeFont = F,
) -> None:
    draw.rounded_rectangle(box, radius=10, fill=fill, outline=outline, width=width)
    draw_center_text(draw, box, text, fnt)


def ellipse(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    text: str,
    fill: str = "#EEF7F0",
    outline: str = "#287A45",
    width: int = 3,
    fnt: ImageFont.FreeTypeFont = F,
) -> None:
    draw.ellipse(box, fill=fill, outline=outline, width=width)
    draw_center_text(draw, box, text, fnt)


def diamond(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    text: str,
    fill: str = "#FFF7E9",
    outline: str = "#B06B00",
    width: int = 3,
    fnt: ImageFont.FreeTypeFont = F_S,
) -> None:
    x1, y1, x2, y2 = box
    pts = ((x1 + x2) // 2, y1), (x2, (y1 + y2) // 2), ((x1 + x2) // 2, y2), (x1, (y1 + y2) // 2)
    draw.polygon(pts, fill=fill, outline=outline)
    for offset in range(width):
        pts2 = (
            ((x1 + x2) // 2, y1 + offset),
            (x2 - offset, (y1 + y2) // 2),
            ((x1 + x2) // 2, y2 - offset),
            (x1 + offset, (y1 + y2) // 2),
        )
        draw.line(pts2 + (pts2[0],), fill=outline, width=1)
    draw_center_text(draw, box, text, fnt)


def parallelogram(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    text: str,
    fill: str = "#F6F8FB",
    outline: str = "#2E5EAA",
    width: int = 3,
    fnt: ImageFont.FreeTypeFont = F,
) -> None:
    x1, y1, x2, y2 = box
    skew = min(45, max(20, (x2 - x1) // 8))
    pts = ((x1 + skew, y1), (x2, y1), (x2 - skew, y2), (x1, y2))
    draw.polygon(pts, fill=fill, outline=outline)
    for offset in range(width):
        pts2 = ((x1 + skew + offset, y1 + offset), (x2 - offset, y1 + offset), (x2 - skew - offset, y2 - offset), (x1 + offset, y2 - offset))
        draw.line(pts2 + (pts2[0],), fill=outline, width=1)
    draw_center_text(draw, box, text, fnt)


def terminator(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str) -> None:
    draw.rounded_rectangle(box, radius=(box[3] - box[1]) // 2, fill="#EEF7F0", outline="#287A45", width=3)
    draw_center_text(draw, box, text, F)


def process_box(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    text: str,
    fill: str = "#F6F8FB",
    outline: str = "#2E5EAA",
    fnt: ImageFont.FreeTypeFont = F,
) -> None:
    draw.rectangle(box, fill=fill, outline=outline, width=3)
    draw_center_text(draw, box, text, fnt)


def dfd_store(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str) -> None:
    x1, y1, x2, y2 = box
    draw.rectangle(box, fill="#F6F8FB", outline="#6B7280", width=3)
    draw.line((x1 + 18, y1, x1 + 18, y2), fill="#6B7280", width=3)
    draw.line((x2 - 18, y1, x2 - 18, y2), fill="#6B7280", width=3)
    draw_center_text(draw, (x1 + 22, y1, x2 - 22, y2), text, F_S)


def module_line(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], color: str = "#475467") -> None:
    draw.line((start, end), fill=color, width=3)


def arrow(
    draw: ImageDraw.ImageDraw,
    start: tuple[int, int],
    end: tuple[int, int],
    text: str | None = None,
    color: str = "#344054",
    width: int = 3,
    fnt: ImageFont.FreeTypeFont = F_XS,
    text_offset: tuple[int, int] = (0, -24),
) -> None:
    draw.line((start, end), fill=color, width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    length = 16
    spread = math.pi / 7
    p1 = (end[0] - length * math.cos(angle - spread), end[1] - length * math.sin(angle - spread))
    p2 = (end[0] - length * math.cos(angle + spread), end[1] - length * math.sin(angle + spread))
    draw.polygon((end, p1, p2), fill=color)
    if text:
        mx = (start[0] + end[0]) // 2 + text_offset[0]
        my = (start[1] + end[1]) // 2 + text_offset[1]
        w, h = text_size(draw, text, fnt)
        draw.rounded_rectangle((mx - 6, my - 4, mx + w + 6, my + h + 4), radius=5, fill="white")
        draw.text((mx, my), text, font=fnt, fill=color)


def dashed_line(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], color: str = "#475467", width: int = 2) -> None:
    x1, y1 = start
    x2, y2 = end
    dx, dy = x2 - x1, y2 - y1
    dist = max(1, math.hypot(dx, dy))
    steps = int(dist // 20)
    for i in range(steps):
        if i % 2 == 0:
            a = i / steps
            b = min(1, (i + 1) / steps)
            draw.line((x1 + dx * a, y1 + dy * a, x1 + dx * b, y1 + dy * b), fill=color, width=width)


def dashed_arrow(
    draw: ImageDraw.ImageDraw,
    start: tuple[int, int],
    end: tuple[int, int],
    text: str | None = None,
    color: str = "#475467",
    fnt: ImageFont.FreeTypeFont = F_XS,
    text_offset: tuple[int, int] = (0, -24),
) -> None:
    dashed_line(draw, start, end, color=color, width=2)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    length = 15
    spread = math.pi / 7
    p1 = (end[0] - length * math.cos(angle - spread), end[1] - length * math.sin(angle - spread))
    p2 = (end[0] - length * math.cos(angle + spread), end[1] - length * math.sin(angle + spread))
    draw.polygon((end, p1, p2), fill=color)
    if text:
        mx = (start[0] + end[0]) // 2 + text_offset[0]
        my = (start[1] + end[1]) // 2 + text_offset[1]
        w, h = text_size(draw, text, fnt)
        draw.rounded_rectangle((mx - 8, my - 4, mx + w + 8, my + h + 4), radius=5, fill="white")
        draw.text((mx, my), text, font=fnt, fill=color)


def cylinder(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str) -> None:
    x1, y1, x2, y2 = box
    draw.rectangle((x1, y1 + 16, x2, y2 - 16), fill="#F2F4F7", outline="#6B7280", width=3)
    draw.ellipse((x1, y1, x2, y1 + 32), fill="#F2F4F7", outline="#6B7280", width=3)
    draw.arc((x1, y2 - 32, x2, y2), 0, 180, fill="#6B7280", width=3)
    draw.line((x1, y1 + 16, x1, y2 - 16), fill="#6B7280", width=3)
    draw.line((x2, y1 + 16, x2, y2 - 16), fill="#6B7280", width=3)
    draw_center_text(draw, box, text, F_S)


def actor(draw: ImageDraw.ImageDraw, center: tuple[int, int], label: str) -> None:
    x, y = center
    draw.ellipse((x - 18, y - 50, x + 18, y - 14), outline="#111827", width=3)
    draw.line((x, y - 14, x, y + 46), fill="#111827", width=3)
    draw.line((x - 34, y + 8, x + 34, y + 8), fill="#111827", width=3)
    draw.line((x, y + 46, x - 28, y + 88), fill="#111827", width=3)
    draw.line((x, y + 46, x + 28, y + 88), fill="#111827", width=3)
    w, _ = text_size(draw, label, F)
    draw.text((x - w / 2, y + 96), label, font=F, fill="#111827")


def draw_class_diagram() -> Path:
    img, draw = new_canvas(1800, 1220, "宿舍报修管理系统 UML 类图")

    classes = {
        "学生 Student": (80, 130, 470, 380, ["学号 studentId", "姓名 name", "宿舍号 dormNo", "联系方式 phone"], ["提交报修()", "查询进度()", "评价维修()"]),
        "报修工单 RepairOrder": (700, 130, 1120, 430, ["工单号 orderId", "故障类型 type", "故障描述 description", "状态 status", "提交时间 createdAt"], ["创建工单()", "分派维修员()", "更新状态()", "关闭工单()"]),
        "维修员 Maintainer": (1340, 130, 1730, 380, ["工号 maintainerId", "姓名 name", "专业技能 skill", "联系电话 phone"], ["接单()", "填写维修结果()", "申请退回()"]),
        "宿舍 Dormitory": (80, 650, 470, 880, ["宿舍号 dormNo", "楼栋 building", "房间号 roomNo", "床位数 capacity"], ["查询宿舍信息()", "维护宿舍资料()"]),
        "派工记录 DispatchRecord": (700, 650, 1120, 910, ["派工号 dispatchId", "派工时间 dispatchTime", "到场时间 arriveTime", "完成时间 finishTime"], ["记录派工()", "记录完成()", "生成统计()"]),
        "管理员 Admin": (1340, 650, 1730, 880, ["账号 adminId", "姓名 name", "角色 role"], ["审核工单()", "派工()", "维护基础数据()", "统计分析()"]),
    }

    def class_box(box: tuple[int, int, int, int], name: str, attrs: list[str], methods: list[str]) -> None:
        x1, y1, x2, y2 = box
        draw.rounded_rectangle(box, radius=8, fill="#F8FAFC", outline="#365A9C", width=3)
        draw.rectangle((x1, y1, x2, y1 + 48), fill="#EAF1FF", outline="#365A9C", width=3)
        draw_center_text(draw, (x1, y1, x2, y1 + 48), name, F_H)
        split = y1 + 48 + 28 * len(attrs) + 18
        draw.line((x1, split, x2, split), fill="#365A9C", width=2)
        y = y1 + 62
        for item in attrs:
            draw.text((x1 + 18, y), "+ " + item, font=F_S, fill="#172033")
            y += 28
        y = split + 14
        for item in methods:
            draw.text((x1 + 18, y), "+ " + item, font=F_S, fill="#172033")
            y += 28

    for name, (x1, y1, x2, y2, attrs, methods) in classes.items():
        class_box((x1, y1, x2, y2), name, attrs, methods)

    def assoc(start: tuple[int, int], end: tuple[int, int], label: str, a_mult: str, b_mult: str, offset: tuple[int, int] = (0, -24)) -> None:
        draw.line((start, end), fill="#344054", width=3)
        mx = (start[0] + end[0]) // 2 + offset[0]
        my = (start[1] + end[1]) // 2 + offset[1]
        w, h = text_size(draw, label, F_XS)
        draw.rounded_rectangle((mx - 6, my - 4, mx + w + 6, my + h + 4), radius=5, fill="white")
        draw.text((mx, my), label, font=F_XS, fill="#344054")
        draw.text((start[0] + 8, start[1] - 26), a_mult, font=F_XS, fill="#344054")
        draw.text((end[0] - 48, end[1] - 26), b_mult, font=F_XS, fill="#344054")

    assoc((470, 255), (700, 255), "提交", "1", "0..*")
    assoc((1120, 255), (1340, 255), "分派给", "0..*", "0..1")
    assoc((275, 650), (275, 380), "入住", "1", "1..*")
    assoc((910, 430), (910, 650), "产生", "1", "0..1")
    assoc((1340, 760), (1120, 760), "管理", "1", "0..*", offset=(-20, -24))
    assoc((1340, 330), (1120, 720), "处理", "1", "0..*", offset=(-90, -10))
    assoc((470, 760), (700, 760), "关联宿舍", "1", "0..*")

    draw.text((80, 1070), "说明：类图使用三栏类框、关联名称和多重度；学生提交工单，管理员派工，维修员处理并形成派工记录。", font=F, fill="#172033")
    return save(img, "01_uml_class_diagram.png")


def draw_context_dfd() -> Path:
    img, draw = new_canvas(1600, 960, "顶层 DFD：宿舍报修管理系统")
    rect(draw, (90, 190, 330, 310), "学生", "#FFF7E9", "#B06B00")
    rect(draw, (90, 620, 330, 740), "宿管/管理员", "#FFF7E9", "#B06B00")
    rect(draw, (1260, 190, 1500, 310), "维修员", "#FFF7E9", "#B06B00")
    rect(draw, (1260, 620, 1500, 740), "后勤管理部门", "#FFF7E9", "#B06B00")
    ellipse(draw, (560, 330, 1040, 620), "P0\n宿舍报修管理系统", "#EEF7FF", "#275EAD", fnt=F_H)

    arrow(draw, (330, 250), (560, 410), "报修信息")
    arrow(draw, (560, 480), (330, 260), "受理结果/进度")
    arrow(draw, (330, 680), (560, 535), "审核规则/基础数据")
    arrow(draw, (560, 555), (330, 700), "统计报表")
    arrow(draw, (1040, 410), (1260, 250), "派工单")
    arrow(draw, (1260, 270), (1040, 490), "维修反馈")
    arrow(draw, (1040, 545), (1260, 680), "维修统计")
    arrow(draw, (1260, 700), (1040, 575), "管理要求")
    return save(img, "02_context_dfd.png")


def draw_level0_dfd() -> Path:
    img, draw = new_canvas(2100, 1450, "0 层 DFD：宿舍报修管理系统")
    rect(draw, (70, 170, 270, 270), "学生", "#FFF7E9", "#B06B00", fnt=F_S)
    rect(draw, (70, 1040, 270, 1140), "管理员", "#FFF7E9", "#B06B00", fnt=F_S)
    rect(draw, (1830, 170, 2030, 270), "维修员", "#FFF7E9", "#B06B00", fnt=F_S)
    rect(draw, (1830, 1040, 2030, 1140), "后勤部门", "#FFF7E9", "#B06B00", fnt=F_S)

    ellipse(draw, (430, 150, 750, 300), "P1\n报修申请受理", "#EEF7FF", "#275EAD", fnt=F)
    ellipse(draw, (430, 580, 750, 730), "P2\n工单审核", "#EEF7FF", "#275EAD", fnt=F)
    ellipse(draw, (930, 580, 1250, 730), "P3\n派工处理", "#EEF7FF", "#275EAD", fnt=F)
    ellipse(draw, (1380, 150, 1700, 300), "P4\n维修反馈确认", "#EEF7FF", "#275EAD", fnt=F)
    ellipse(draw, (1380, 1030, 1700, 1180), "P5\n统计查询", "#EEF7FF", "#275EAD", fnt=F)

    dfd_store(draw, (430, 960, 760, 1080), "D1 学生/宿舍档案")
    dfd_store(draw, (880, 960, 1210, 1080), "D2 报修工单库")
    dfd_store(draw, (1320, 960, 1650, 1080), "D3 派工/维修记录")

    arrow(draw, (270, 220), (430, 225), "报修申请")
    arrow(draw, (750, 225), (1380, 225), "已登记工单")
    arrow(draw, (1540, 300), (1540, 1030), "完成信息", text_offset=(15, -10))
    arrow(draw, (1830, 220), (1700, 225), "维修结果")
    arrow(draw, (1700, 225), (1830, 220), "派工任务")
    arrow(draw, (590, 300), (590, 580), "待审核工单")
    arrow(draw, (750, 655), (930, 655), "有效工单")
    arrow(draw, (1250, 655), (1830, 250), "派工单")
    arrow(draw, (1830, 1085), (1700, 1105), "报表需求")
    arrow(draw, (1380, 1105), (270, 1090), "统计报表")
    arrow(draw, (160, 1040), (520, 730), "审核意见")
    arrow(draw, (590, 730), (590, 960), "查验档案")
    arrow(draw, (1035, 730), (1035, 960), "更新工单")
    arrow(draw, (1540, 300), (1485, 960), "写入维修记录")
    arrow(draw, (1210, 1020), (1380, 1085), "工单数据")
    arrow(draw, (1650, 1020), (1380, 1120), "维修数据")
    arrow(draw, (880, 1000), (750, 690), "读取工单")
    arrow(draw, (760, 1020), (430, 665), "档案数据")

    draw.text((70, 1320), "DFD 说明：矩形为外部实体，椭圆为加工，双竖线矩形为数据存储，箭头为数据流；未混用流程图控制符号。", font=F, fill="#172033")
    return save(img, "03_level0_dfd.png")


def draw_state_diagram() -> Path:
    img, draw = new_canvas(1780, 980, "核心业务实体状态图：报修工单")
    states = {
        "待提交": (80, 410, 260, 520),
        "待审核": (390, 410, 590, 520),
        "待派工": (720, 410, 920, 520),
        "维修中": (1050, 410, 1250, 520),
        "待确认": (1380, 410, 1580, 520),
        "已完成": (1380, 690, 1580, 800),
        "已驳回": (390, 690, 590, 800),
        "已取消": (80, 690, 260, 800),
    }
    draw.ellipse((35, 445, 75, 485), fill="#111827")
    for label, box in states.items():
        rect(draw, box, label, "#F6F8FB", "#365A9C", fnt=F)
    arrow(draw, (75, 465), (80, 465), None)
    arrow(draw, (260, 465), (390, 465), "提交报修")
    arrow(draw, (590, 465), (720, 465), "审核通过")
    arrow(draw, (920, 465), (1050, 465), "分派维修员")
    arrow(draw, (1250, 465), (1380, 465), "填写完成结果")
    arrow(draw, (1480, 520), (1480, 690), "学生确认")
    arrow(draw, (170, 520), (170, 690), "未提交前取消", text_offset=(15, -8))
    arrow(draw, (490, 520), (490, 690), "信息不完整/不合规")
    arrow(draw, (1380, 740), (1250, 500), "要求返修", text_offset=(-85, -12))
    draw.ellipse((1665, 725, 1725, 785), outline="#111827", width=4)
    draw.ellipse((1680, 740, 1710, 770), fill="#111827")
    arrow(draw, (1580, 745), (1665, 755), "归档")
    draw.text((80, 885), "状态说明：报修工单从提交、审核、派工、维修、确认到归档闭环；驳回、取消、返修为异常或补充路径。", font=F, fill="#172033")
    return save(img, "04_repair_order_state.png")


def draw_structure_chart() -> Path:
    img, draw = new_canvas(1900, 1180, "由 0 层 DFD 转换得到的系统结构图")
    rect(draw, (745, 120, 1155, 210), "宿舍报修管理系统", "#EAF1FF", "#275EAD", fnt=F_H)
    level1 = [
        ("报修申请模块", (80, 360, 360, 450)),
        ("工单审核模块", (430, 360, 710, 450)),
        ("派工处理模块", (780, 360, 1060, 450)),
        ("维修反馈模块", (1130, 360, 1410, 450)),
        ("统计查询模块", (1480, 360, 1760, 450)),
    ]
    root_x, root_y = 950, 210
    bus_y = 300
    module_line(draw, (root_x, root_y), (root_x, bus_y))
    module_line(draw, (220, bus_y), (1620, bus_y))
    for name, box in level1:
        rect(draw, box, name, "#F6F8FB", "#365A9C")
        cx = (box[0] + box[2]) // 2
        module_line(draw, (cx, bus_y), (cx, box[1]))
    subs = {
        "报修申请模块": ["录入报修信息", "校验宿舍档案", "生成报修工单"],
        "工单审核模块": ["读取待审工单", "判断是否受理", "退回或提交派工"],
        "派工处理模块": ["匹配维修类型", "选择维修员", "生成派工记录"],
        "维修反馈模块": ["登记维修结果", "学生确认评价", "关闭或返修"],
        "统计查询模块": ["按状态统计", "按宿舍/类型查询", "导出维修报表"],
    }
    for name, box in level1:
        x1, y1, x2, y2 = box
        cx = (x1 + x2) // 2
        module_line(draw, (cx, y2), (cx, 610))
        for i, sub in enumerate(subs[name]):
            top = 610 + i * 120
            sbox = (cx - 125, top, cx + 125, top + 82)
            rect(draw, sbox, sub, "#FFFFFF", "#7A8AA0", width=2, fnt=F_S)
            if i > 0:
                module_line(draw, (cx, top - 38), (cx, top), "#7A8AA0")
    draw.text((80, 1060), "转换依据：0 层 DFD 中 P1-P5 映射为一级功能模块，数据存储访问通过模块接口完成。", font=F, fill="#172033")
    return save(img, "05_structure_chart.png")


def draw_flowchart() -> Path:
    img, draw = new_canvas(1500, 1900, "详细设计图 1：报修申请受理模块程序流程图")
    boxes = {
        "start": (610, 130, 890, 220),
        "input": (520, 300, 980, 395),
        "complete": (520, 480, 980, 575),
        "query": (520, 680, 980, 775),
        "valid_student": (520, 860, 980, 955),
        "loop": (520, 1040, 980, 1135),
        "urgent": (520, 1220, 980, 1315),
        "save": (520, 1400, 980, 1495),
        "notify": (520, 1580, 980, 1675),
        "end": (610, 1760, 890, 1850),
    }
    terminator(draw, boxes["start"], "开始")
    parallelogram(draw, boxes["input"], "输入宿舍号、故障类型、描述、附件")
    diamond(draw, boxes["complete"], "信息是否完整？")
    parallelogram(draw, (1040, 490, 1370, 565), "输出：提示补全信息", fnt=F_S)
    process_box(draw, boxes["query"], "查询学生与宿舍档案")
    diamond(draw, boxes["valid_student"], "档案是否匹配？")
    parallelogram(draw, (160, 870, 455, 945), "输出：拒绝提交并提示原因", fnt=F_S)
    process_box(draw, boxes["loop"], "FOR 每个附件：检查格式与大小")
    diamond(draw, boxes["urgent"], "是否紧急故障？")
    process_box(draw, (1040, 1230, 1370, 1305), "设置优先级为高", fnt=F_S)
    process_box(draw, (160, 1230, 455, 1305), "设置优先级为普通", fnt=F_S)
    process_box(draw, boxes["save"], "生成工单并写入工单库")
    parallelogram(draw, boxes["notify"], "输出：通知管理员待审核")
    terminator(draw, boxes["end"], "结束")
    centers = {k: ((v[0] + v[2]) // 2, (v[1] + v[3]) // 2) for k, v in boxes.items()}
    for a, b in [("start", "input"), ("input", "complete"), ("query", "valid_student"), ("loop", "urgent"), ("save", "notify"), ("notify", "end")]:
        arrow(draw, (centers[a][0], boxes[a][3]), (centers[b][0], boxes[b][1]), None)
    arrow(draw, (750, boxes["complete"][3]), (750, boxes["query"][1]), "是")
    arrow(draw, (980, 528), (1040, 528), "否")
    arrow(draw, (1205, 490), (1205, 340), None)
    arrow(draw, (1205, 340), (980, 340), "重新填写")
    arrow(draw, (520, 907), (455, 907), "否")
    arrow(draw, (750, boxes["valid_student"][3]), (750, boxes["loop"][1]), "是")
    arrow(draw, (750, boxes["urgent"][3]), (750, boxes["save"][1]), "汇合")
    arrow(draw, (980, 1267), (1040, 1267), "是")
    arrow(draw, (520, 1267), (455, 1267), "否")
    arrow(draw, (1205, 1305), (810, 1400), None)
    arrow(draw, (308, 1305), (690, 1400), None)
    arrow(draw, (308, 945), (308, 1810), "结束提交")
    arrow(draw, (308, 1810), (610, 1810), None)
    return save(img, "06_acceptance_flowchart.png")


def draw_ns_chart() -> Path:
    img, draw = new_canvas(1500, 1500, "详细设计图 2：派工处理模块 N-S 图")
    outer = (180, 160, 1320, 1380)
    draw.rectangle(outer, fill="#FFFFFF", outline="#365A9C", width=4)
    def ns_box(box: tuple[int, int, int, int], text: str, fill: str = "#FFFFFF", fnt: ImageFont.FreeTypeFont = F) -> None:
        draw.rectangle(box, fill=fill, outline="#365A9C", width=3)
        draw_center_text(draw, box, text, fnt)

    ns_box((180, 160, 1320, 280), "读取待派工工单列表")
    ns_box((180, 280, 1320, 400), "按紧急程度和提交时间排序")

    loop = (180, 400, 1320, 1180)
    draw.rectangle(loop, fill="#FFFFFF", outline="#365A9C", width=3)
    ns_box((180, 400, 1320, 480), "FOR 每一张待派工工单", "#EAF1FF")
    ns_box((230, 520, 1270, 630), "根据故障类型筛选具备技能且空闲的维修员", fnt=F_S)

    # N-S branch block: condition header + two structured branch columns.
    cond = (230, 670, 1270, 1080)
    draw.rectangle(cond, fill="#FFFFFF", outline="#7A8AA0", width=3)
    draw.rectangle((230, 670, 1270, 740), fill="#FFF7E9", outline="#7A8AA0", width=3)
    draw_center_text(draw, (230, 670, 1270, 740), "IF 存在可派维修员", F_S)
    draw.line((750, 740, 750, 1080), fill="#7A8AA0", width=3)
    draw.rectangle((230, 740, 750, 800), fill="#F6F8FB", outline="#7A8AA0", width=2)
    draw.rectangle((750, 740, 1270, 800), fill="#F6F8FB", outline="#7A8AA0", width=2)
    draw_center_text(draw, (230, 740, 750, 800), "THEN", F_S)
    draw_center_text(draw, (750, 740, 1270, 800), "ELSE", F_S)
    draw_center_text(draw, (250, 820, 730, 1060), "选择负载最低维修员\n生成派工单\n更新工单状态为“维修中”", F_S)
    draw_center_text(draw, (770, 820, 1250, 1060), "工单标记为“待协调”\n通知管理员人工处理\n保留待派工状态", F_S)

    ns_box((180, 1180, 1320, 1290), "保存派工记录并通知维修员、学生")
    ns_box((180, 1290, 1320, 1380), "结束")
    return save(img, "07_dispatch_ns_chart.png")


def draw_use_case() -> Path:
    img, draw = new_canvas(1900, 1220, "宿舍报修管理系统 UML 用例图")
    actor(draw, (130, 260), "学生")
    actor(draw, (130, 760), "管理员")
    actor(draw, (1750, 260), "维修员")
    actor(draw, (1750, 760), "后勤部门")

    boundary = (360, 125, 1540, 1080)
    draw.rounded_rectangle(boundary, radius=16, outline="#365A9C", width=4)
    draw.text((390, 145), "系统边界：宿舍报修管理系统", font=F_H, fill="#172033")
    cases = {
        "提交报修": (470, 240, 720, 325),
        "查询进度": (470, 390, 720, 475),
        "确认评价": (470, 540, 720, 625),
        "审核工单": (850, 240, 1100, 325),
        "派工": (850, 390, 1100, 475),
        "维护基础数据": (850, 540, 1100, 625),
        "接收派工": (1210, 240, 1460, 325),
        "填写维修结果": (1210, 390, 1460, 475),
        "申请返修说明": (1210, 540, 1460, 625),
        "统计查询": (850, 760, 1100, 845),
        "导出报表": (1210, 760, 1460, 845),
        "填写报修信息": (470, 760, 720, 845),
    }
    for name, box in cases.items():
        ellipse(draw, box, name, "#F6F8FB", "#365A9C", fnt=F_S)
    # Associations.
    for target in ["提交报修", "查询进度", "确认评价"]:
        x1, y1, _, y2 = cases[target]
        draw.line((190, 260, x1, (y1 + y2) // 2), fill="#344054", width=2)
    for target in ["审核工单", "派工", "维护基础数据", "统计查询"]:
        x1, y1, _, y2 = cases[target]
        draw.line((190, 760, x1, (y1 + y2) // 2), fill="#344054", width=2)
    for target in ["接收派工", "填写维修结果", "申请返修说明"]:
        _, y1, x2, y2 = cases[target]
        draw.line((1710, 260, x2, (y1 + y2) // 2), fill="#344054", width=2)
    for target in ["统计查询", "导出报表"]:
        _, y1, x2, y2 = cases[target]
        draw.line((1710, 760, x2, (y1 + y2) // 2), fill="#344054", width=2)
    dashed_arrow(draw, (595, 325), (595, 760), "<<include>>", text_offset=(15, -8))
    dashed_arrow(draw, (1210, 800), (1100, 800), "<<extend>>", text_offset=(-95, -30))
    dashed_arrow(draw, (1335, 540), (1335, 475), "<<extend>>", text_offset=(18, -6))
    return save(img, "08_use_case_diagram.png")


def draw_all_diagrams() -> dict[str, Path]:
    return {
        "类图": draw_class_diagram(),
        "顶层DFD": draw_context_dfd(),
        "0层DFD": draw_level0_dfd(),
        "状态图": draw_state_diagram(),
        "系统结构图": draw_structure_chart(),
        "流程图": draw_flowchart(),
        "NS图": draw_ns_chart(),
        "用例图": draw_use_case(),
    }


PROJECT_PURPOSE = [
    "巩固软件工程核心知识点：将课堂所学的需求分析、总体设计、详细设计、编码、测试等理论知识，转化为实际操作，熟练掌握 DFD、UML 图（类图、状态图、用例图）、系统结构图、详细设计图形（流程图 / 盒图 / PAD 图）等工具的使用。",
    "培养完整的软件设计思维：学会从需求出发，逐步推进系统设计、编码与测试，理解各阶段的衔接逻辑，建立“需求—设计—编码—测试”的完整闭环思维，避免设计与实现脱节。",
    "提升实操与问题解决能力：通过自选系统、完成各环节设计与编码，锻炼模块划分、逻辑梳理、测试用例设计的实操能力，学会解决设计过程中出现的图形规范、逻辑矛盾等实际问题。",
]


TEMPLATE_REQUIREMENTS = [
    "自选系统限定为中小型事务管理信息系统，从以下四类中选择其一：图书借阅系统、学生教务管理系统、超市订单管理系统、宿舍报修管理系统，禁止自定义其他简单系统。",
    "所有图形（E-R图、类图、DFD、状态图、系统结构图、流程图/盒图/PAD图、用例图）需符合软件工程绘图规范，标注清晰、逻辑严谨。",
    "所有设计内容需前后一致，编码、测试需对应前面设计的模块/功能，不得脱节。",
]


REPORT_SECTIONS: list[tuple[str, list[str]]] = [
    (
        "一、需求分析",
        [
            "本报告选择“宿舍报修管理系统”作为分析与设计对象。该系统面向高校宿舍日常维修管理场景，主要解决学生报修渠道分散、维修派工不透明、处理进度难以追踪、后勤统计依赖人工整理等问题。系统的核心业务围绕“学生提交报修—管理员审核—维修员处理—学生确认评价—后勤统计分析”展开，是典型的中小型事务管理信息系统。",
            "系统采用增量模型进行开发。宿舍报修业务需求相对清晰，但在实际运行中可能逐步增加移动端提醒、维修材料管理、满意度统计等功能。增量模型可以先完成报修、审核、派工、反馈这一主流程，再逐步扩展统计分析、评价管理和基础数据维护等功能。与瀑布模型相比，增量模型便于分阶段交付和根据师生反馈调整；与螺旋模型相比，本系统规模较小、风险较低，不需要投入过高的风险分析成本。",
            "系统核心功能包括：1）报修申请管理：学生填写宿舍号、故障类型、故障描述和附件后提交报修申请，系统校验学生与宿舍档案并生成工单；2）工单审核与派工：管理员审核报修信息，判断是否受理，并根据故障类型、紧急程度和维修员负载进行派工；3）维修处理与反馈：维修员接收派工任务，到场处理后填写维修结果，学生确认后关闭工单；4）统计查询：管理员和后勤部门按照工单状态、故障类型、楼栋宿舍和维修时长查询并导出统计结果。",
            "参与者包括学生、宿管或管理员、维修员和后勤管理部门。主要数据对象包括学生、宿舍、报修工单、派工记录、维修员和管理员。核心业务实体为“报修工单”，其状态变化贯穿系统主流程。",
        ],
    ),
    (
        "二、总体设计",
        [
            "根据需求分析阶段的 0 层 DFD，系统可以采用事务型数据流设计方法转换为层次化功能结构。报修申请、工单审核、派工处理、维修反馈、统计查询分别对应 0 层 DFD 中的 P1 至 P5 加工。系统顶层模块负责协调各功能模块，数据存储访问通过统一接口完成，从而保证工单状态、派工记录和统计数据的一致性。",
            "模块接口设计如下：报修申请模块输出待审核工单；工单审核模块读取工单和学生宿舍档案，输出有效工单或驳回结果；派工处理模块读取有效工单和维修员信息，输出派工记录；维修反馈模块更新维修记录和工单状态；统计查询模块读取工单库与维修记录，生成统计报表。",
        ],
    ),
    (
        "三、详细设计",
        [
            "详细设计选择两个包含分支与循环逻辑的模块。第一个模块为“报修申请受理模块”，使用程序流程图描述。该模块需要判断学生填写的信息是否完整、学生宿舍档案是否匹配，并循环检查附件格式和数量，最后根据故障类型设置紧急程度并生成工单。",
            "第二个模块为“派工处理模块”，使用 N-S 图描述。该模块先读取待派工工单并按紧急程度排序，然后对每一张工单循环处理；在循环体内根据故障类型筛选维修员，如果没有可派维修员，则工单进入待协调状态，否则选择当前负载最低的维修员并生成派工单。",
        ],
    ),
    (
        "四、程序编码",
        [
            "本节选择“报修申请受理模块”编写 PDL 伪代码。伪代码与第三部分程序流程图保持一致，包含输入校验、宿舍档案匹配、附件循环检查、紧急程度判断、工单生成和通知管理员等步骤。",
        ],
    ),
    (
        "五、软件测试",
        [
            "测试对象选择“报修申请受理模块”。该模块包含嵌套分支和附件循环检查，适合同时进行白盒测试与黑盒测试。白盒测试依据 PDL 伪代码设计，重点覆盖语句覆盖、判定-条件覆盖、条件组合覆盖和独立路径覆盖；黑盒测试从用户输入角度进行等价类划分与边界值分析。",
            "用例图覆盖系统主要参与者与核心用例，可直接作为系统测试和验收测试的功能依据。",
        ],
    ),
]


PDL = """PROCEDURE SubmitRepair(studentId, dormNo, faultType, description, attachments)
    IF studentId is empty OR dormNo is empty OR faultType is empty OR description is empty THEN
        RETURN "提交失败：必填信息不完整"
    END IF

    studentRecord = QueryStudentDormRecord(studentId)
    IF studentRecord does not exist OR studentRecord.dormNo != dormNo THEN
        RETURN "提交失败：学生与宿舍档案不匹配"
    END IF

    validAttachmentCount = 0
    FOR each file IN attachments DO
        IF file.type NOT IN {jpg, png, jpeg} OR file.size > 5MB THEN
            RETURN "提交失败：附件格式或大小不符合要求"
        ELSE
            validAttachmentCount = validAttachmentCount + 1
        END IF
    END FOR

    IF validAttachmentCount > 5 THEN
        RETURN "提交失败：附件数量超过限制"
    END IF

    IF faultType IN {"漏水", "断电", "门锁损坏"} THEN
        priority = "高"
    ELSE
        priority = "普通"
    END IF

    orderId = GenerateRepairOrder(studentId, dormNo, faultType, description, priority, attachments)
    NotifyAdmin(orderId)
    RETURN "提交成功：" + orderId
END PROCEDURE"""


WHITE_BOX_ROWS = [
    ["覆盖标准", "测试用例", "覆盖内容", "预期结果"],
    ["语句覆盖", "studentId=S001，dormNo=3-502，faultType=漏水，description=卫生间漏水，attachments=[a.jpg]", "执行完整成功路径，覆盖输入校验、档案查询、附件循环、紧急判断、生成工单、通知管理员", "返回“提交成功：工单号”，工单优先级为高"],
    ["判定-条件覆盖", "TC1：description为空；TC2：studentId=S001但dormNo=3-601；TC3：附件类型为txt；TC4：faultType=灯管损坏且附件合法", "使必填项判定、档案匹配判定、附件合法判定、紧急故障判定的真假结果均至少出现一次", "TC1提示信息不完整；TC2提示档案不匹配；TC3提示附件不合法；TC4提交成功且优先级普通"],
    ["条件组合覆盖", "必填项组合：学生号空/不空、宿舍号空/不空、类型空/不空、描述空/不空；附件组合：合法类型且≤5MB、非法类型、超5MB；故障类型组合：紧急/非紧急", "覆盖关键复合条件的有效与无效组合，重点验证 OR 条件下任一必填项为空均拒绝提交", "无效组合均被拒绝并给出对应提示；全部有效时进入后续处理"],
    ["独立路径覆盖", "P1：信息不完整；P2：档案不匹配；P3：附件不合法；P4：附件数量超过5；P5：紧急故障成功；P6：普通故障成功", "覆盖控制流图的主要独立路径，包括异常返回路径和两条成功路径", "各路径返回与路径条件一致的结果，成功路径生成工单并通知管理员"],
]


EQUIVALENCE_ROWS = [
    ["输入项", "有效等价类", "无效等价类", "代表用例及预期结果"],
    ["学生号", "已存在且处于正常住宿状态的学生号", "空值、不存在、已退宿学生号", "S001 有效；空值返回必填信息不完整；S999 返回档案不匹配"],
    ["宿舍号", "与学生档案一致的宿舍号", "空值、格式错误、与学生不匹配", "3-502 有效；A502 格式错误；3-601 档案不匹配"],
    ["故障类型", "系统允许的类型，如漏水、断电、门锁损坏、灯管损坏", "空值、未定义类型", "漏水成功进入高优先级；空值拒绝提交"],
    ["故障描述", "5 至 200 字的明确描述", "空值、少于5字、超过200字", "“卫生间水管持续漏水”有效；“坏了”提示描述过短"],
    ["附件", "0 至 5 个 jpg/png/jpeg 且单个≤5MB", "超过5个、类型非法、单个超过5MB", "1个jpg有效；6个附件提示数量超限；txt提示类型非法"],
]


BOUNDARY_ROWS = [
    ["边界对象", "一般性边界测试", "健壮性边界测试", "预期结果"],
    ["故障描述长度（5-200字）", "5、6、199、200字", "4、201字", "边界内接受，低于5或高于200拒绝"],
    ["附件数量（0-5个）", "0、1、4、5个", "6个", "0至5个接受，6个拒绝"],
    ["附件大小（≤5MB）", "4.99MB、5MB", "5.01MB", "不超过5MB接受，超过5MB拒绝"],
    ["宿舍号格式", "1-101、12-1208", "空值、1#101、A-101", "符合楼栋-房间格式接受，其余拒绝或提示格式错误"],
]


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(text) <= 20 else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(10.5)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_table(doc: Document, rows: Sequence[Sequence[str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            set_cell_text(table.cell(i, j), value, bold=(i == 0))
            if i == 0:
                shade_cell(table.cell(i, j), "D9EAF7")


def set_doc_style(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.45)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.7)
    section.right_margin = Cm(2.5)

    styles = doc.styles
    styles["Normal"].font.name = "宋体"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(12)
    styles["Normal"].paragraph_format.line_spacing = 1.5

    for style_name, size in [("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)]:
        style = styles[style_name]
        style.font.name = "黑体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("- PAGE -")
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(10.5)


def add_paragraph(doc: Document, text: str, first_line: bool = True, size: float = 12, bold: bool = False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    if first_line:
        p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)


def add_centered_run(doc: Document, text: str, size: float, font_name: str = "宋体", bold: bool = False, spacing_after: float = 0) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(spacing_after)
    r = p.add_run(text)
    r.bold = bold
    r.font.name = font_name
    r._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    r.font.size = Pt(size)


def add_field_line(doc: Document, label: str, value: str = "", size: float = 16) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.8
    display = f"{label}{value}" if value else f"{label}{'_' * 22}"
    r = p.add_run(display)
    r.font.name = "宋体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    r.font.size = Pt(size)


def add_section_label(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.name = "黑体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    r.font.size = Pt(14)


def add_code_block(doc: Document, code: str) -> None:
    for line in code.splitlines():
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.line_spacing = 1.1
        run = p.add_run(line)
        run.font.name = "Courier New"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(9.5)


def add_figure(doc: Document, title: str, path: Path, width_cm: float = 15.5) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    cap = doc.add_paragraph(title)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in cap.runs:
        r.font.name = "宋体"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        r.font.size = Pt(10.5)


def add_cover(doc: Document) -> None:
    for _ in range(1):
        doc.add_paragraph()
    add_centered_run(doc, "Sichuan Top Vocational College of Information Technology", 14, "Times New Roman", False, 8)
    for _ in range(2):
        doc.add_paragraph()
    add_centered_run(doc, "期  末  考  查  报  告", 26, "黑体", True, 18)
    add_centered_run(doc, "（《软件工程导论》）", 18, "宋体", False, 30)

    add_field_line(doc, "姓    名： ")
    add_field_line(doc, "学    号： ")
    add_field_line(doc, "系    别： ", "_信息工程学院_")
    add_field_line(doc, "专    业： ", "_软件技术_    _")
    add_field_line(doc, "年    级： ", "___2024级_____")
    add_field_line(doc, "班    级： ", "___ ___________")
    add_field_line(doc, "指导教师： ", "___项叙淋_ ____")

    for _ in range(2):
        doc.add_paragraph()
    add_centered_run(doc, "2026 年    月    日 至   2026 年   月   日", 14, "宋体")
    add_centered_run(doc, "所 在 单 位 ：  __         _级    系    专业    班", 12, "宋体")
    doc.add_page_break()


def add_main_report(doc: Document, images: dict[str, Path]) -> None:
    add_section_label(doc, "项目名称：")
    add_paragraph(doc, "宿舍报修管理系统分析与设计", first_line=False, size=12)

    add_section_label(doc, "项目目的：")
    for idx, para in enumerate(PROJECT_PURPOSE, 1):
        add_paragraph(doc, f"{idx}、{para}", first_line=False, size=12)

    add_section_label(doc, "项目内容及要求：")
    add_paragraph(doc, "说明：", first_line=False, size=12, bold=True)
    for idx, para in enumerate(TEMPLATE_REQUIREMENTS, 1):
        add_paragraph(doc, f"{idx}、{para}", first_line=False, size=12)

    doc.add_heading("一、需求分析", level=1)
    for para in REPORT_SECTIONS[0][1]:
        add_paragraph(doc, para)
    doc.add_heading("1. E-R 图或 UML 类图", level=2)
    add_paragraph(doc, "本报告选择绘制 UML 类图。类图包含学生、宿舍、报修工单、维修员、管理员、派工记录六个主要类，能够体现系统的静态结构。")
    add_figure(doc, "图 1 宿舍报修管理系统 UML 类图", images["类图"], 15.8)
    doc.add_heading("2. 分层数据流图（DFD）", level=2)
    add_paragraph(doc, "顶层 DFD 将整个系统抽象为一个加工 P0，外部实体包括学生、管理员、维修员和后勤管理部门。")
    add_figure(doc, "图 2 顶层 DFD", images["顶层DFD"], 15.8)
    add_paragraph(doc, "0 层 DFD 将 P0 分解为报修申请受理、工单审核、派工处理、维修反馈确认、统计查询五个加工，并设置学生/宿舍档案、报修工单库、派工/维修记录三个数据存储。")
    add_figure(doc, "图 3 0 层 DFD", images["0层DFD"], 16.0)
    doc.add_heading("3. 核心业务实体状态图", level=2)
    add_paragraph(doc, "核心业务实体选择“报修工单”。工单状态从待提交开始，经待审核、待派工、维修中、待确认，最终进入已完成状态。异常状态包括已驳回、已取消和返修。")
    add_figure(doc, "图 4 报修工单 UML 状态图", images["状态图"], 15.8)

    doc.add_heading("二、总体设计", level=1)
    for para in REPORT_SECTIONS[1][1]:
        add_paragraph(doc, para)
    add_figure(doc, "图 5 系统功能结构图", images["系统结构图"], 16.0)

    doc.add_heading("三、详细设计", level=1)
    for para in REPORT_SECTIONS[2][1]:
        add_paragraph(doc, para)
    add_figure(doc, "图 6 报修申请受理模块程序流程图", images["流程图"], 13.5)
    add_figure(doc, "图 7 派工处理模块 N-S 图", images["NS图"], 14.5)

    doc.add_heading("四、程序编码", level=1)
    for para in REPORT_SECTIONS[3][1]:
        add_paragraph(doc, para)
    add_code_block(doc, PDL)

    doc.add_heading("五、软件测试", level=1)
    for para in REPORT_SECTIONS[4][1]:
        add_paragraph(doc, para)
    doc.add_heading("1. 白盒测试", level=2)
    add_paragraph(doc, "白盒测试以报修申请受理模块的 PDL 伪代码为依据。该模块的关键判定包括必填信息是否完整、学生宿舍档案是否匹配、附件是否合法、附件数量是否超限以及故障类型是否紧急。")
    add_table(doc, WHITE_BOX_ROWS)
    doc.add_heading("2. 黑盒测试", level=2)
    add_paragraph(doc, "黑盒测试从输入域出发，不依赖程序内部结构。等价类划分如下表所示。")
    add_table(doc, EQUIVALENCE_ROWS)
    add_paragraph(doc, "边界值分析包括一般性测试和健壮性测试，如下表所示。")
    add_table(doc, BOUNDARY_ROWS)
    doc.add_heading("3. 用例图设计", level=2)
    add_paragraph(doc, "系统用例图包含学生、管理员、维修员、后勤部门四类参与者，覆盖报修、查询、审核、派工、维修反馈、统计导出等核心用例。")
    add_figure(doc, "图 8 宿舍报修管理系统 UML 用例图", images["用例图"], 16.0)

    doc.add_heading("结论", level=1)
    add_paragraph(doc, "宿舍报修管理系统围绕报修工单形成完整业务闭环。需求分析中的类图、DFD 和状态图明确了系统对象、数据流和状态变化；总体设计将 0 层 DFD 转换为层次化功能结构；详细设计、PDL 伪代码和测试用例均围绕报修申请受理与派工处理展开，保证了前后设计内容的一致性。系统能够满足学生便捷报修、管理员规范派工、维修员及时反馈和后勤部门统计管理的基本需求。")

    add_section_label(doc, "个人总结：")
    add_paragraph(doc, "通过本次宿舍报修管理系统的分析与设计，我进一步理解了软件工程各阶段之间的衔接关系。需求分析阶段明确系统边界和数据流，总体设计阶段完成模块划分，详细设计阶段细化关键模块逻辑，编码与测试阶段则验证设计是否可实现、是否前后一致。本次设计也让我认识到，图形规范、数据命名和测试用例之间必须保持一致，否则容易出现设计与实现脱节的问题。", first_line=True)
    doc.add_paragraph()
    add_paragraph(doc, "学生签名：____________________", first_line=False)
    add_paragraph(doc, "指导教师签字（签章）：____________________", first_line=False)


def write_markdown(images: dict[str, Path]) -> None:
    lines = [
        "# 《软件工程导论》期末考查报告",
        "",
        "题目：宿舍报修管理系统分析与设计",
        "",
        "> 封面信息（姓名、学号、班级、专业、指导教师）需由本人填写。",
        "",
    ]
    for title, paras in REPORT_SECTIONS:
        lines.append(f"## {title}")
        lines.extend(["", *paras, ""])
    lines.extend(["## 图表清单", ""])
    for name, path in images.items():
        lines.append(f"- {name}: `{path.relative_to(ROOT)}`")
    lines.extend(["", "## PDL 伪代码", "", "```text", PDL, "```", ""])
    lines.extend(["## 白盒测试", ""])
    for row in WHITE_BOX_ROWS:
        lines.append(" | ".join(row))
    lines.extend(["", "## 黑盒测试：等价类划分", ""])
    for row in EQUIVALENCE_ROWS:
        lines.append(" | ".join(row))
    lines.extend(["", "## 黑盒测试：边界值分析", ""])
    for row in BOUNDARY_ROWS:
        lines.append(" | ".join(row))
    MD_PATH.write_text("\n".join(lines), encoding="utf-8")


def build_docx(images: dict[str, Path]) -> None:
    doc = Document()
    set_doc_style(doc)
    add_cover(doc)
    add_main_report(doc, images)
    doc.save(DOCX_PATH)
    shutil.copyfile(DOCX_PATH, ASCII_DOCX_PATH)


def main() -> None:
    ensure_dirs()
    images = draw_all_diagrams()
    write_markdown(images)
    build_docx(images)
    print(f"Wrote {DOCX_PATH}")
    print(f"Wrote {ASCII_DOCX_PATH}")
    print(f"Wrote {MD_PATH}")
    print(f"Wrote {len(images)} diagram images under {IMG_DIR}")


if __name__ == "__main__":
    main()
